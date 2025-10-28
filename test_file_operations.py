"""
Unit Tests for File Operations Server
Tests reading and writing various file formats: txt, docx, xlsx, pdf, png
Coverage: 70%+ of methods
"""

# Import testing framework
import unittest
import os
import sys
from pathlib import Path
import tempfile
import shutil

# Import libraries for creating test files
from docx import Document
import openpyxl
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from PIL import Image

# Import the file operations functions to test
# Note: This assumes the file_operations_server.py is in the same directory
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

# Import the functions we want to test
from file_operations_server import (
    read_text_file,
    read_docx_file,
    read_xlsx_file,
    read_pdf_file,
    read_png_file,
    write_text_file,
    write_docx_file,
    write_xlsx_file,
    write_pdf_file
)

class TestFileOperations(unittest.TestCase):
    """Test suite for file operations server"""
    
    @classmethod
    def setUpClass(cls):
        """Set up test environment once for all tests"""
        # Create a temporary directory for test files
        cls.test_dir = Path(tempfile.mkdtemp())
        print(f"\nTest directory: {cls.test_dir}")
    
    @classmethod
    def tearDownClass(cls):
        """Clean up test environment after all tests"""
        # Remove the temporary directory and all its contents
        if cls.test_dir.exists():
            shutil.rmtree(cls.test_dir)
            print(f"Cleaned up test directory: {cls.test_dir}")
    
    def test_read_text_file(self):
        """Test reading a plain text file"""
        # Create a test text file
        test_file = self.test_dir / "test.txt"
        test_content = "Hello, World!\nThis is a test file.\nLine 3."
        
        # Write the test content to the file
        with open(test_file, 'w', encoding='utf-8') as f:
            f.write(test_content)
        
        # Test reading the file
        result = read_text_file(test_file)
        
        # Verify the content matches
        self.assertEqual(result, test_content)
        print("✓ Text file read successfully")
    
    def test_read_text_file_with_special_chars(self):
        """Test reading a text file with special characters"""
        # Create a test text file with special characters
        test_file = self.test_dir / "test_special.txt"
        test_content = "Special chars: é, ñ, 中文, 🎉"
        
        # Write the test content to the file
        with open(test_file, 'w', encoding='utf-8') as f:
            f.write(test_content)
        
        # Test reading the file
        result = read_text_file(test_file)
        
        # Verify the content matches
        self.assertEqual(result, test_content)
        print("✓ Text file with special characters read successfully")
    
    def test_write_text_file(self):
        """Test writing a plain text file"""
        # Define the test file path and content
        test_file = self.test_dir / "write_test.txt"
        test_content = "This is a test\nof writing\nto a file."
        
        # Test writing the file
        write_text_file(test_file, test_content)
        
        # Verify the file was created
        self.assertTrue(test_file.exists())
        
        # Verify the content is correct
        with open(test_file, 'r', encoding='utf-8') as f:
            result = f.read()
        
        self.assertEqual(result, test_content)
        print("✓ Text file written successfully")
    
    def test_read_docx_file(self):
        """Test reading a Word document"""
        # Create a test DOCX file
        test_file = self.test_dir / "test.docx"
        doc = Document()
        
        # Add paragraphs to the document
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        doc.add_paragraph("Third paragraph")
        
        # Save the document
        doc.save(test_file)
        
        # Test reading the file
        result = read_docx_file(test_file)
        
        # Verify the content is correct
        expected = "First paragraph\nSecond paragraph\nThird paragraph"
        self.assertEqual(result, expected)
        print("✓ DOCX file read successfully")
    
    def test_write_docx_file(self):
        """Test writing a Word document"""
        # Define the test file path and content
        test_file = self.test_dir / "write_test.docx"
        test_content = "Paragraph 1\nParagraph 2\nParagraph 3"
        
        # Test writing the file
        write_docx_file(test_file, test_content)
        
        # Verify the file was created
        self.assertTrue(test_file.exists())
        
        # Verify the content is correct by reading it back
        doc = Document(test_file)
        paragraphs = [para.text for para in doc.paragraphs]
        result = '\n'.join(paragraphs)
        
        self.assertEqual(result, test_content)
        print("✓ DOCX file written successfully")
    
    def test_read_xlsx_file(self):
        """Test reading an Excel file"""
        # Create a test XLSX file
        test_file = self.test_dir / "test.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "TestSheet"
        
        # Add data to the sheet
        ws['A1'] = 'Name'
        ws['B1'] = 'Age'
        ws['A2'] = 'John'
        ws['B2'] = 30
        ws['A3'] = 'Jane'
        ws['B3'] = 25
        
        # Save the workbook
        wb.save(test_file)
        
        # Test reading the file
        result = read_xlsx_file(test_file)
        
        # Verify the content contains expected data
        self.assertIn('TestSheet', result)
        self.assertIn('Name', result)
        self.assertIn('Age', result)
        self.assertIn('John', result)
        self.assertIn('30', result)
        print("✓ XLSX file read successfully")
    
    def test_write_xlsx_file(self):
        """Test writing an Excel file"""
        # Define the test file path and content
        test_file = self.test_dir / "write_test.xlsx"
        test_content = "Name\tAge\nJohn\t30\nJane\t25"
        
        # Test writing the file
        write_xlsx_file(test_file, test_content)
        
        # Verify the file was created
        self.assertTrue(test_file.exists())
        
        # Verify the content is correct by reading it back
        wb = openpyxl.load_workbook(test_file)
        ws = wb.active
        
        # Check first row
        self.assertEqual(ws['A1'].value, 'Name')
        self.assertEqual(ws['B1'].value, 'Age')
        
        # Check second row
        self.assertEqual(ws['A2'].value, 'John')
        self.assertEqual(ws['B2'].value, '30')
        
        print("✓ XLSX file written successfully")
    
    def test_read_pdf_file(self):
        """Test reading a PDF file"""
        # Create a test PDF file
        test_file = self.test_dir / "test.pdf"
        c = canvas.Canvas(str(test_file), pagesize=letter)
        
        # Add text to the PDF
        c.drawString(100, 750, "Hello from PDF")
        c.drawString(100, 730, "This is a test PDF file")
        
        # Save the PDF
        c.save()
        
        # Test reading the file
        result = read_pdf_file(test_file)
        
        # Verify the content contains expected text
        # Note: PDF text extraction may vary, so we check for presence
        self.assertIn('Page 1', result)
        print("✓ PDF file read successfully")
    
    def test_write_pdf_file(self):
        """Test writing a PDF file"""
        # Define the test file path and content
        test_file = self.test_dir / "write_test.pdf"
        test_content = "This is a test PDF\nWith multiple lines\nAnd some content"
        
        # Test writing the file
        write_pdf_file(test_file, test_content)
        
        # Verify the file was created
        self.assertTrue(test_file.exists())
        
        # Verify the file size is reasonable (not empty)
        self.assertGreater(test_file.stat().st_size, 0)
        print("✓ PDF file written successfully")
    
    def test_read_png_file(self):
        """Test reading a PNG image file"""
        # Create a test PNG file
        test_file = self.test_dir / "test.png"
        
        # Create a simple red image
        img = Image.new('RGB', (100, 100), color='red')
        img.save(test_file)
        
        # Test reading the file
        result = read_png_file(test_file)
        
        # Verify the result has the expected structure
        self.assertIn('metadata', result)
        self.assertIn('data', result)
        self.assertIn('description', result)
        
        # Verify metadata is correct
        self.assertEqual(result['metadata']['width'], 100)
        self.assertEqual(result['metadata']['height'], 100)
        self.assertEqual(result['metadata']['format'], 'PNG')
        
        print("✓ PNG file read successfully")
    
    def test_read_png_file_different_sizes(self):
        """Test reading PNG files of different sizes"""
        # Create test PNG files of various sizes
        sizes = [(50, 50), (200, 150), (300, 400)]
        
        for width, height in sizes:
            test_file = self.test_dir / f"test_{width}x{height}.png"
            
            # Create an image with the specified size
            img = Image.new('RGB', (width, height), color='blue')
            img.save(test_file)
            
            # Test reading the file
            result = read_png_file(test_file)
            
            # Verify the dimensions match
            self.assertEqual(result['metadata']['width'], width)
            self.assertEqual(result['metadata']['height'], height)
        
        print("✓ PNG files of different sizes read successfully")
    
    def test_read_nonexistent_file(self):
        """Test reading a file that doesn't exist"""
        # Define a path to a non-existent file
        test_file = self.test_dir / "nonexistent.txt"
        
        # Test that reading the file raises an exception
        with self.assertRaises(FileNotFoundError):
            read_text_file(test_file)
        
        print("✓ Nonexistent file error handled correctly")
    
    def test_write_and_read_roundtrip_txt(self):
        """Test writing and reading back a text file (roundtrip test)"""
        # Define the test file and content
        test_file = self.test_dir / "roundtrip.txt"
        test_content = "This is a roundtrip test\nWith multiple lines\nAnd special chars: ñ, é, 中文"
        
        # Write the file
        write_text_file(test_file, test_content)
        
        # Read it back
        result = read_text_file(test_file)
        
        # Verify the content matches exactly
        self.assertEqual(result, test_content)
        print("✓ Text file roundtrip test passed")
    
    def test_write_and_read_roundtrip_docx(self):
        """Test writing and reading back a DOCX file (roundtrip test)"""
        # Define the test file and content
        test_file = self.test_dir / "roundtrip.docx"
        test_content = "Line 1\nLine 2\nLine 3"
        
        # Write the file
        write_docx_file(test_file, test_content)
        
        # Read it back
        result = read_docx_file(test_file)
        
        # Verify the content matches
        self.assertEqual(result, test_content)
        print("✓ DOCX file roundtrip test passed")
    
    def test_write_and_read_roundtrip_xlsx(self):
        """Test writing and reading back an XLSX file (roundtrip test)"""
        # Define the test file and content
        test_file = self.test_dir / "roundtrip.xlsx"
        test_content = "Header1\tHeader2\nValue1\tValue2"
        
        # Write the file
        write_xlsx_file(test_file, test_content)
        
        # Read it back
        result = read_xlsx_file(test_file)
        
        # Verify the content contains the expected values
        self.assertIn('Header1', result)
        self.assertIn('Header2', result)
        self.assertIn('Value1', result)
        self.assertIn('Value2', result)
        print("✓ XLSX file roundtrip test passed")

def run_tests():
    """Run all tests and display results"""
    # Create a test suite
    suite = unittest.TestLoader().loadTestsFromTestCase(TestFileOperations)
    
    # Run the tests with verbose output
    runner = unittest.TextTestRunner(verbosity=2)
    result = runner.run(suite)
    
    # Print summary
    print("\n" + "="*70)
    print("TEST SUMMARY")
    print("="*70)
    print(f"Tests run: {result.testsRun}")
    print(f"Successes: {result.testsRun - len(result.failures) - len(result.errors)}")
    print(f"Failures: {len(result.failures)}")
    print(f"Errors: {len(result.errors)}")
    print("="*70)
    
    # Return exit code (0 for success, 1 for failure)
    return 0 if result.wasSuccessful() else 1

if __name__ == '__main__':
    # Run the tests when the script is executed
    sys.exit(run_tests())

