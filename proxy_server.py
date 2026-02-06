#!/usr/bin/env python3
"""
Python FastAPI replacement for the Node.js proxy server.
Provides the same functionality but in Python for better integration with the MCP ecosystem.
"""

import asyncio
import json
import os
import re
import time
import base64
from typing import Dict, List, Optional, Any
from pathlib import Path
from datetime import datetime
from io import BytesIO
from urllib.parse import urlparse, urlunparse

import httpx
from fastapi import FastAPI, HTTPException, Request
from fastapi.exceptions import RequestValidationError
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from starlette.middleware.base import BaseHTTPMiddleware
from pydantic import BaseModel
import uvicorn

# Import dotenv to load .env file
try:
    from dotenv import load_dotenv
    DOTENV_AVAILABLE = True
except ImportError:
    DOTENV_AVAILABLE = False
    print("[WARN] python-dotenv not available. Install with: pip install python-dotenv")

# Import file operations libraries
try:
    from docx import Document  # python-docx for Word documents
    import openpyxl  # openpyxl for Excel files
    from openpyxl.styles import Font, Alignment  # For Excel formatting
    import PyPDF2  # PyPDF2 for PDF reading
    from PIL import Image  # Pillow for image operations
    FILE_OPS_AVAILABLE = True
    print("[OK] File operations libraries loaded successfully")
except ImportError as e:
    print(f"[WARN] File operations libraries not available: {e}")
    FILE_OPS_AVAILABLE = False

# Import AutoGen components for team-based chat
try:
    from autogen_agentchat.teams import SelectorGroupChat
    from autogen_core import Component, FunctionCall, ComponentLoader
    AUTOGEN_AVAILABLE = True
    print("[OK] AutoGen imports successful")
except ImportError as e:
    print(f"[WARN] AutoGen not available: {e}")
    AUTOGEN_AVAILABLE = False
    SelectorGroupChat = None
    Component = None
    ComponentLoader = None

# Import MCP SDK (with error handling)
try:
    from mcp import ClientSession, stdio_client, StdioServerParameters
    MCP_AVAILABLE = True
except ImportError as e:
    print(f"MCP import error: {e}")
    MCP_AVAILABLE = False
    ClientSession = None
    stdio_client = None
    StdioServerParameters = None

# Import memory system (with error handling)
MEMORY_AVAILABLE = False
MemoryManager = None
memory_import_error = None

try:
    # Check for required dependencies first
    try:
        import numpy
    except ImportError:
        raise ImportError("numpy is required for the memory system. Install with: pip install numpy")
    
    from memory import MemoryManager
    MEMORY_AVAILABLE = True
    print("[OK] Memory system imports successful")
except ImportError as e:
    memory_import_error = str(e)
    print(f"[WARN] Memory system not available: {e}")
    if "numpy" in str(e).lower():
        print("   Install numpy with: pip install numpy")
    MEMORY_AVAILABLE = False
    MemoryManager = None
except Exception as e:
    memory_import_error = str(e)
    print(f"[WARN] Memory system not available (unexpected error): {e}")
    MEMORY_AVAILABLE = False
    MemoryManager = None

# Import philosopher mode
try:
    from philosopher_mode import PhilosopherMode
    PHILOSOPHER_MODE_AVAILABLE = True
    print("[OK] Philosopher mode imports successful")
except ImportError as e:
    print(f"[WARN] Philosopher mode not available: {e}")
    PHILOSOPHER_MODE_AVAILABLE = False
    PhilosopherMode = None
except Exception as e:
    print(f"[WARN] Philosopher mode not available (unexpected error): {e}")
    PHILOSOPHER_MODE_AVAILABLE = False
    PhilosopherMode = None

# Load environment variables from .env file in project root
if DOTENV_AVAILABLE:
    # Try loading from current directory first
    env_path = Path(__file__).parent / '.env'
    if env_path.exists():
        load_dotenv(env_path)
        print(f"✅ Loaded environment variables from {env_path}")
    else:
        # Try loading from parent directory as fallback
        parent_env = Path(__file__).parent.parent / '.env'
        if parent_env.exists():
            load_dotenv(parent_env)
            print(f"✅ Loaded environment variables from {parent_env}")
        else:
            print(f"⚠️  No .env file found. Using system environment variables.")
            print(f"   Looked in: {env_path} and {parent_env}")
else:
    print("⚠️  python-dotenv not available. Using system environment variables only.")

# Pydantic models for request/response validation
class ServerConfig(BaseModel):
    id: Optional[str] = None
    name: Optional[str] = None
    command: Optional[str] = None
    apiKey: Optional[str] = None
    model: Optional[str] = None
    url: Optional[str] = None
    wsUrl: Optional[str] = None
    status: Optional[str] = None
    enabled: Optional[bool] = None
    action: Optional[str] = None

class ToolCallRequest(BaseModel):
    toolName: str
    parameters: Optional[Dict[str, Any]] = None

# Pydantic models for file operations
class ReadFileRequest(BaseModel):
    filename: str  # Name of the file to read

class WriteFileRequest(BaseModel):
    filename: str  # Name of the file to write
    content: str  # Content to write to the file
    format: Optional[str] = "txt"  # File format (txt, docx, xlsx, pdf)

class FileResponse(BaseModel):
    success: bool  # Whether the operation was successful
    message: str  # Human-readable message
    data: Optional[Dict[str, Any]] = None  # Optional data payload

class TelegramChatMessage(BaseModel):
    role: str
    content: str


class TelegramChatRequest(BaseModel):
    message: str
    conversation_id: Optional[str] = None
    user_id: Optional[str] = None
    history: Optional[List[TelegramChatMessage]] = None
    system_prompt: Optional[str] = None
    model: Optional[str] = None
    temperature: Optional[float] = None
    max_output_tokens: Optional[int] = None


class TelegramChatResponse(BaseModel):
    reply: str
    conversation_id: str
    usage: Optional[Dict[str, Any]] = None

# Pydantic models for memory operations
class MemoryStoreRequest(BaseModel):
    text: str
    category: Optional[str] = None
    source: Optional[str] = None
    metadata: Optional[Dict[str, Any]] = None

class MemorySearchRequest(BaseModel):
    query: str
    limit: Optional[int] = None
    similarity_threshold: Optional[float] = None
    category: Optional[str] = None

class MemoryExtractRequest(BaseModel):
    messages: List[Dict[str, str]]
    max_memories: Optional[int] = 3

class MemoryResponse(BaseModel):
    success: bool
    message: str
    data: Optional[Dict[str, Any]] = None

# Pydantic models for philosopher mode
class PhilosopherStartRequest(BaseModel):
    conversation_id: Optional[str] = None
    user_id: Optional[str] = None

class PhilosopherStopRequest(BaseModel):
    conversation_id: Optional[str] = None
    user_id: Optional[str] = None

class PhilosopherContemplateRequest(BaseModel):
    conversation_id: Optional[str] = None
    user_id: Optional[str] = None
    question: Optional[str] = None  # Optional: if not provided, generate one

class PhilosopherResponse(BaseModel):
    success: bool
    message: str
    data: Optional[Dict[str, Any]] = None

# Global state (similar to the Node.js version)
mcp_clients = {}
mcp_servers = {}
SERVERS_FILE = Path(__file__).parent / "mcp_servers.json"
TEAM_CONFIG_FILE = Path(__file__).parent / "team-config.json"
SCRATCH_DIR = Path(__file__).parent / "scratch"

# Telegram chat session storage (simple in-memory cache)
telegram_conversations: Dict[str, List[Dict[str, str]]] = {}

# Philosopher mode state storage (per conversation)
philosopher_mode_active: Dict[str, bool] = {}
philosopher_mode_instances: Dict[str, Any] = {}

# Telegram/OpenAI configuration
TELEGRAM_DEFAULT_MODEL = os.getenv("TELEGRAM_MODEL", os.getenv("OPENAI_MODEL", "gpt-4o-mini"))
TELEGRAM_SYSTEM_PROMPT = os.getenv(
    "TELEGRAM_SYSTEM_PROMPT",
    "You are Eva, a helpful AI assistant that responds concisely for Telegram users.",
)
TELEGRAM_HISTORY_LIMIT = int(os.getenv("TELEGRAM_HISTORY_LIMIT", "12"))
TELEGRAM_CHAT_TIMEOUT = float(os.getenv("TELEGRAM_CHAT_TIMEOUT", "30"))
TELEGRAM_OPENAI_BASE_URL = os.getenv("TELEGRAM_OPENAI_BASE_URL") or os.getenv(
    "OPENAI_API_BASE", "https://api.openai.com/v1"
)
TELEGRAM_OPENAI_CHAT_PATH = os.getenv("TELEGRAM_OPENAI_CHAT_PATH", "/chat/completions")
OPENAI_ORG_ID = os.getenv("OPENAI_ORG_ID") or os.getenv("OPENAI_ORGANIZATION")
OPENAI_PROJECT_ID = os.getenv("OPENAI_PROJECT_ID")

# Helper utilities for Telegram integration
def build_openai_url(path: str) -> str:
    """Return absolute URL for OpenAI-compatible endpoints."""

    if path.startswith("http://") or path.startswith("https://"):
        return path

    base = TELEGRAM_OPENAI_BASE_URL.rstrip("/")
    if not path.startswith("/"):
        path = f"/{path}"
    return f"{base}{path}"


def trim_telegram_history(history: List[Dict[str, str]]) -> None:
    """Trim stored history to the configured limit (in-place)."""

    max_messages = max(TELEGRAM_HISTORY_LIMIT, 1) * 2
    if len(history) > max_messages:
        del history[: len(history) - max_messages]


# Create scratch directory if it doesn't exist
SCRATCH_DIR.mkdir(parents=True, exist_ok=True)

# Global AutoGen team instance
autogen_team = None

# Global memory manager instance
memory_manager = None

# Initialize memory manager if available
if MEMORY_AVAILABLE:
    try:
        memory_enabled = os.getenv("MEMORY_ENABLED", "true").lower() == "true"
        if memory_enabled:
            memory_manager = MemoryManager()
            print(f"✅ Memory system initialized with {memory_manager.count()} existing memories")
        else:
            print("⚠️  Memory system disabled via MEMORY_ENABLED=false")
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"⚠️  Failed to initialize memory system: {e}")
        print(f"   Full traceback:\n{error_trace}")
        memory_manager = None

# MCP Client Manager class to handle transport lifecycle
class MCPClientManager:
    """Manages MCP client and transport lifecycle."""

    def __init__(self, server_config: Dict[str, Any]):
        self.server_config = server_config
        self.client = None
        self.transport = None
        self.read = None
        self.write = None

    async def connect(self):
        """Connect to MCP server and initialize client."""
        if not MCP_AVAILABLE:
            raise Exception("MCP SDK not available")

        try:
            command_parts = self.server_config['command'].strip().split()
            if not command_parts:
                raise ValueError('Invalid command format')

            print(f"🔧 Creating MCP client with command: {command_parts[0]} and args: {command_parts[1:]}")

            # Prepare environment variables
            env = os.environ.copy()

            # Add API keys from server config if available
            if self.server_config.get('apiKey'):
                # Determine which API key to use based on the model
                model = self.server_config.get('model', '').lower()
                if 'gemini' in model:
                    env['GOOGLE_API_KEY'] = self.server_config['apiKey']
                    env['MCP_MODEL_PROVIDER'] = 'google'
                elif 'claude' in model:
                    env['ANTHROPIC_API_KEY'] = self.server_config['apiKey']
                    env['MCP_MODEL_PROVIDER'] = 'anthropic'
                else:
                    env['OPENAI_API_KEY'] = self.server_config['apiKey']
                    env['MCP_MODEL_PROVIDER'] = 'openai'

            # Add model configuration
            if self.server_config.get('model'):
                env['MCP_MODEL_NAME'] = self.server_config['model']

            # Set browser configuration for MCP server
            env.setdefault('BROWSER_USE_HEADLESS', 'true')
            env.setdefault('BROWSER_USE_DISABLE_SECURITY', 'false')

            print(f"🔐 Environment variables for MCP server: {list(env.keys())}")

            # Create server parameters
            server_params = StdioServerParameters(
                command=command_parts[0],
                args=command_parts[1:],
                env=env
            )

            # Create stdio client transport using async context manager
            try:
                transport_cm = stdio_client(server_params)
                async with transport_cm as (self.read, self.write):
                    # Create client session
                    self.client = ClientSession(self.read, self.write)

                    # Initialize the client
                    await self.client.initialize()

                    print("✅ MCP client setup complete")
                    return self.client
            except Exception as e:
                print(f"Failed to create stdio client: {e}")
                raise

        except Exception as e:
            print(f"MCP connection error: {e}")
            raise Exception(f"Failed to connect to MCP server: {str(e)}")

    async def disconnect(self):
        """Disconnect from MCP server."""
        if self.client:
            await self.client.close()
        if self.transport:
            # The transport context manager will handle cleanup
            pass

# FastAPI app
app = FastAPI(title="CATBot Proxy Server", version="2.0.0")

# Startup event to verify app initialization
@app.on_event("startup")
async def startup_event():
    """Log that the application has started successfully."""
    import sys
    print("🚀 FastAPI application startup event fired", flush=True)
    sys.stdout.flush()
    print(f"🚀 App routes registered: {len(app.routes)} routes", flush=True)
    sys.stdout.flush()
    # List all registered routes
    for route in app.routes:
        if hasattr(route, 'path') and hasattr(route, 'methods'):
            print(f"   Route: {list(route.methods)} {route.path}", flush=True)
            sys.stdout.flush()

# Request logging middleware to debug CORS issues
class RequestLoggingMiddleware(BaseHTTPMiddleware):
    """Middleware to log all incoming requests for debugging."""
    async def dispatch(self, request: Request, call_next):
        import sys
        # Safely log the incoming request with error handling
        try:
            method = getattr(request, 'method', 'UNKNOWN')
            path = getattr(request.url, 'path', 'unknown') if hasattr(request, 'url') else 'unknown'
            query = getattr(request.url, 'query', '') if hasattr(request, 'url') and hasattr(request.url, 'query') else ''
            origin = request.headers.get('origin', 'none') if hasattr(request, 'headers') else 'none'
            print(f"🌐 [{method}] {path}?{query}", flush=True)
            sys.stdout.flush()
            print(f"   Origin: {origin}", flush=True)
            sys.stdout.flush()
            if hasattr(request, 'headers'):
                print(f"   Headers: {dict(request.headers)}", flush=True)
                sys.stdout.flush()
        except Exception as log_error:
            # If logging fails, continue anyway - don't break the request
            print(f"⚠️ Error logging request: {log_error}", flush=True)
            sys.stdout.flush()
            import traceback
            print(traceback.format_exc(), flush=True)
            sys.stdout.flush()
        
        try:
            # Process the request
            response = await call_next(request)
            # Log successful response
            try:
                method = getattr(request, 'method', 'UNKNOWN')
                path = getattr(request.url, 'path', 'unknown') if hasattr(request, 'url') else 'unknown'
                status = getattr(response, 'status_code', 'unknown')
                print(f"✅ [{method}] {path} -> {status}", flush=True)
                sys.stdout.flush()
            except Exception as log_err:
                print(f"⚠️ Error logging response: {log_err}", flush=True)
                sys.stdout.flush()
            return response
        except Exception as e:
            # Log the exception
            try:
                method = getattr(request, 'method', 'UNKNOWN')
                path = getattr(request.url, 'path', 'unknown') if hasattr(request, 'url') else 'unknown'
                print(f"❌ [{method}] {path} -> Exception: {e}", flush=True)
                sys.stdout.flush()
                import traceback
                print(traceback.format_exc(), flush=True)
                sys.stdout.flush()
            except Exception:
                print("❌ Error in exception logging", flush=True)
                sys.stdout.flush()
            # Re-raise the exception so it can be handled by exception handlers
            raise

# Add request logging middleware first (before CORS)
app.add_middleware(RequestLoggingMiddleware)

# Add CORS middleware
# Allow specific origins for development and production
# Note: Using allow_credentials=False allows more flexible origin matching
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allow all origins for development (be more permissive to debug)
    allow_credentials=False,  # Set to False to allow more flexible CORS handling
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS", "PATCH"],
    allow_headers=["*"],
    expose_headers=["*"],
)

def build_cors_headers(request: Request) -> Dict[str, str]:
    """Build CORS headers for the request origin. Supports both localhost and remote access."""
    try:
        # Safely get origin from request headers, with fallback
        if hasattr(request, 'headers'):
            origin = request.headers.get("origin")
            # If no origin header (e.g., same-origin request), allow all origins for network access
            if not origin:
                origin = "*"
        else:
            origin = "*"
    except Exception:
        # If we can't access headers, allow all origins for network access
        origin = "*"
    
    return {
        "Access-Control-Allow-Origin": origin,
        "Access-Control-Allow-Methods": "GET, POST, PUT, DELETE, OPTIONS, PATCH",
        "Access-Control-Allow-Headers": "*",
    }

# Global exception handler to ensure CORS headers are always included
@app.exception_handler(HTTPException)
async def http_exception_handler(request: Request, exc: HTTPException):
    """Handle HTTP exceptions and ensure CORS headers are included."""
    # Safely build CORS headers - if this fails, use minimal headers
    try:
        cors_headers = build_cors_headers(request)
    except Exception as header_error:
        print(f"⚠️ Error building CORS headers in HTTPException handler: {header_error}")
        # Use minimal safe headers if build_cors_headers fails
        cors_headers = {
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Methods": "GET, POST, PUT, DELETE, OPTIONS, PATCH",
            "Access-Control-Allow-Headers": "*",
        }
    
    return JSONResponse(
        status_code=exc.status_code,
        content={"detail": exc.detail},
        headers=cors_headers,
    )

@app.exception_handler(RequestValidationError)
async def validation_exception_handler(request: Request, exc: RequestValidationError):
    """Handle validation errors and ensure CORS headers are included."""
    # Safely build CORS headers - if this fails, use minimal headers
    try:
        cors_headers = build_cors_headers(request)
    except Exception as header_error:
        print(f"⚠️ Error building CORS headers in ValidationException handler: {header_error}")
        # Use minimal safe headers if build_cors_headers fails
        cors_headers = {
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Methods": "GET, POST, PUT, DELETE, OPTIONS, PATCH",
            "Access-Control-Allow-Headers": "*",
        }
    
    return JSONResponse(
        status_code=422,
        content={"detail": exc.errors()},
        headers=cors_headers,
    )

@app.exception_handler(Exception)
async def general_exception_handler(request: Request, exc: Exception):
    """Handle all other exceptions and ensure CORS headers are included."""
    import sys
    import traceback
    error_trace = traceback.format_exc()
    print(f"❌ Unhandled exception in general_exception_handler: {exc}", flush=True)
    sys.stdout.flush()
    print(error_trace, flush=True)
    sys.stdout.flush()
    
    # Safely build CORS headers - if this fails, use minimal headers
    try:
        cors_headers = build_cors_headers(request)
    except Exception as header_error:
        print(f"⚠️ Error building CORS headers: {header_error}", flush=True)
        sys.stdout.flush()
        import traceback
        print(traceback.format_exc(), flush=True)
        sys.stdout.flush()
        # Use minimal safe headers if build_cors_headers fails
        cors_headers = {
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Methods": "GET, POST, PUT, DELETE, OPTIONS, PATCH",
            "Access-Control-Allow-Headers": "*",
        }
    
    try:
        response = JSONResponse(
            status_code=500,
            content={"detail": f"Internal server error: {str(exc)}"},
            headers=cors_headers,
        )
        print(f"✅ Created error response with status 500", flush=True)
        sys.stdout.flush()
        return response
    except Exception as response_error:
        print(f"❌ Error creating error response: {response_error}", flush=True)
        sys.stdout.flush()
        import traceback
        print(traceback.format_exc(), flush=True)
        sys.stdout.flush()
        # Last resort - return a simple response
        from fastapi.responses import PlainTextResponse
        return PlainTextResponse(
            content=f"Internal server error: {str(exc)}",
            status_code=500,
            headers=cors_headers,
        )

# Helper function to clean HTML text (similar to Node.js version)
def clean_text(text: str) -> str:
    """Clean HTML text by removing tags and decoding entities."""
    if not text:
        return ""

    # Remove HTML tags
    text = re.sub(r'</?[^>]+(>|$)', '', text)

    # Decode HTML entities
    html_entities = {
        '&amp;': '&', '&lt;': '<', '&gt;': '>', '&quot;': '"', '&#039;': "'",
        '&rsquo;': "'", '&lsquo;': "'", '&rdquo;': '"', '&ldquo;': '"',
        '&ndash;': '-', '&mdash;': '—'
    }

    for entity, replacement in html_entities.items():
        text = text.replace(entity, replacement)

    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

# Helper function to parse dates (similar to Node.js version)
def parse_date(date_str: Optional[str]) -> Optional[float]:
    """Parse date string to timestamp."""
    if not date_str:
        return None
    try:
        return datetime.fromisoformat(date_str.replace('Z', '+00:00')).timestamp()
    except (ValueError, AttributeError):
        return None

# Load servers from disk
def load_servers():
    """Load MCP servers from JSON file."""
    global mcp_servers
    try:
        if SERVERS_FILE.exists():
            with open(SERVERS_FILE, 'r', encoding='utf-8') as f:
                servers = json.load(f)
                mcp_servers = {server['id']: server for server in servers}
                print(f"Loaded {len(servers)} MCP servers from disk")
    except Exception as e:
        print(f"No existing servers file found, starting with empty state: {e}")

# Save servers to disk
def save_servers():
    """Save MCP servers to JSON file."""
    global mcp_servers
    try:
        servers = list(mcp_servers.values())
        with open(SERVERS_FILE, 'w', encoding='utf-8') as f:
            json.dump(servers, f, indent=2, ensure_ascii=False)
        print(f"Saved {len(servers)} MCP servers to disk")
    except Exception as e:
        print(f"Error saving servers to disk: {e}")

# Load AutoGen team from config
def load_autogen_team():
    """Load AutoGen team from team-config.json."""
    global autogen_team
    
    if not AUTOGEN_AVAILABLE:
        print("⚠️  AutoGen not available, skipping team load")
        return None
    
    try:
        if not TEAM_CONFIG_FILE.exists():
            print(f"⚠️  Team config file not found: {TEAM_CONFIG_FILE}")
            return None
            
        print(f"📂 Loading AutoGen team from {TEAM_CONFIG_FILE}...")
        
        with open(TEAM_CONFIG_FILE, 'r', encoding='utf-8') as f:
            team_config = json.load(f)
        
        # Load the team from the configuration using ComponentLoader
        loader = ComponentLoader()
        team = loader.load_component(team_config)
        
        print(f"✅ AutoGen team loaded successfully: {team_config.get('label', 'Unknown')}")
        return team
        
    except Exception as e:
        import traceback
        print(f"❌ Error loading AutoGen team: {e}")
        print(traceback.format_exc())
        return None

# Load servers on startup (with error handling to prevent startup failures)
try:
    load_servers()
    print(f"✅ Loaded {len(mcp_servers)} MCP servers from disk")
except Exception as e:
    import traceback
    print(f"⚠️ Warning: Could not load servers on startup: {e}")
    print(traceback.format_exc())
    # Continue anyway - server should still work without pre-loaded servers

# Load AutoGen team on startup (with error handling to prevent startup failures)
try:
    autogen_team = load_autogen_team()
    if autogen_team is not None:
        print("✅ AutoGen team loaded successfully on startup")
except Exception as e:
    import traceback
    print(f"⚠️ Warning: Could not load AutoGen team on startup: {e}")
    print(traceback.format_exc())
    # Continue anyway - server should still work without AutoGen team
    autogen_team = None

# Web proxy endpoint for fetching content
@app.get("/v1/proxy/fetch")
async def proxy_fetch(url: str):
    """Fetch web content through proxy."""
    if not url:
        raise HTTPException(status_code=400, detail="URL parameter is required")

    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
            'Connection': 'keep-alive',
            'Upgrade-Insecure-Requests': '1'
        }

        async with httpx.AsyncClient(timeout=10.0) as client:
            response = await client.get(url, headers=headers)

        return {"content": response.text}

    except Exception as e:
        print(f"Proxy error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch content: {str(e)}")

# Web search endpoint
@app.get("/v1/proxy/search")
async def proxy_search(query: str):
    """Search the web using Brave Search API or DuckDuckGo fallback."""
    if not query:
        raise HTTPException(status_code=400, detail="Search query is required")

    # Try Brave Search first
    brave_api_key = os.getenv('BRAVE_API_KEY')
    if not brave_api_key:
        print("⚠️  BRAVE_API_KEY not configured. Falling back to DuckDuckGo.")
    else:
        try:
            print(f"🔍 Using Brave Search API for query: {query[:50]}...")
            async with httpx.AsyncClient(timeout=15.0) as client:
                response = await client.get(
                    'https://api.search.brave.com/res/v1/web/search',
                    headers={
                        'Accept': 'application/json',
                        'Accept-Encoding': 'gzip',
                        'X-Subscription-Token': brave_api_key
                    },
                    params={
                        'q': query,
                        'count': 10,
                        'search_lang': 'en',
                        'safesearch': 'moderate',
                        'freshness': 'past_month'
                    }
                )

            if response.status_code == 200:
                data = response.json()
                if data.get('web', {}).get('results'):
                    results = []
                    for result in data['web']['results']:
                        # Parse date information
                        date_str = result.get('age') or result.get('published')
                        parsed_date = parse_date(date_str) if date_str else None

                        results.append({
                            'url': result['url'],
                            'title': clean_text(result.get('title', '')),
                            'snippet': clean_text(result.get('description', '')),
                            'date': parsed_date
                        })

                    # Sort by date (newest first) and filter valid results
                    results = [r for r in results if r['title'] and r['snippet']]
                    results.sort(key=lambda x: parse_date(x.get('date')) or 0, reverse=True)

                    print(f"✅ Brave Search returned {len(results)} results")
                    return {"results": results[:5], "source": "brave"}  # Return top 5 results
                else:
                    print(f"⚠️  Brave Search returned no results in response")
            elif response.status_code == 401:
                print(f"❌ Brave Search API authentication failed (401). Check your BRAVE_API_KEY.")
                raise HTTPException(
                    status_code=500,
                    detail="Brave Search API authentication failed. Please check your BRAVE_API_KEY configuration."
                )
            elif response.status_code == 429:
                print(f"⚠️  Brave Search API rate limit exceeded (429). Falling back to DuckDuckGo.")
            else:
                print(f"⚠️  Brave Search API returned status {response.status_code}. Falling back to DuckDuckGo.")
                try:
                    error_data = response.json()
                    print(f"   Error details: {error_data}")
                except:
                    print(f"   Error text: {response.text[:200]}")

        except httpx.RequestError as e:
            error_msg = str(e) if str(e) else f"Network error: {type(e).__name__}"
            print(f"❌ Brave Search network error: {error_msg}")
            print("   Falling back to DuckDuckGo...")
        except httpx.HTTPStatusError as e:
            print(f"❌ Brave Search HTTP error: {e.response.status_code if e.response else 'Unknown'}")
            print("   Falling back to DuckDuckGo...")
        except HTTPException:
            # Re-raise HTTPExceptions (like auth errors) - don't fall back
            raise
        except Exception as e:
            error_msg = str(e) if str(e) else f"Unknown error: {type(e).__name__}"
            print(f"❌ Brave Search failed: {error_msg}")
            print("   Falling back to DuckDuckGo...")
            import traceback
            print(traceback.format_exc())

    # Fallback to DuckDuckGo (only if Brave Search is not available or failed)
    print("🦆 Falling back to DuckDuckGo search...")
    try:
        search_url = f"https://html.duckduckgo.com/html/?q={query}"

        async with httpx.AsyncClient(timeout=15.0) as client:
            response = await client.get(
                search_url,
                headers={
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                    'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
                    'Accept-Language': 'en-US,en;q=0.5',
                    'Referer': 'https://duckduckgo.com/'
                },
                follow_redirects=True
            )
            
            # Check response status
            if response.status_code != 200:
                raise HTTPException(
                    status_code=500,
                    detail=f"DuckDuckGo search returned HTTP {response.status_code}. The search service may be temporarily unavailable."
                )

        results = []
        html = response.text

        # Extract results using multiple regex patterns (DuckDuckGo HTML structure can vary)
        patterns = [
            # Pattern 1: Modern DuckDuckGo structure
            r'<div class="links_main links_deep result__body">.*?<a class="result__a" href="([^"]+)".*?>(.*?)</a>.*?<a class="result__snippet".*?>(.*?)</a>',
            # Pattern 2: Alternative structure
            r'<div class="result__body">.*?<a class="result__url" href="([^"]+)".*?>(.*?)</a>.*?<div class="result__snippet">(.*?)</div>',
            # Pattern 3: Another variant
            r'<div class="result__body">.*?<a class="result__a" href="([^"]+)".*?>(.*?)</a>.*?<div class="result__snippet">(.*?)</div>',
            # Pattern 4: Try to find any links with result classes
            r'<a[^>]*class="[^"]*result[^"]*"[^>]*href="([^"]+)"[^>]*>(.*?)</a>.*?<div[^>]*class="[^"]*snippet[^"]*"[^>]*>(.*?)</div>',
            # Pattern 5: More generic pattern
            r'<a[^>]*href="([^"]+)"[^>]*class="[^"]*result__a[^"]*"[^>]*>(.*?)</a>.*?<span[^>]*class="[^"]*result__snippet[^"]*"[^>]*>(.*?)</span>'
        ]

        for pattern in patterns:
            matches = re.finditer(pattern, html, re.DOTALL)
            for match in matches:
                if len(results) >= 5:  # Limit to 5 results
                    break

                try:
                    url, title, snippet = match.groups()
                    url = url.replace('&amp;', '&')
                    
                    # Clean up the extracted text
                    title = clean_text(title)
                    snippet = clean_text(snippet)

                    # Skip if URL is invalid or is a DuckDuckGo internal link
                    if url and 'duckduckgo.com' not in url and title and snippet:
                        results.append({
                            'url': url,
                            'title': title,
                            'snippet': snippet
                        })
                except (ValueError, IndexError) as e:
                    # Skip malformed matches
                    continue

            if len(results) >= 5:
                break

        # If no results found with regex, try a simpler approach
        if len(results) == 0:
            # Log the HTML snippet for debugging (first 1000 chars)
            print(f"⚠️  DuckDuckGo search: No results found with regex patterns. HTML preview: {html[:1000]}")
            # Return empty results rather than failing
            return {"results": [], "source": "duckduckgo", "message": "No results found. DuckDuckGo HTML structure may have changed."}

        print(f"✅ DuckDuckGo returned {len(results)} results")
        return {"results": results, "source": "duckduckgo"}

    except httpx.RequestError as e:
        error_msg = str(e) if str(e) else f"Network error: {type(e).__name__}"
        print(f"Search error (network): {error_msg}")
        raise HTTPException(status_code=500, detail=f"Failed to perform search: Network error - {error_msg}")
    except httpx.HTTPStatusError as e:
        error_msg = f"HTTP {e.response.status_code}: {e.response.text[:200]}" if e.response else str(e)
        print(f"Search error (HTTP): {error_msg}")
        raise HTTPException(status_code=500, detail=f"Failed to perform search: {error_msg}")
    except Exception as e:
        error_msg = str(e) if str(e) else f"Unknown error: {type(e).__name__}"
        print(f"Search error: {error_msg}")
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Failed to perform search: {error_msg}")

# News API search endpoint
@app.get("/v1/proxy/news")
async def proxy_news_search(query: str):
    """Search news articles using News API."""
    if not query:
        raise HTTPException(status_code=400, detail="Search query is required")

    # Get News API key from environment variable
    news_api_key = os.getenv('NEWS_API_KEY')
    if not news_api_key:
        raise HTTPException(
            status_code=503,
            detail="NEWS_API_KEY is not configured. Please set it in your .env file."
        )

    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            response = await client.get(
                'https://newsapi.org/v2/everything',
                headers={
                    'Accept': 'application/json',
                },
                params={
                    'q': query,
                    'apiKey': news_api_key,
                    'sortBy': 'publishedAt',
                    'language': 'en',
                    'pageSize': 100  # Maximum allowed by News API
                }
            )

        if response.status_code == 200:
            data = response.json()
            articles = data.get('articles', [])
            
            if not articles:
                return {
                    "success": False,
                    "message": f"No articles found for search term \"{query}\"",
                    "articles": []
                }

            # Format articles for response
            formatted_articles = []
            for article in articles:
                formatted_articles.append({
                    'title': clean_text(article.get('title', '')),
                    'url': article.get('url', ''),
                    'description': clean_text(article.get('description', '')),
                    'publishedAt': article.get('publishedAt', ''),
                    'source': article.get('source', {}).get('name', 'Unknown')
                })

            return {
                "success": True,
                "message": f"Found {len(formatted_articles)} articles",
                "articles": formatted_articles,
                "totalResults": data.get('totalResults', len(formatted_articles))
            }
        else:
            error_data = response.json() if response.headers.get('content-type', '').startswith('application/json') else {}
            error_message = error_data.get('message', f"News API returned status {response.status_code}")
            print(f"News API error: {error_message}")
            raise HTTPException(
                status_code=response.status_code,
                detail=f"News API error: {error_message}"
            )

    except httpx.HTTPStatusError as e:
        print(f"News API HTTP error: {e}")
        raise HTTPException(
            status_code=e.response.status_code,
            detail=f"News API request failed: {str(e)}"
        )
    except Exception as e:
        print(f"News API error: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to fetch news: {str(e)}"
        )

# AutoGen team chat endpoint (integrated directly)
@app.post("/v1/proxy/autogen")
async def autogen_chat(request: Request):
    """Run AutoGen team conversation directly (no separate service needed)."""
    global autogen_team
    
    try:
        body = await request.json()
        input_text = body.get('input')
        
        if not input_text:
            raise HTTPException(status_code=400, detail="Input parameter is required")

        print(f"🤖 AutoGen team request: {input_text[:100]}...")

        # Check if AutoGen is available and team is loaded
        if not AUTOGEN_AVAILABLE:
            raise HTTPException(
                status_code=503,
                detail="AutoGen not available. Please install: pip install autogen-agentchat autogen-ext"
            )
        
        if autogen_team is None:
            # Try to load the team
            print("🔄 Loading AutoGen team for the first time...")
            autogen_team = load_autogen_team()
            if autogen_team is None:
                raise HTTPException(
                    status_code=503,
                    detail="AutoGen team not loaded. Check team-config.json exists and is valid."
                )
        else:
            # Check if the team config file has been modified and reload if needed
            try:
                config_mtime = TEAM_CONFIG_FILE.stat().st_mtime
                if not hasattr(autogen_team, '_config_mtime') or autogen_team._config_mtime != config_mtime:
                    print("🔄 Team config file changed, reloading AutoGen team...")
                    new_team = load_autogen_team()
                    if new_team is not None:
                        autogen_team = new_team
                        autogen_team._config_mtime = config_mtime
                    else:
                        print("⚠️  Failed to reload team, using existing team")
            except Exception as e:
                print(f"⚠️  Error checking team config modification time: {e}")

        try:
            # Run the team conversation
            print(f"🚀 Running AutoGen team with input: {input_text[:100]}...")
            result = await autogen_team.run(task=input_text)
            
            # Extract messages and final answer from result
            messages = []
            if hasattr(result, 'messages'):
                messages = [
                    {
                        "source": msg.source if hasattr(msg, 'source') else 'unknown',
                        "content": msg.content if hasattr(msg, 'content') else str(msg)
                    }
                    for msg in result.messages
                ]
            
            # Format all messages into a readable conversation summary
            conversation_summary = "=== AutoGen Team Workflow ===\n\n"
            
            if messages:
                for i, msg in enumerate(messages, 1):
                    source = msg.get('source', 'unknown')
                    content = msg.get('content', '')
                    conversation_summary += f"[{i}] {source}:\n{content}\n\n"
                
                conversation_summary += "=== End of Workflow ===\n\n"
                conversation_summary += "Please review the above conversation and provide a concise summary of the final result."
            else:
                conversation_summary = "No messages returned from AutoGen team."
            
            print(f"✅ AutoGen team completed with {len(messages)} messages")
            
            return {
                "output": conversation_summary,
                "response": conversation_summary,  # Alternative field name for compatibility
                "messages": messages,
                "message_count": len(messages)
            }
            
        except Exception as e:
            import traceback
            error_trace = traceback.format_exc()
            print(f"❌ AutoGen team execution error: {e}")
            print(f"   Traceback: {error_trace}")
            raise HTTPException(
                status_code=500,
                detail=f"AutoGen team execution failed: {str(e)}"
            )

    except HTTPException:
        raise
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"❌ AutoGen endpoint error: {e}")
        print(f"   Traceback: {error_trace}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to process AutoGen request: {str(e)}"
        )

# MCP server management endpoints
@app.post("/v1/mcp/servers")
async def manage_servers(server_config: ServerConfig):
    """Manage MCP servers (create, update, clear)."""
    print(f"Received server config: {server_config.dict()}")

    global mcp_clients, mcp_servers

    try:
        # Handle clear action
        if server_config.action == 'clear':
            # Disconnect all connected servers
            for server_id, client in mcp_clients.items():
                try:
                    await client.close()
                except Exception as e:
                    print(f"Error closing client {server_id}: {e}")

            # Clear all servers and clients
            mcp_servers.clear()
            mcp_clients.clear()

            print("Cleared all MCP servers and clients")

            # Save to disk
            save_servers()

            return {"message": "All MCP servers cleared successfully"}

        # Validate required fields
        if not server_config.id or not server_config.name:
            raise HTTPException(status_code=400, detail="Missing required fields: id, name")

        # Update existing server or add new one
        existing_server = mcp_servers.get(server_config.id)
        if existing_server:
            # Update existing server
            mcp_servers[server_config.id] = {
                **existing_server,
                **server_config.dict(),
                'status': existing_server.get('status', 'disconnected')  # Preserve connection status
            }
            print(f"Updated MCP server: {server_config.name} ({server_config.id})")
        else:
            # Add new server
            mcp_servers[server_config.id] = {
                **server_config.dict(),
                'status': 'disconnected'
            }
            print(f"Added MCP server: {server_config.name} ({server_config.id})")

        # Save to disk
        save_servers()

        return {"message": "Server saved successfully"}

    except HTTPException:
        raise
    except Exception as e:
        print(f"Error saving MCP server: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to save MCP server: {str(e)}")

@app.get("/v1/mcp/servers")
async def get_servers():
    """Get all MCP servers."""
    try:
        servers = list(mcp_servers.values())
        return {"servers": servers}
    except Exception as e:
        print(f"Error getting MCP servers: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to get MCP servers: {str(e)}")

def setup_browser_llm():
    """Set up the language model for browser automation."""
    model_provider = os.getenv("MCP_MODEL_PROVIDER", "google").lower()
    model_name = os.getenv("MCP_MODEL_NAME", "gemini-flash-latest")
    temperature = float(os.getenv("MCP_TEMPERATURE", "0.1"))

    if model_provider == "openai":
        from langchain_openai import ChatOpenAI
        return ChatOpenAI(
            model=model_name,
            temperature=temperature,
            api_key=os.getenv("OPENAI_API_KEY")
        )
    elif model_provider == "anthropic":
        from langchain_anthropic import ChatAnthropic
        return ChatAnthropic(
            model=model_name,
            temperature=temperature,
            api_key=os.getenv("ANTHROPIC_API_KEY")
        )
    elif model_provider == "google":
        from langchain_google_genai import ChatGoogleGenerativeAI
        return ChatGoogleGenerativeAI(
            model=model_name,
            temperature=temperature,
            api_key=os.getenv("GOOGLE_API_KEY")
        )
    else:
        # Default to Google/Gemini
        from langchain_google_genai import ChatGoogleGenerativeAI
        return ChatGoogleGenerativeAI(
            model=model_name,
            temperature=temperature,
            api_key=os.getenv("GOOGLE_API_KEY", "dummy-key")
        )

async def create_mcp_client(server_config: Dict[str, Any]):
    """Create MCP client connection (similar to Node.js version)."""
    manager = MCPClientManager(server_config)
    client = await manager.connect()
    return client

@app.post("/v1/mcp/servers/{server_id}/connect")
async def connect_server(server_id: str):
    """Connect to an MCP server."""
    try:
        print(f"Attempting to connect to server: {server_id}")

        # Check if server exists
        server = mcp_servers.get(server_id)
        if not server:
            print(f"Server not found: {server_id}")
            raise HTTPException(status_code=404, detail="Server not found")

        print(f"Found server: {server['name']} ({server_id})")

        # For MCP Browser Use server, just mark as connected since we handle it directly
        if "mcp-browser-use" in server.get("name", "").lower():
            server['status'] = 'connected'
            mcp_servers[server_id] = server
            print(f"Successfully connected to MCP Browser Use server: {server['name']}")
            return {"message": "Server connected successfully"}

        if server_id in mcp_clients:
            print(f"Server already connected: {server_id}")
            raise HTTPException(status_code=409, detail="Server is already connected")

        if not MCP_AVAILABLE:
            raise HTTPException(status_code=503, detail="MCP SDK not available")

        print(f"Creating MCP client for server: {server['name']}")
        client = await create_mcp_client(server)
        mcp_clients[server_id] = client

        server['status'] = 'connected'
        mcp_servers[server_id] = server

        print(f"Successfully connected to MCP server: {server['name']}")
        return {"message": "Server connected successfully"}

    except HTTPException:
        raise
    except Exception as e:
        print(f"Error connecting to MCP server: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to connect to MCP server: {str(e)}")

@app.post("/v1/mcp/servers/{server_id}/disconnect")
async def disconnect_server(server_id: str):
    """Disconnect from an MCP server."""
    try:
        global mcp_clients, mcp_servers

        # Check if this is the MCP Browser Use server
        server = mcp_servers.get(server_id)
        if server and "mcp-browser-use" in server.get("name", "").lower():
            # Just mark as disconnected since we don't have a real connection
            server['status'] = 'disconnected'
            mcp_servers[server_id] = server
            return {"message": "Server disconnected successfully"}

        if not MCP_AVAILABLE:
            raise HTTPException(status_code=503, detail="MCP SDK not available")

        client = mcp_clients.get(server_id)
        if not client:
            raise HTTPException(status_code=404, detail="Server is not connected")

        # Create a temporary manager to handle disconnect
        if server:
            manager = MCPClientManager(server)
            await manager.disconnect()

        del mcp_clients[server_id]

        server = mcp_servers.get(server_id)
        if server:
            server['status'] = 'disconnected'
            mcp_servers[server_id] = server

        return {"message": "Server disconnected successfully"}

    except HTTPException:
        raise
    except Exception as e:
        print(f"Error disconnecting MCP server: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to disconnect MCP server: {str(e)}")

# Browser automation tool endpoint (integrated directly)
@app.post("/v1/mcp/servers/{server_id}/tools/call")
async def call_tool(server_id: str, request: ToolCallRequest):
    """Call a tool on an MCP server - now handles browser automation directly."""
    if not MCP_AVAILABLE:
        raise HTTPException(status_code=503, detail="MCP SDK not available")

    try:
        print(f"🔧 [TOOLS/CALL] Server: {server_id}")

        client = mcp_clients.get(server_id)
        if not client:
            print(f"❌ [TOOLS/CALL] Server {server_id} not found or not connected")
            raise HTTPException(status_code=404, detail="Server is not connected")

        tool_name = request.toolName
        parameters = request.parameters or {}

        print(f"🔍 [TOOLS/CALL] Tool name: {tool_name}")
        print(f"🔍 [TOOLS/CALL] Parameters: {parameters}")

        if tool_name == "run_browser_agent":
            # Handle browser automation directly
            instruction = parameters.get("instruction", "")
            max_steps = parameters.get("max_steps", 10)
            use_vision = parameters.get("use_vision", True)

            try:
                # Import browser-use components (lazy imports)
                Agent = __import__('browser_use', fromlist=['Agent']).Agent
                BrowserSession = __import__('browser_use', fromlist=['BrowserSession']).BrowserSession

                # Set up browser configuration
                headless = os.getenv("BROWSER_USE_HEADLESS", "true").lower() == "true"
                disable_security = os.getenv("BROWSER_USE_DISABLE_SECURITY", "false").lower() == "true"

                # Create browser session
                browser = BrowserSession(
                    headless=headless,
                    disable_security=disable_security
                )

                # Set up LLM based on environment configuration
                llm = setup_browser_llm()

                # Create and run the agent
                agent = Agent(
                    task=instruction,
                    llm=llm,
                    browser=browser,
                    max_steps=max_steps,
                    use_vision=use_vision
                )

                # Run the agent
                result = await agent.run()

                # Close browser
                await browser.close()

                return {
                    "result": {
                        "content": [
                            {
                                "type": "text",
                                "text": f"Browser automation completed successfully.\n\nResult: {result}"
                            }
                        ]
                    }
                }

            except Exception as e:
                import traceback
                error_details = traceback.format_exc()
                return {
                    "result": {
                        "content": [
                            {
                                "type": "text",
                                "text": f"Browser automation failed: {str(e)}\n\nDetails: {error_details}"
                            }
                        ]
                    }
                }
        else:
            # Handle other tools through the MCP client
            if not tool_name:
                print("❌ [TOOLS/CALL] toolName is required but missing")
                raise HTTPException(status_code=400, detail="toolName is required")

            # Make MCP request to call tool
            result = await client.request(
                method="tools/call",
                params={
                    "name": tool_name,
                    "arguments": parameters
                }
            )

            return {"result": result}

    except HTTPException:
        raise
    except Exception as e:
        print(f"💥 [TOOLS/CALL] Error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to call tool on MCP server: {str(e)}")

@app.post("/v1/mcp/servers/{server_id}/tools/list")
async def list_tools(server_id: str):
    """List tools available on an MCP server."""
    if not MCP_AVAILABLE:
        raise HTTPException(status_code=503, detail="MCP SDK not available")

    try:
        print(f"🔍 [TOOLS/LIST] Server: {server_id}")

        # Check if this is the MCP Browser Use server
        server = mcp_servers.get(server_id)
        if server and "mcp-browser-use" in server.get("name", "").lower():
            # Return the browser automation tool directly
            return {
                "result": {
                    "tools": [
                        {
                            "name": "run_browser_agent",
                            "description": "Control a web browser using natural language commands. Executes browser automation tasks and returns results.",
                            "inputSchema": {
                                "type": "object",
                                "properties": {
                                    "instruction": {
                                        "type": "string",
                                        "description": "Natural language instruction for browser automation (e.g., 'Navigate to Google and search for cats')",
                                    },
                                    "max_steps": {
                                        "type": "integer",
                                        "description": "Maximum number of steps the agent should take",
                                        "default": 10
                                    },
                                    "use_vision": {
                                        "type": "boolean",
                                        "description": "Whether to use vision for understanding page content",
                                        "default": True
                                    }
                                },
                                "required": ["instruction"]
                            }
                        }
                    ]
                }
            }

        global mcp_clients

        client = mcp_clients.get(server_id)
        if not client:
            print(f"❌ [TOOLS/LIST] Server {server_id} not found or not connected")
            raise HTTPException(status_code=404, detail="Server is not connected")

        # Make MCP request to list tools
        result = await client.request(
            method="tools/list",
            params={}
        )

        print(f"📨 [TOOLS/LIST] Raw response from MCP server: {result}")

        # Validate response structure
        if not result:
            print("❌ [TOOLS/LIST] No result returned from MCP server")
        elif 'tools' not in result:
            print(f"❌ [TOOLS/LIST] Missing 'tools' field in response: {list(result.keys())}")
        elif not isinstance(result['tools'], list):
            print(f"❌ [TOOLS/LIST] 'tools' field is not an array: {type(result['tools'])}")
        else:
            print(f"✅ [TOOLS/LIST] Found {len(result['tools'])} tools in response")
            for i, tool in enumerate(result['tools']):
                print(f"  Tool {i}: {tool.get('name', 'unnamed')}")
                if 'name' not in tool:
                    print(f"    ❌ Missing name for tool {i}")
                if 'description' not in tool:
                    print(f"    ⚠️  Missing description for tool {i}")
                if 'inputSchema' not in tool:
                    print(f"    ⚠️  Missing inputSchema for tool {i}")
                else:
                    schema = tool['inputSchema']
                    print(f"    ✅ inputSchema type: {schema.get('type')}")
                    if 'properties' in schema:
                        print(f"    ✅ Has {len(schema['properties'])} properties")
                    else:
                        print("    ⚠️  No properties in inputSchema")

        return {"result": result}

    except HTTPException:
        raise
    except Exception as e:
        print(f"💥 [TOOLS/LIST] Error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to list tools on MCP server: {str(e)}")

# Root endpoint
@app.get("/")
async def root():
    """Root endpoint."""
    return {"message": "CATBot Proxy Server", "version": "2.0.0"}

@app.post("/v1/telegram/chat", response_model=TelegramChatResponse)
async def telegram_chat_endpoint(request: TelegramChatRequest):
    """Process a Telegram chat message via OpenAI-compatible API."""

    message_text = (request.message or "").strip()
    if not message_text:
        raise HTTPException(status_code=400, detail="message is required")

    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise HTTPException(status_code=503, detail="OPENAI_API_KEY is not configured on the server")

    conversation_id = request.conversation_id or request.user_id or "default"
    history = telegram_conversations.setdefault(conversation_id, [])

    if request.history is not None:
        history = [
            {"role": msg.role, "content": msg.content}
            for msg in request.history
            if msg.content
        ]
        telegram_conversations[conversation_id] = history
        trim_telegram_history(history)

    history.append({"role": "user", "content": message_text})
    trim_telegram_history(history)

    system_prompt = request.system_prompt
    if system_prompt is None:
        system_prompt = TELEGRAM_SYSTEM_PROMPT

    # Retrieve relevant memories if memory system is available
    memory_context = ""
    if MEMORY_AVAILABLE and memory_manager:
        try:
            # Search for relevant memories based on the current message
            # Use a lower threshold (0.3) for automatic retrieval to catch more relevant memories
            relevant_memories = await memory_manager.search_memories(
                query=message_text,
                limit=5,
                similarity_threshold=0.3,  # Lower threshold for better recall
            )
            
            # Log search results for debugging
            if relevant_memories:
                print(f"Found {len(relevant_memories)} relevant memories for query: '{message_text[:50]}...'")
                for mem in relevant_memories:
                    print(f"  - {mem.get('text', '')} (similarity: {mem.get('similarity', 0):.3f})")
            else:
                print(f"No memories found for query: '{message_text[:50]}...' (threshold: 0.3)")
            
            # Build memory context if memories found
            if relevant_memories:
                memory_context = "\n\nRelevant context from previous conversations:\n"
                for i, mem in enumerate(relevant_memories, 1):
                    memory_context += f"{i}. {mem.get('text', '')}\n"
                memory_context += "\nUse this context to provide more personalized and relevant responses."
        except Exception as e:
            print(f"Warning: Failed to retrieve memories: {e}")
            import traceback
            print(traceback.format_exc())

    # Add memory context to system prompt
    if memory_context:
        system_prompt = system_prompt + memory_context

    messages: List[Dict[str, str]] = []
    if system_prompt:
        messages.append({"role": "system", "content": system_prompt})
    messages.extend(history)

    model_name = request.model or TELEGRAM_DEFAULT_MODEL
    if not model_name:
        raise HTTPException(status_code=400, detail="No model configured for Telegram chat")

    payload: Dict[str, Any] = {
        "model": model_name,
        "messages": messages,
    }

    if request.temperature is not None:
        payload["temperature"] = request.temperature

    if request.max_output_tokens is not None:
        payload["max_tokens"] = request.max_output_tokens

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }

    if OPENAI_ORG_ID:
        headers["OpenAI-Organization"] = OPENAI_ORG_ID

    if OPENAI_PROJECT_ID:
        headers["OpenAI-Project"] = OPENAI_PROJECT_ID

    url = build_openai_url(TELEGRAM_OPENAI_CHAT_PATH)

    try:
        async with httpx.AsyncClient(timeout=TELEGRAM_CHAT_TIMEOUT) as client:
            response = await client.post(url, headers=headers, json=payload)
    except httpx.RequestError as exc:
        print(f"Telegram chat request error: {exc}")
        raise HTTPException(status_code=502, detail="Failed to contact language model service") from exc

    if response.status_code != 200:
        print(f"Telegram chat API error {response.status_code}: {response.text}")
        detail = response.text
        try:
            error_json = response.json()
            detail = (
                error_json.get("error", {}).get("message")
                or error_json.get("message")
                or detail
            )
        except ValueError:
            pass
        raise HTTPException(status_code=response.status_code, detail=detail)

    data = response.json()
    reply = None
    choices = data.get("choices") or []
    if choices:
        message = choices[0].get("message") or {}
        reply = message.get("content")

    if not reply:
        reply = "I couldn't generate a response right now. Please try again shortly."

    history.append({"role": "assistant", "content": reply})
    trim_telegram_history(history)

    # Extract and store memories if memory system is available and auto-extract is enabled
    if MEMORY_AVAILABLE and memory_manager:
        auto_extract = os.getenv("MEMORY_AUTO_EXTRACT", "true").lower() == "true"
        if auto_extract:
            try:
                # Extract memories from the conversation (last few messages)
                # Include both user message and assistant response
                recent_messages = history[-4:] if len(history) >= 4 else history
                await memory_manager.extract_memories_from_conversation(
                    messages=recent_messages,
                    max_memories=3,
                )
            except Exception as e:
                print(f"Warning: Failed to extract memories: {e}")

    usage = data.get("usage") if isinstance(data, dict) else None

    return TelegramChatResponse(
        reply=reply,
        conversation_id=conversation_id,
        usage=usage,
    )


@app.delete("/v1/telegram/chat/{conversation_id}")
async def telegram_clear_conversation(conversation_id: str):
    """Clear cached Telegram conversation history for a user."""

    removed = telegram_conversations.pop(conversation_id, None) is not None
    return {"conversation_id": conversation_id, "cleared": removed}

# ============================================================================
# MEMORY SYSTEM ENDPOINTS
# ============================================================================

@app.post("/v1/memory/store", response_model=MemoryResponse)
async def store_memory(request: MemoryStoreRequest):
    """Store a memory explicitly."""
    if not MEMORY_AVAILABLE or not memory_manager:
        error_detail = "Memory system is not available."
        if not MEMORY_AVAILABLE:
            error_detail += f" Import failed: {memory_import_error}. Check /v1/memory/status for details."
        elif not memory_manager:
            error_detail += " Memory manager initialization failed. Check server logs and /v1/memory/status endpoint."
        raise HTTPException(
            status_code=503,
            detail=error_detail
        )
    
    try:
        # Store the memory
        memory_id = await memory_manager.store_memory(
            text=request.text,
            category=request.category,
            source=request.source or "explicit",
            metadata=request.metadata,
        )
        
        return MemoryResponse(
            success=True,
            message=f"Memory stored successfully",
            data={"memory_id": memory_id}
        )
    except Exception as e:
        print(f"Error storing memory: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to store memory: {str(e)}")

@app.post("/v1/memory/search", response_model=MemoryResponse)
async def search_memories(request: MemorySearchRequest):
    """Search memories by query."""
    if not MEMORY_AVAILABLE or not memory_manager:
        raise HTTPException(
            status_code=503,
            detail="Memory system is not available. Check MEMORY_ENABLED setting."
        )
    
    try:
        # Search for relevant memories
        results = await memory_manager.search_memories(
            query=request.query,
            limit=request.limit,
            similarity_threshold=request.similarity_threshold,
            category=request.category,
        )
        
        return MemoryResponse(
            success=True,
            message=f"Found {len(results)} relevant memories",
            data={"memories": results, "count": len(results)}
        )
    except Exception as e:
        print(f"Error searching memories: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to search memories: {str(e)}")

@app.get("/v1/memory/list", response_model=MemoryResponse)
async def list_memories(limit: Optional[int] = None):
    """List recent memories."""
    if not MEMORY_AVAILABLE or not memory_manager:
        raise HTTPException(
            status_code=503,
            detail="Memory system is not available. Check MEMORY_ENABLED setting."
        )
    
    try:
        # List memories
        memories = memory_manager.list_memories(limit=limit)
        
        return MemoryResponse(
            success=True,
            message=f"Retrieved {len(memories)} memories",
            data={"memories": memories, "count": len(memories), "total": memory_manager.count()}
        )
    except Exception as e:
        print(f"Error listing memories: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to list memories: {str(e)}")

@app.get("/v1/memory/{memory_id}", response_model=MemoryResponse)
async def get_memory(memory_id: str):
    """Get a specific memory by ID."""
    if not MEMORY_AVAILABLE or not memory_manager:
        raise HTTPException(
            status_code=503,
            detail="Memory system is not available. Check MEMORY_ENABLED setting."
        )
    
    try:
        # Get memory by ID
        memory = memory_manager.get_memory(memory_id)
        
        if not memory:
            raise HTTPException(status_code=404, detail=f"Memory {memory_id} not found")
        
        return MemoryResponse(
            success=True,
            message="Memory retrieved successfully",
            data={"memory": memory}
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error getting memory: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to get memory: {str(e)}")

@app.post("/v1/memory/extract", response_model=MemoryResponse)
async def extract_memories(request: MemoryExtractRequest):
    """Extract and store memories from a conversation."""
    if not MEMORY_AVAILABLE or not memory_manager:
        error_detail = "Memory system is not available."
        if not MEMORY_AVAILABLE:
            error_detail += f" Import failed: {memory_import_error}. Check /v1/memory/status for details."
        elif not memory_manager:
            error_detail += " Memory manager initialization failed. Check server logs and /v1/memory/status endpoint."
        raise HTTPException(
            status_code=503,
            detail=error_detail
        )
    
    try:
        # Check if auto-extract is enabled
        auto_extract = os.getenv("MEMORY_AUTO_EXTRACT", "true").lower() == "true"
        if not auto_extract:
            return MemoryResponse(
                success=False,
                message="Automatic memory extraction is disabled via MEMORY_AUTO_EXTRACT=false",
                data={"extracted": 0}
            )
        
        # Extract memories from conversation
        max_memories = request.max_memories or 1
        memory_ids = await memory_manager.extract_memories_from_conversation(
            messages=request.messages,
            max_memories=max_memories,
        )
        
        return MemoryResponse(
            success=True,
            message=f"Extracted and stored {len(memory_ids)} memories",
            data={"extracted": len(memory_ids), "memory_ids": memory_ids}
        )
    except Exception as e:
        print(f"Error extracting memories: {e}")
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Failed to extract memories: {str(e)}")

@app.get("/v1/memory/status")
async def memory_status():
    """Get the status of the memory system."""
    status = {
        "available": MEMORY_AVAILABLE,
        "enabled": os.getenv("MEMORY_ENABLED", "true").lower() == "true",
        "initialized": memory_manager is not None,
        "memory_count": memory_manager.count() if memory_manager else 0,
    }
    if not MEMORY_AVAILABLE:
        status["error"] = f"Memory module not available (import failed: {memory_import_error})"
        if memory_import_error and "numpy" in memory_import_error.lower():
            status["fix"] = "Install numpy with: pip install numpy"
    elif not status["enabled"]:
        status["error"] = "Memory system disabled via MEMORY_ENABLED=false"
    elif not status["initialized"]:
        status["error"] = "Memory manager failed to initialize (check server logs)"
    return status

@app.delete("/v1/memory/{memory_id}", response_model=MemoryResponse)
async def delete_memory(memory_id: str):
    """Delete a memory by ID."""
    if not MEMORY_AVAILABLE or not memory_manager:
        error_detail = "Memory system is not available."
        if not MEMORY_AVAILABLE:
            error_detail += " Memory module import failed."
        elif not memory_manager:
            error_detail += " Memory manager initialization failed. Check server logs and /v1/memory/status endpoint."
        raise HTTPException(
            status_code=503,
            detail=error_detail
        )
    
    try:
        # Delete memory
        deleted = memory_manager.delete_memory(memory_id)
        
        if not deleted:
            raise HTTPException(status_code=404, detail=f"Memory {memory_id} not found")
        
        return MemoryResponse(
            success=True,
            message=f"Memory {memory_id} deleted successfully",
            data={"memory_id": memory_id}
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error deleting memory: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to delete memory: {str(e)}")

# ============================================================================
# END MEMORY SYSTEM ENDPOINTS
# ============================================================================

# ============================================================================
# PHILOSOPHER MODE HELPER FUNCTIONS
# ============================================================================

async def get_all_available_tools() -> List[Dict]:
    """Get all available tools from all connected MCP servers and built-in proxy tools."""
    print(f"[PHILOSOPHER] get_all_available_tools called - MCP_AVAILABLE: {MCP_AVAILABLE}")
    
    all_tools = []
    
    # Add built-in proxy tools (web search, web scraper, news API)
    # These are always available if the server is running
    
    # 1. Web Search Tool
    all_tools.append({
        "name": "web_search",
        "description": "Search the web using Brave Search API (with DuckDuckGo fallback). Returns top search results with URLs, titles, and snippets.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "The search query to execute (e.g., 'latest AI developments 2024')",
                }
            },
            "required": ["query"]
        },
        "server_id": "proxy_server"
    })
    print("[PHILOSOPHER] Added web_search tool")
    
    # 2. Web Scraper/Fetcher Tool
    all_tools.append({
        "name": "web_scraper",
        "description": "Fetch and scrape content from a web URL. Returns the HTML content of the webpage.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "url": {
                    "type": "string",
                    "description": "The URL to fetch and scrape (must include http:// or https://)",
                }
            },
            "required": ["url"]
        },
        "server_id": "proxy_server"
    })
    print("[PHILOSOPHER] Added web_scraper tool")
    
    # 3. News API Tool (only if API key is configured)
    news_api_key = os.getenv('NEWS_API_KEY')
    if news_api_key:
        all_tools.append({
            "name": "news_search",
            "description": "Search for recent news articles using News API. Returns articles with titles, URLs, descriptions, and publication dates.",
            "inputSchema": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "The search query for news articles (e.g., 'artificial intelligence')",
                    }
                },
                "required": ["query"]
            },
            "server_id": "proxy_server"
        })
        print("[PHILOSOPHER] Added news_search tool")
    else:
        print("[PHILOSOPHER] NEWS_API_KEY not configured, skipping news_search tool")
    
    # Add MCP tools if MCP is available
    if MCP_AVAILABLE:
        # Debug: Check what servers are available
        print(f"[PHILOSOPHER] Checking MCP tools - mcp_clients: {len(mcp_clients)} clients, mcp_servers: {len(mcp_servers)} servers")
        print(f"[PHILOSOPHER] Connected client IDs: {list(mcp_clients.keys())}")
        print(f"[PHILOSOPHER] Server IDs in mcp_servers: {list(mcp_servers.keys())}")
        
        # First, check mcp_servers for connected browser-use servers
        # (Browser-use servers are marked as connected but may not be in mcp_clients)
        for server_id, server in mcp_servers.items():
            server_status = server.get("status", "disconnected")
            server_name = server.get("name", "").lower()
            
            # Check if this is a connected browser-use server
            if server_status == "connected" and "mcp-browser-use" in server_name:
                print(f"[PHILOSOPHER] Found connected browser-use server: {server_id}")
                # Add browser automation tool
                all_tools.append({
                    "name": "run_browser_agent",
                    "description": "Control a web browser using natural language commands. Executes browser automation tasks and returns results.",
                    "inputSchema": {
                        "type": "object",
                        "properties": {
                            "instruction": {
                                "type": "string",
                                "description": "Natural language instruction for browser automation (e.g., 'Navigate to Google and search for cats')",
                            },
                            "max_steps": {
                                "type": "integer",
                                "description": "Maximum number of steps the agent should take",
                                "default": 10
                            },
                            "use_vision": {
                                "type": "boolean",
                                "description": "Whether to use vision for understanding page content",
                                "default": True
                            }
                        },
                        "required": ["instruction"]
                    },
                    "server_id": server_id
                })
        
        # Get tools from each connected MCP client (non-browser-use servers)
        for server_id, client in mcp_clients.items():
            try:
                print(f"[PHILOSOPHER] Processing MCP server: {server_id}")
                # Check if this is the browser-use server (shouldn't be, but check anyway)
                server = mcp_servers.get(server_id)
                print(f"[PHILOSOPHER] Server config for {server_id}: {server}")
                
                if server and "mcp-browser-use" in server.get("name", "").lower():
                    # Skip - already handled above
                    print(f"[PHILOSOPHER] Skipping browser-use server {server_id} (already handled)")
                    continue
                else:
                    # Get tools from MCP server
                    print(f"[PHILOSOPHER] Requesting tools/list from MCP server {server_id}")
                    result = await client.request(method="tools/list", params={})
                    print(f"[PHILOSOPHER] tools/list response from {server_id}: {result}")
                    if result and "tools" in result:
                        print(f"[PHILOSOPHER] Found {len(result['tools'])} tools from server {server_id}")
                        for tool in result["tools"]:
                            tool["server_id"] = server_id
                            all_tools.append(tool)
                    else:
                        print(f"[PHILOSOPHER] No tools in response from server {server_id}")
            except Exception as e:
                print(f"[PHILOSOPHER] Error getting tools from server {server_id}: {e}")
                import traceback
                print(traceback.format_exc())
                continue
    else:
        print("[PHILOSOPHER] MCP not available, skipping MCP tools")
    
    print(f"[PHILOSOPHER] Total tools collected: {len(all_tools)}")
    return all_tools

async def execute_tool_for_philosopher(tool_name: str, parameters: Dict) -> str:
    """Execute a tool for philosopher mode. Returns result as string."""
    print(f"[PHILOSOPHER] Executing tool: {tool_name} with parameters: {parameters}")
    
    # Handle built-in proxy server tools
    if tool_name == "web_search":
        try:
            query = parameters.get("query", "")
            if not query:
                return "Error: 'query' parameter is required for web_search"
            
            # Call the proxy search endpoint
            result = await proxy_search(query)
            
            # Format results as readable text
            if result and "results" in result:
                results = result["results"]
                if not results:
                    return f"No search results found for query: {query}"
                
                formatted_results = []
                for i, item in enumerate(results, 1):
                    formatted_results.append(
                        f"{i}. {item.get('title', 'No title')}\n"
                        f"   URL: {item.get('url', 'No URL')}\n"
                        f"   {item.get('snippet', 'No description')}"
                    )
                
                return f"Search results for '{query}':\n\n" + "\n\n".join(formatted_results)
            else:
                return f"Search returned no results for query: {query}"
        except HTTPException as e:
            return f"Error executing web_search: {e.detail}"
        except Exception as e:
            return f"Error executing web_search: {str(e)}"
    
    elif tool_name == "web_scraper":
        try:
            url = parameters.get("url", "")
            if not url:
                return "Error: 'url' parameter is required for web_scraper"
            
            # Call the proxy fetch endpoint
            result = await proxy_fetch(url)
            
            # Return the content (may be HTML)
            if result and "content" in result:
                content = result["content"]
                # Truncate if too long (keep first 10000 characters)
                if len(content) > 10000:
                    return f"Web content from {url} (truncated):\n\n{content[:10000]}...\n\n[Content truncated to 10000 characters]"
                return f"Web content from {url}:\n\n{content}"
            else:
                return f"No content retrieved from URL: {url}"
        except HTTPException as e:
            return f"Error executing web_scraper: {e.detail}"
        except Exception as e:
            return f"Error executing web_scraper: {str(e)}"
    
    elif tool_name == "news_search":
        try:
            query = parameters.get("query", "")
            if not query:
                return "Error: 'query' parameter is required for news_search"
            
            # Call the proxy news search endpoint
            result = await proxy_news_search(query)
            
            # Format results as readable text
            if result and "articles" in result:
                articles = result["articles"]
                if not articles:
                    return f"No news articles found for query: {query}"
                
                formatted_articles = []
                for i, article in enumerate(articles[:10], 1):  # Limit to top 10
                    formatted_articles.append(
                        f"{i}. {article.get('title', 'No title')}\n"
                        f"   Source: {article.get('source', 'Unknown')}\n"
                        f"   Published: {article.get('publishedAt', 'Unknown date')}\n"
                        f"   URL: {article.get('url', 'No URL')}\n"
                        f"   {article.get('description', 'No description')}"
                    )
                
                total = result.get("totalResults", len(articles))
                return f"News articles for '{query}' (showing {len(formatted_articles)} of {total}):\n\n" + "\n\n".join(formatted_articles)
            else:
                return f"News search returned no articles for query: {query}"
        except HTTPException as e:
            return f"Error executing news_search: {e.detail}"
        except Exception as e:
            return f"Error executing news_search: {str(e)}"
    
    # Handle MCP tools
    elif MCP_AVAILABLE:
        # Find which server has this tool
        all_tools = await get_all_available_tools()
        server_id = None
        
        for tool in all_tools:
            if tool.get("name") == tool_name:
                server_id = tool.get("server_id")
                break
        
        if not server_id:
            return f"Error: Tool '{tool_name}' not found on any connected server"
        
        # Execute the tool using the existing call_tool logic
        try:
            request = ToolCallRequest(toolName=tool_name, parameters=parameters)
            result = await call_tool(server_id, request)
            
            # Extract result content
            if result and "result" in result:
                result_data = result["result"]
                if "content" in result_data:
                    content = result_data["content"]
                    if isinstance(content, list):
                        # Extract text from content array
                        text_parts = []
                        for item in content:
                            if isinstance(item, dict) and item.get("type") == "text":
                                text_parts.append(item.get("text", ""))
                        return "\n".join(text_parts)
                    return str(content)
                return str(result_data)
            
            return str(result)
        except Exception as e:
            return f"Error executing tool {tool_name}: {str(e)}"
    else:
        return f"Error: Tool '{tool_name}' is not a built-in tool and MCP is not available"

# ============================================================================
# PHILOSOPHER MODE ENDPOINTS
# ============================================================================

@app.post("/v1/philosopher/start", response_model=PhilosopherResponse)
async def philosopher_start(request: PhilosopherStartRequest):
    """Enable philosopher mode for a conversation."""
    if not PHILOSOPHER_MODE_AVAILABLE or not PhilosopherMode:
        raise HTTPException(
            status_code=503,
            detail="Philosopher mode is not available. Check server logs for details."
        )
    
    # Check if philosopher mode is enabled
    philosopher_enabled = os.getenv("PHILOSOPHER_MODE_ENABLED", "true").lower() == "true"
    if not philosopher_enabled:
        raise HTTPException(
            status_code=503,
            detail="Philosopher mode is disabled via PHILOSOPHER_MODE_ENABLED=false"
        )
    
    # Get conversation ID
    conversation_id = request.conversation_id or request.user_id or "default"
    
    # Check if already active
    if philosopher_mode_active.get(conversation_id, False):
        return PhilosopherResponse(
            success=True,
            message="Philosopher mode is already active for this conversation",
            data={"conversation_id": conversation_id, "active": True}
        )
    
    try:
        # Get API configuration
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise HTTPException(status_code=503, detail="OPENAI_API_KEY is not configured")
        
        api_base = os.getenv("OPENAI_API_BASE", "https://api.openai.com/v1")
        # Use OPENAI_MODEL directly (not TELEGRAM_DEFAULT_MODEL) since philosopher mode is not Telegram-specific
        model = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
        
        # Create philosopher mode instance with tool support
        philosopher = PhilosopherMode(
            api_key=api_key,
            api_base=api_base,
            model=model,
            memory_manager=memory_manager if MEMORY_AVAILABLE else None,
            max_cycles=int(os.getenv("PHILOSOPHER_MAX_CYCLES", "10")),
            similarity_threshold=float(os.getenv("PHILOSOPHER_SIMILARITY_THRESHOLD", "0.3")),
            memory_limit=int(os.getenv("PHILOSOPHER_MEMORY_LIMIT", "10")),
            conversation_history_limit=int(os.getenv("PHILOSOPHER_CONVERSATION_HISTORY_LIMIT", "3")),
            tool_executor=execute_tool_for_philosopher,
            get_tools_func=get_all_available_tools,
            diversification_threshold=int(os.getenv("PHILOSOPHER_DIVERSIFICATION_THRESHOLD", "7")),
        )
        
        # Store instance and activate mode
        philosopher_mode_instances[conversation_id] = philosopher
        philosopher_mode_active[conversation_id] = True
        
        return PhilosopherResponse(
            success=True,
            message="Philosopher mode activated",
            data={"conversation_id": conversation_id, "active": True}
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error starting philosopher mode: {e}")
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Failed to start philosopher mode: {str(e)}")

@app.post("/v1/philosopher/stop", response_model=PhilosopherResponse)
async def philosopher_stop(request: PhilosopherStopRequest):
    """Disable philosopher mode for a conversation."""
    # Get conversation ID
    conversation_id = request.conversation_id or request.user_id or "default"
    
    # Check if active
    if not philosopher_mode_active.get(conversation_id, False):
        return PhilosopherResponse(
            success=True,
            message="Philosopher mode is not active for this conversation",
            data={"conversation_id": conversation_id, "active": False}
        )
    
    try:
        # Deactivate mode
        philosopher_mode_active[conversation_id] = False
        # Remove instance (optional - could keep for reuse)
        philosopher_mode_instances.pop(conversation_id, None)
        
        return PhilosopherResponse(
            success=True,
            message="Philosopher mode deactivated",
            data={"conversation_id": conversation_id, "active": False}
        )
    except Exception as e:
        print(f"Error stopping philosopher mode: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to stop philosopher mode: {str(e)}")

@app.get("/v1/philosopher/status")
async def philosopher_status(conversation_id: Optional[str] = None, user_id: Optional[str] = None):
    """Check if philosopher mode is active for a conversation."""
    # Get conversation ID
    conv_id = conversation_id or user_id or "default"
    
    is_active = philosopher_mode_active.get(conv_id, False)
    
    return {
        "active": is_active,
        "conversation_id": conv_id,
        "available": PHILOSOPHER_MODE_AVAILABLE,
        "enabled": os.getenv("PHILOSOPHER_MODE_ENABLED", "true").lower() == "true"
    }

@app.post("/v1/philosopher/contemplate", response_model=PhilosopherResponse)
async def philosopher_contemplate(request: PhilosopherContemplateRequest):
    """Execute a single contemplation cycle."""
    if not PHILOSOPHER_MODE_AVAILABLE or not PhilosopherMode:
        raise HTTPException(
            status_code=503,
            detail="Philosopher mode is not available. Check server logs for details."
        )
    
    # Get conversation ID
    conversation_id = request.conversation_id or request.user_id or "default"
    
    # Check if mode is active
    if not philosopher_mode_active.get(conversation_id, False):
        raise HTTPException(
            status_code=400,
            detail="Philosopher mode is not active for this conversation. Start it first with /v1/philosopher/start"
        )
    
    # Get philosopher instance
    philosopher = philosopher_mode_instances.get(conversation_id)
    if not philosopher:
        raise HTTPException(
            status_code=500,
            detail="Philosopher mode instance not found. Try restarting philosopher mode."
        )
    
    try:
        # Generate question if not provided
        if request.question:
            question = request.question
        else:
            question = await philosopher.generate_contemplation_question()
            if not question:
                raise HTTPException(status_code=500, detail="Failed to generate contemplation question")
        
        # Execute contemplation
        result = await philosopher.contemplate_question(question)
        
        # Store contemplation in memory
        memory_id = await philosopher.store_contemplation(
            question=result["question"],
            conclusion=result["conclusion"],
            cycle_count=result["cycle_count"]
        )
        
        return PhilosopherResponse(
            success=True,
            message="Contemplation completed",
            data={
                "question": result["question"],
                "conclusion": result["conclusion"],
                "contemplation_steps": result["contemplation_steps"],
                "cycle_count": result["cycle_count"],
                "memory_id": memory_id,
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error during contemplation: {e}")
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Failed to contemplate: {str(e)}")

# ============================================================================
# END PHILOSOPHER MODE ENDPOINTS
# ============================================================================

# Simple test endpoint to verify requests are reaching the server
@app.get("/test")
async def test_endpoint():
    """Simple test endpoint that should always work."""
    import sys
    print("🧪 TEST endpoint called", flush=True)
    sys.stdout.flush()
    return {"message": "test successful", "timestamp": time.time()}

# Health check endpoint
@app.get("/health")
async def health_check():
    """Health check endpoint."""
    # Add explicit logging to debug
    import sys
    print("🏥 Health check endpoint called", flush=True)
    sys.stdout.flush()
    try:
        result = {"status": "healthy", "timestamp": time.time()}
        print(f"🏥 Health check returning: {result}", flush=True)
        sys.stdout.flush()
        return result
    except Exception as e:
        import traceback
        print(f"❌ Health check error: {e}", flush=True)
        print(traceback.format_exc(), flush=True)
        sys.stdout.flush()
        raise

# Whisper proxy endpoint to handle CORS
@app.post("/v1/audio/transcriptions")
async def proxy_whisper(request: Request):
    """Proxy Whisper transcription requests to handle CORS."""
    try:
        # Get the form data from the request
        form_data = await request.form()
        
        # Get the Whisper endpoint (defaulting to localhost:8001)
        whisper_endpoint = os.getenv('WHISPER_ENDPOINT', 'http://localhost:8001/v1/audio/transcriptions')
        
        print(f"📝 Proxying Whisper request to: {whisper_endpoint}")
        
        # Get Authorization header from the request
        auth_header = request.headers.get('Authorization', '')
        
        # Prepare the files and data for forwarding
        files = {}
        data = {}
        
        for key, value in form_data.items():
            if hasattr(value, 'read'):  # This is a file
                # Read the file content
                file_content = await value.read()
                files[key] = (value.filename, file_content, value.content_type)
                print(f"  📎 File: {value.filename} ({len(file_content)} bytes)")
            else:  # This is a regular form field
                data[key] = value
                print(f"  📄 Field: {key} = {value}")
        
        # Forward the request to the Whisper service
        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.post(
                whisper_endpoint,
                files=files,
                data=data,
                headers={'Authorization': auth_header} if auth_header else {}
            )
        
        print(f"✅ Whisper response status: {response.status_code}")
        print(f"📄 Response content type: {response.headers.get('content-type', 'unknown')}")
        
        # Check if the response is successful
        if response.status_code != 200:
            print(f"❌ Whisper service returned error: {response.status_code}")
            print(f"   Response text: {response.text}")
            return JSONResponse(
                content={"error": f"Whisper service error: {response.text}"},
                status_code=response.status_code
            )
        
        # Try to parse the JSON response
        try:
            response_data = response.json()
            print(f"✅ Parsed JSON response: {response_data}")
            return JSONResponse(content=response_data, status_code=200)
        except Exception as json_error:
            print(f"❌ Failed to parse JSON response: {json_error}")
            print(f"   Raw response text: {response.text[:200]}")
            # Return the raw text if JSON parsing fails
            return JSONResponse(
                content={"text": response.text},
                status_code=200
            )
    
    except httpx.ConnectError as e:
        print(f"❌ Connection error: Could not connect to Whisper service at {whisper_endpoint}")
        print(f"   Make sure the Whisper service is running on port 8001")
        raise HTTPException(
            status_code=503,
            detail=f"Could not connect to Whisper service. Make sure it's running on port 8001."
        )
    except Exception as e:
        print(f"❌ Whisper proxy error: {e}")
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Failed to proxy Whisper request: {str(e)}")

# TTS voices proxy endpoint to handle CORS
@app.get("/v1/proxy/tts/voices")
async def proxy_tts_voices(endpoint: str):
    """Proxy TTS voices requests to handle CORS. Tries /voices first, then /v1/audio/voices."""
    if not endpoint:
        raise HTTPException(status_code=400, detail="Endpoint parameter is required")
    
    try:
        # Normalize the endpoint URL (remove trailing slash)
        base_endpoint = endpoint.rstrip('/')
        
        # Extract base URL (origin: protocol + host + port) to avoid path duplication
        try:
            # Parse the endpoint URL to extract the origin (protocol + host + port)
            if not base_endpoint.startswith('http://') and not base_endpoint.startswith('https://'):
                # If no protocol, assume http://
                base_endpoint = f"http://{base_endpoint}"
            
            parsed_url = urlparse(base_endpoint)
            # Reconstruct base URL with scheme, hostname, and port (if present)
            base_url = f"{parsed_url.scheme}://{parsed_url.netloc}"
        except Exception as e:
            # Fallback to simple string replacement if URL parsing fails
            # Extract origin manually using regex
            match = re.match(r'^(https?://[^/]+)', base_endpoint)
            if match:
                base_url = match.group(1)
            else:
                # Last resort: remove /v1 and any path
                base_url = base_endpoint.split('/')[0] if '/' in base_endpoint else base_endpoint
                if not base_url.startswith('http'):
                    base_url = f"http://{base_url}"
        
        # Try /voices first (Chatterbox style)
        voices_url_primary = f"{base_url}/voices"
        print(f"🎤 Trying primary TTS voices endpoint: {voices_url_primary}")
        
        response = None
        response_data = None
        
        async with httpx.AsyncClient(timeout=10.0) as client:
            try:
                # Try the primary endpoint first
                response = await client.get(
                    voices_url_primary,
                    headers={
                        'Content-Type': 'application/json',
                        'Accept': 'application/json'
                    }
                )
                
                print(f"✅ Primary TTS voices response status: {response.status_code}")
                
                # If primary endpoint succeeds, use it
                if response.status_code == 200:
                    try:
                        response_data = response.json()
                        print(f"✅ Parsed TTS voices JSON response from primary endpoint")
                        return JSONResponse(content=response_data, status_code=200)
                    except Exception as json_error:
                        print(f"❌ Failed to parse JSON response: {json_error}")
                        print(f"   Raw response text: {response.text[:200]}")
                        # Return the raw text if JSON parsing fails
                        return JSONResponse(
                            content={"text": response.text},
                            status_code=200
                        )
            except (httpx.ConnectError, httpx.HTTPStatusError) as e:
                # Primary endpoint failed, try fallback
                print(f"⚠️ Primary endpoint failed: {e}")
                response = None
            
            # If primary failed, try /v1/audio/voices (OpenAI-compatible style)
            if not response or response.status_code != 200:
                voices_url_fallback = f"{base_url}/v1/audio/voices"
                print(f"🎤 Trying fallback TTS voices endpoint: {voices_url_fallback}")
                
                try:
                    response = await client.get(
                        voices_url_fallback,
                        headers={
                            'Content-Type': 'application/json',
                            'Accept': 'application/json'
                        }
                    )
                    
                    print(f"✅ Fallback TTS voices response status: {response.status_code}")
                    
                    # Check if the fallback response is successful
                    if response.status_code == 200:
                        try:
                            response_data = response.json()
                            print(f"✅ Parsed TTS voices JSON response from fallback endpoint")
                            return JSONResponse(content=response_data, status_code=200)
                        except Exception as json_error:
                            print(f"❌ Failed to parse JSON response: {json_error}")
                            print(f"   Raw response text: {response.text[:200]}")
                            # Return the raw text if JSON parsing fails
                            return JSONResponse(
                                content={"text": response.text},
                                status_code=200
                            )
                    else:
                        # Fallback also failed
                        print(f"❌ Fallback TTS service returned error: {response.status_code}")
                        print(f"   Response text: {response.text[:200]}")
                        raise HTTPException(
                            status_code=response.status_code,
                            detail=f"TTS service error: {response.text[:200]}"
                        )
                except httpx.ConnectError as e:
                    print(f"❌ Connection error: Could not connect to TTS service at {voices_url_fallback}")
                    raise HTTPException(
                        status_code=503,
                        detail=f"Could not connect to TTS service. Tried {voices_url_primary} and {voices_url_fallback}"
                    )
                except httpx.HTTPStatusError as e:
                    print(f"❌ HTTP error from fallback TTS service: {e.response.status_code}")
                    raise HTTPException(
                        status_code=e.response.status_code,
                        detail=f"TTS service returned error: {str(e)}"
                    )
        
        # If we get here, both attempts failed
        if response and response.status_code != 200:
            raise HTTPException(
                status_code=response.status_code,
                detail=f"TTS service error: {response.text[:200]}"
            )
        else:
            raise HTTPException(
                status_code=503,
                detail=f"Could not connect to TTS service. Tried {voices_url_primary} and {base_url}/v1/audio/voices"
            )
    
    except HTTPException:
        # Re-raise HTTP exceptions as-is
        raise
    except httpx.ConnectError as e:
        print(f"❌ Connection error: Could not connect to TTS service at {endpoint}")
        raise HTTPException(
            status_code=503,
            detail=f"Could not connect to TTS service at {endpoint}"
        )
    except httpx.HTTPStatusError as e:
        print(f"❌ HTTP error from TTS service: {e.response.status_code}")
        raise HTTPException(
            status_code=e.response.status_code,
            detail=f"TTS service returned error: {str(e)}"
        )
    except Exception as e:
        print(f"❌ TTS voices proxy error: {e}")
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Failed to proxy TTS voices request: {str(e)}")

# TTS speech proxy endpoint to handle CORS and streaming
@app.post("/v1/proxy/tts/speech")
async def proxy_tts_speech(request: Request, endpoint: Optional[str] = None):
    """Proxy TTS speech requests to handle CORS. Supports streaming responses."""
    try:
        # Get endpoint from query parameter or try to extract from request
        if not endpoint:
            # Try to get from query params
            endpoint = request.query_params.get('endpoint', '')
        
        # If still no endpoint, try to get from TTS_ENDPOINT env var or use default
        if not endpoint:
            tts_endpoint = os.getenv('TTS_ENDPOINT', 'http://localhost:4123/v1')
            endpoint = tts_endpoint.rstrip('/')
        
        # Normalize the endpoint URL (remove trailing slash)
        base_endpoint = endpoint.rstrip('/')
        
        # Extract base URL (origin: protocol + host + port) to avoid path duplication
        try:
            # Parse the endpoint URL to extract the origin (protocol + host + port)
            if not base_endpoint.startswith('http://') and not base_endpoint.startswith('https://'):
                # If no protocol, assume http://
                base_endpoint = f"http://{base_endpoint}"
            
            parsed_url = urlparse(base_endpoint)
            # Reconstruct base URL with scheme, hostname, and port (if present)
            base_url = f"{parsed_url.scheme}://{parsed_url.netloc}"
        except Exception as e:
            # Fallback to simple string replacement if URL parsing fails
            # Extract origin manually using regex
            match = re.match(r'^(https?://[^/]+)', base_endpoint)
            if match:
                base_url = match.group(1)
            else:
                # Last resort: remove /v1 and any path
                base_url = base_endpoint.split('/')[0] if '/' in base_endpoint else base_endpoint
                if not base_url.startswith('http'):
                    base_url = f"http://{base_url}"
        
        # Construct the speech endpoint URL
        speech_url = f"{base_url}/v1/audio/speech"
        
        print(f"🎤 Proxying TTS speech request to: {speech_url}")
        
        # Get the request body
        try:
            request_body = await request.json()
        except Exception:
            request_body = {}
        
        # Get headers from the original request
        # Forward Accept header to support both SSE (Chatterbox) and binary audio (VibeVoice/OpenAI-compatible)
        # If client requests SSE, forward it; otherwise let TTS service decide (defaults to binary audio)
        accept_header = request.headers.get('Accept', '')
        forward_headers = {
            'Content-Type': 'application/json',
        }
        
        # Forward Accept header if present (allows Chatterbox to return SSE, VibeVoice will ignore and return binary)
        if accept_header:
            forward_headers['Accept'] = accept_header
        
        # Forward Authorization header if present
        auth_header = request.headers.get('Authorization')
        if auth_header:
            forward_headers['Authorization'] = auth_header
        
        # Create a streaming response generator
        # We need to peek at the response to get the actual content-type
        # Since we can't change media_type after StreamingResponse is created,
        # we'll use a wrapper that can handle both SSE and binary audio
        
        class TTSStreamWrapper:
            """Wrapper to handle both SSE and binary audio responses."""
            def __init__(self):
                self.content_type = None
                self.first_chunk = True
                
            async def __aiter__(self):
                try:
                    async with httpx.AsyncClient(timeout=60.0) as client:
                        async with client.stream(
                            'POST',
                            speech_url,
                            json=request_body,
                            headers=forward_headers
                        ) as response:
                            print(f"✅ TTS speech response status: {response.status_code}")
                            
                            # Log request details for debugging
                            print(f"📤 TTS request body: {json.dumps(request_body, indent=2)[:500]}")
                            print(f"📤 TTS request headers: {forward_headers}")
                            
                            # Capture the actual content type from the TTS service
                            self.content_type = response.headers.get('content-type', 'audio/mpeg')
                            print(f"📦 TTS response content-type: {self.content_type}")
                            
                            # Check if the response is successful
                            if response.status_code != 200:
                                error_text = await response.aread()
                                print(f"❌ TTS service returned error: {response.status_code}")
                                print(f"   Response text: {error_text[:200]}")
                                # Yield error as bytes
                                if isinstance(error_text, bytes):
                                    yield error_text
                                else:
                                    yield str(error_text).encode('utf-8')
                                return
                            
                            # Stream the response chunk by chunk as bytes
                            # This preserves binary audio data (MP3, WAV, etc.) or SSE text
                            async for chunk in response.aiter_bytes():
                                if chunk:
                                    yield chunk
                                    
                except httpx.ConnectError as e:
                    print(f"❌ Connection error: Could not connect to TTS service at {speech_url}")
                    error_msg = f"Error: Could not connect to TTS service at {speech_url}"
                    yield error_msg.encode('utf-8')
                except Exception as e:
                    print(f"❌ TTS speech proxy error: {e}")
                    import traceback
                    print(traceback.format_exc())
                    error_msg = f"Error: Failed to proxy TTS speech request: {str(e)}"
                    yield error_msg.encode('utf-8')
        
        stream_wrapper = TTSStreamWrapper()
        
        # Determine initial content-type based on Accept header
        # If client requests SSE, use that; otherwise default to audio/mpeg
        # The actual content-type will be in the response headers
        initial_content_type = 'audio/mpeg'  # Default
        if accept_header and 'text/event-stream' in accept_header:
            initial_content_type = 'text/event-stream'  # Use SSE if requested
        
        # Use StreamingResponse with the initial content-type
        # The actual content-type from the TTS service will be in the response headers
        # JavaScript will check res.headers.get('content-type') to get the actual type
        # Note: FastAPI's StreamingResponse sets Content-Type based on media_type parameter,
        # but the actual response from the TTS service should have its own content-type header
        # that JavaScript can read. However, if Chatterbox is returning audio/mpeg instead of
        # text/event-stream, that means Chatterbox itself is not respecting the stream_format: 'sse'
        # parameter or the Accept header.
        
        response_obj = StreamingResponse(
            stream_wrapper,
            media_type=initial_content_type,  # Initial guess based on Accept header
            status_code=200
        )
        
        return response_obj
    
    except Exception as e:
        print(f"❌ TTS speech proxy error: {e}")
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Failed to proxy TTS speech request: {str(e)}")

# ============================================================================
# FILE OPERATIONS ENDPOINTS
# ============================================================================

def read_text_file(filepath: Path) -> str:
    """Read a plain text file and return its content"""
    try:
        # Read the file with UTF-8 encoding
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        # Fallback to latin-1 if UTF-8 fails
        with open(filepath, 'r', encoding='latin-1') as f:
            return f.read()

def read_docx_file(filepath: Path) -> str:
    """Read a Word document and return its text content"""
    # Load the document using python-docx
    doc = Document(filepath)
    # Extract text from all paragraphs
    paragraphs = [para.text for para in doc.paragraphs]
    # Join paragraphs with newlines
    return '\n'.join(paragraphs)

def read_xlsx_file(filepath: Path) -> str:
    """Read an Excel file and return its content as formatted text"""
    # Load the workbook
    wb = openpyxl.load_workbook(filepath, data_only=True)
    result = []
    
    # Process each sheet in the workbook
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        result.append(f"=== Sheet: {sheet_name} ===\n")
        
        # Process each row in the sheet
        for row in sheet.iter_rows(values_only=True):
            # Filter out None values and convert to strings
            row_data = [str(cell) if cell is not None else '' for cell in row]
            # Join cells with tabs for better formatting
            result.append('\t'.join(row_data))
        
        result.append('\n')  # Add blank line between sheets
    
    # Join all lines with newlines
    return '\n'.join(result)

def read_pdf_file(filepath: Path) -> str:
    """Read a PDF file and return its text content"""
    result = []
    # Open the PDF file in binary mode
    with open(filepath, 'rb') as f:
        # Create a PDF reader object
        pdf_reader = PyPDF2.PdfReader(f)
        # Extract text from each page
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text = page.extract_text()
            result.append(f"=== Page {page_num + 1} ===\n{text}\n")
    
    # Join all pages with newlines
    return '\n'.join(result)

def read_png_file(filepath: Path) -> Dict[str, Any]:
    """Read a PNG image and return metadata and base64-encoded data"""
    # Open the image using PIL
    img = Image.open(filepath)
    
    # Get image metadata
    metadata = {
        'width': img.width,
        'height': img.height,
        'format': img.format,
        'mode': img.mode
    }
    
    # Convert image to base64 for transmission
    buffered = BytesIO()
    img.save(buffered, format=img.format)
    img_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')
    
    return {
        'metadata': metadata,
        'data': img_base64,
        'description': f"Image: {img.width}x{img.height} pixels, format: {img.format}"
    }

def write_text_file(filepath: Path, content: str) -> None:
    """Write content to a plain text file"""
    # Write the content with UTF-8 encoding
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(content)

def write_docx_file(filepath: Path, content: str) -> None:
    """Write content to a Word document"""
    # Create a new document
    doc = Document()
    
    # Split content into paragraphs and add each to the document
    for paragraph in content.split('\n'):
        doc.add_paragraph(paragraph)
    
    # Save the document
    doc.save(filepath)

def write_xlsx_file(filepath: Path, content: str) -> None:
    """Write content to an Excel file"""
    # Create a new workbook
    wb = openpyxl.Workbook()
    # Get the active sheet
    ws = wb.active
    ws.title = "Sheet1"
    
    # Split content into rows
    rows = content.split('\n')
    
    # Write each row to the Excel file
    for row_idx, row_content in enumerate(rows, start=1):
        # Split row by tabs or commas
        if '\t' in row_content:
            cells = row_content.split('\t')
        else:
            cells = row_content.split(',')
        
        # Write each cell
        for col_idx, cell_content in enumerate(cells, start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=cell_content.strip())
            
            # Apply formatting to the first row (header)
            if row_idx == 1:
                cell.font = Font(bold=True)
                cell.alignment = Alignment(horizontal='center')
    
    # Auto-adjust column widths
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)  # Cap at 50 characters
        ws.column_dimensions[column_letter].width = adjusted_width
    
    # Save the workbook
    wb.save(filepath)

def write_pdf_file(filepath: Path, content: str) -> None:
    """Write content to a PDF file using reportlab"""
    try:
        # Import reportlab for PDF creation
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas as pdf_canvas
        from reportlab.lib.units import inch
        
        # Create a canvas for PDF generation
        c = pdf_canvas.Canvas(str(filepath), pagesize=letter)
        width, height = letter
        
        # Set up text formatting
        text_object = c.beginText(1 * inch, height - 1 * inch)
        text_object.setFont("Helvetica", 12)
        
        # Split content into lines and add to PDF
        lines = content.split('\n')
        for line in lines:
            # Handle long lines by wrapping
            if len(line) > 80:
                words = line.split(' ')
                current_line = ''
                for word in words:
                    if len(current_line + word) < 80:
                        current_line += word + ' '
                    else:
                        text_object.textLine(current_line)
                        current_line = word + ' '
                if current_line:
                    text_object.textLine(current_line)
            else:
                text_object.textLine(line)
            
            # Add new page if needed (simple check)
            if text_object.getY() < 1 * inch:
                c.drawText(text_object)
                c.showPage()
                text_object = c.beginText(1 * inch, height - 1 * inch)
                text_object.setFont("Helvetica", 12)
        
        # Draw the text and save the PDF
        c.drawText(text_object)
        c.save()
        
    except ImportError:
        # If reportlab is not available, raise an error
        raise HTTPException(
            status_code=500,
            detail="PDF writing requires reportlab library. Install with: pip install reportlab"
        )

@app.post("/v1/files/read", response_model=FileResponse)
async def read_file(request: ReadFileRequest):
    """
    Read a file from the scratch directory
    Supports: txt, docx, xlsx, pdf, png
    """
    if not FILE_OPS_AVAILABLE:
        raise HTTPException(
            status_code=503,
            detail="File operations not available. Install: pip install python-docx openpyxl PyPDF2 reportlab Pillow"
        )
    
    try:
        # Construct the full file path
        filepath = SCRATCH_DIR / request.filename
        
        # Check if file exists
        if not filepath.exists():
            return FileResponse(
                success=False,
                message=f"File not found: {request.filename}"
            )
        
        # Determine file extension
        file_ext = filepath.suffix.lower()
        
        # Read file based on extension
        if file_ext == '.txt':
            content = read_text_file(filepath)
            return FileResponse(
                success=True,
                message=f"Successfully read {request.filename}",
                data={'content': content, 'type': 'text'}
            )
        
        elif file_ext == '.docx':
            content = read_docx_file(filepath)
            return FileResponse(
                success=True,
                message=f"Successfully read {request.filename}",
                data={'content': content, 'type': 'text'}
            )
        
        elif file_ext in ['.xlsx', '.xls']:
            content = read_xlsx_file(filepath)
            return FileResponse(
                success=True,
                message=f"Successfully read {request.filename}",
                data={'content': content, 'type': 'text'}
            )
        
        elif file_ext == '.pdf':
            content = read_pdf_file(filepath)
            return FileResponse(
                success=True,
                message=f"Successfully read {request.filename}",
                data={'content': content, 'type': 'text'}
            )
        
        elif file_ext in ['.png', '.jpg', '.jpeg']:
            image_data = read_png_file(filepath)
            return FileResponse(
                success=True,
                message=f"Successfully read {request.filename}",
                data={'content': image_data['description'], 'type': 'image', 'image_data': image_data}
            )
        
        else:
            # Unsupported file type
            return FileResponse(
                success=False,
                message=f"Unsupported file type: {file_ext}. Supported types: txt, docx, xlsx, pdf, png"
            )
    
    except Exception as e:
        # Handle any errors during file reading
        return FileResponse(
            success=False,
            message=f"Error reading file: {str(e)}"
        )

@app.post("/v1/files/write", response_model=FileResponse)
async def write_file(request: WriteFileRequest):
    """
    Write content to a file in the scratch directory
    Supports: txt, docx, xlsx, pdf
    """
    if not FILE_OPS_AVAILABLE:
        raise HTTPException(
            status_code=503,
            detail="File operations not available. Install: pip install python-docx openpyxl PyPDF2 reportlab Pillow"
        )
    
    try:
        # Construct the full file path
        filepath = SCRATCH_DIR / request.filename
        
        # Determine file extension from filename or format parameter
        file_ext = filepath.suffix.lower()
        if not file_ext:
            # If no extension in filename, use format parameter
            file_ext = f".{request.format.lower()}"
            filepath = SCRATCH_DIR / f"{request.filename}{file_ext}"
        
        # Write file based on extension
        if file_ext == '.txt':
            write_text_file(filepath, request.content)
        
        elif file_ext == '.docx':
            write_docx_file(filepath, request.content)
        
        elif file_ext in ['.xlsx', '.xls']:
            write_xlsx_file(filepath, request.content)
        
        elif file_ext == '.pdf':
            write_pdf_file(filepath, request.content)
        
        else:
            # Unsupported file type
            return FileResponse(
                success=False,
                message=f"Unsupported file type for writing: {file_ext}. Supported types: txt, docx, xlsx, pdf"
            )
        
        # Return success response
        return FileResponse(
            success=True,
            message=f"Successfully wrote {filepath.name}",
            data={'filepath': str(filepath), 'size': filepath.stat().st_size}
        )
    
    except Exception as e:
        # Handle any errors during file writing
        return FileResponse(
            success=False,
            message=f"Error writing file: {str(e)}"
        )

@app.get("/v1/files/list")
async def list_files():
    """List all files in the scratch directory"""
    try:
        # Get all files in scratch directory
        files = []
        for file in SCRATCH_DIR.iterdir():
            if file.is_file():
                files.append({
                    'name': file.name,
                    'size': file.stat().st_size,
                    'modified': file.stat().st_mtime,
                    'extension': file.suffix
                })
        
        # Sort files by modification time (newest first)
        files.sort(key=lambda x: x['modified'], reverse=True)
        
        return {
            'success': True,
            'count': len(files),
            'files': files,
            'scratch_dir': str(SCRATCH_DIR)
        }
    
    except Exception as e:
        # Handle any errors during directory listing
        return {
            'success': False,
            'message': f"Error listing files: {str(e)}"
        }

@app.delete("/v1/files/delete/{filename}")
async def delete_file(filename: str):
    """Delete a file from the scratch directory"""
    try:
        # Construct the full file path
        filepath = SCRATCH_DIR / filename
        
        # Check if file exists
        if not filepath.exists():
            return FileResponse(
                success=False,
                message=f"File not found: {filename}"
            )
        
        # Delete the file
        filepath.unlink()
        
        return FileResponse(
            success=True,
            message=f"Successfully deleted {filename}"
        )
    
    except Exception as e:
        # Handle any errors during file deletion
        return FileResponse(
            success=False,
            message=f"Error deleting file: {str(e)}"
        )

# ============================================================================
# END FILE OPERATIONS ENDPOINTS
# ============================================================================

# ============================================================================
# GOOGLE DRIVE UPLOAD ENDPOINT
# ============================================================================

@app.post("/v1/proxy/upload-to-drive")
async def upload_to_drive(request: Request):
    """
    Upload a file to Google Drive using service account credentials from .env file.
    Credentials are read from environment variables:
    - GOOGLE_DRIVE_PROJECT_ID
    - GOOGLE_DRIVE_PRIVATE_KEY_ID
    - GOOGLE_DRIVE_PRIVATE_KEY
    - GOOGLE_DRIVE_CLIENT_EMAIL
    - GOOGLE_DRIVE_FOLDER_ID
    """
    try:
        # Get form data from request
        form_data = await request.form()
        
        # Get file path from form data
        file_path = form_data.get('filePath')
        if not file_path:
            raise HTTPException(status_code=400, detail="filePath is required")
        
        # Get optional file name
        file_name = form_data.get('fileName')
        
        # Get folder ID from form data or environment variable
        folder_id = form_data.get('folderId') or os.getenv('GOOGLE_DRIVE_FOLDER_ID')
        if not folder_id:
            raise HTTPException(status_code=400, detail="folderId is required (provide in request or set GOOGLE_DRIVE_FOLDER_ID in .env)")
        
        # Read Google Drive credentials from environment variables (loaded from .env file)
        project_id = os.getenv('GOOGLE_DRIVE_PROJECT_ID')
        private_key_id = os.getenv('GOOGLE_DRIVE_PRIVATE_KEY_ID')
        private_key = os.getenv('GOOGLE_DRIVE_PRIVATE_KEY')
        client_email = os.getenv('GOOGLE_DRIVE_CLIENT_EMAIL')
        
        # Validate that all required credentials are present
        if not all([project_id, private_key_id, private_key, client_email]):
            missing = [k for k, v in {
                'GOOGLE_DRIVE_PROJECT_ID': project_id,
                'GOOGLE_DRIVE_PRIVATE_KEY_ID': private_key_id,
                'GOOGLE_DRIVE_PRIVATE_KEY': private_key,
                'GOOGLE_DRIVE_CLIENT_EMAIL': client_email
            }.items() if not v]
            raise HTTPException(
                status_code=500,
                detail=f"Missing Google Drive credentials in .env file: {', '.join(missing)}"
            )
        
        # Construct credentials object from environment variables
        # Replace \\n with actual newlines in private key
        private_key_formatted = private_key.replace('\\n', '\n')
        
        credentials_dict = {
            "type": "service_account",
            "project_id": project_id,
            "private_key_id": private_key_id,
            "private_key": private_key_formatted,
            "client_email": client_email,
            "auth_uri": "https://accounts.google.com/o/oauth2/auth",
            "token_uri": "https://oauth2.googleapis.com/token",
            "auth_provider_x509_cert_url": "https://www.googleapis.com/oauth2/v1/certs",
            "client_x509_cert_url": f"https://www.googleapis.com/robot/v1/metadata/x509/{client_email}"
        }
        
        # Try to import Google Drive API libraries
        try:
            from google.oauth2 import service_account
            from googleapiclient.discovery import build
            from googleapiclient.http import MediaFileUpload
            GOOGLE_DRIVE_AVAILABLE = True
        except ImportError:
            GOOGLE_DRIVE_AVAILABLE = False
            raise HTTPException(
                status_code=503,
                detail="Google Drive API libraries not available. Install with: pip install google-auth google-auth-oauthlib google-auth-httplib2 google-api-python-client"
            )
        
        # Verify file exists
        file_path_obj = Path(file_path)
        if not file_path_obj.exists():
            raise HTTPException(status_code=404, detail=f"File not found: {file_path}")
        
        # Authenticate with Google Drive using service account credentials
        credentials = service_account.Credentials.from_service_account_info(
            credentials_dict,
            scopes=['https://www.googleapis.com/auth/drive.file']
        )
        
        # Build the Drive service
        drive_service = build('drive', 'v3', credentials=credentials)
        
        # Determine file name to use
        upload_file_name = file_name if file_name else file_path_obj.name
        
        # Prepare file metadata
        file_metadata = {
            'name': upload_file_name,
            'parents': [folder_id] if folder_id else []
        }
        
        # Upload file to Google Drive
        media = MediaFileUpload(
            str(file_path_obj),
            resumable=True
        )
        
        # Perform the upload
        file = drive_service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id, name, webViewLink'
        ).execute()
        
        # Return success response with file ID
        return {
            'success': True,
            'fileId': file.get('id'),
            'fileName': file.get('name'),
            'webViewLink': file.get('webViewLink'),
            'message': f"File successfully uploaded to Google Drive with ID: {file.get('id')}"
        }
        
    except HTTPException:
        # Re-raise HTTP exceptions
        raise
    except Exception as e:
        # Handle any other errors
        print(f"❌ Google Drive upload error: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to upload file to Google Drive: {str(e)}"
        )

# ============================================================================
# END GOOGLE DRIVE UPLOAD ENDPOINT
# ============================================================================

if __name__ == "__main__":
    # Start the server
    print("[START] Starting CATBot Proxy Server with File Operations...")
    print(f"[INFO] Scratch directory: {SCRATCH_DIR}")
    uvicorn.run(
        "proxy_server:app",
        host="0.0.0.0",
        port=8002,
        reload=True,
        log_level="info"
    )
