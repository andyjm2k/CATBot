#!/usr/bin/env python3
"""
Python FastAPI replacement for the Node.js proxy server.
Provides the same functionality but in Python for better integration with the MCP ecosystem.
"""

import asyncio
import json
import os
import re
import sys
import time
import base64
import hmac
import hashlib
import secrets
import glob
import socket
from typing import Dict, List, Optional, Any, Set, Tuple
from pathlib import Path
from datetime import datetime, timedelta, timezone
from io import BytesIO
from urllib.parse import urlparse, urlunparse

import httpx
from fastapi import Depends, FastAPI, Header, HTTPException, Request
from fastapi.exceptions import RequestValidationError
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from starlette.middleware.base import BaseHTTPMiddleware
from pydantic import BaseModel, ConfigDict
import uvicorn

# Import dotenv to load .env file
try:
    from dotenv import load_dotenv
    DOTENV_AVAILABLE = True
except ImportError:
    DOTENV_AVAILABLE = False
    print("[WARN] python-dotenv not available. Install with: pip install python-dotenv")

# Import file operations libraries
try:
    from docx import Document  # python-docx for Word documents
    import openpyxl  # openpyxl for Excel files
    from openpyxl.styles import Font, Alignment  # For Excel formatting
    import PyPDF2  # PyPDF2 for PDF reading
    from PIL import Image  # Pillow for image operations
    FILE_OPS_AVAILABLE = True
    print("[OK] File operations libraries loaded successfully")
except ImportError as e:
    print(f"[WARN] File operations libraries not available: {e}")
    FILE_OPS_AVAILABLE = False

# Import AutoGen components for team-based chat
try:
    from autogen_agentchat.teams import SelectorGroupChat
    from autogen_core import Component, FunctionCall, ComponentLoader
    AUTOGEN_AVAILABLE = True
    print("[OK] AutoGen imports successful")
except ImportError as e:
    print(f"[WARN] AutoGen not available: {e}")
    AUTOGEN_AVAILABLE = False
    SelectorGroupChat = None
    Component = None
    ComponentLoader = None

# Import MCP SDK (with error handling)
try:
    from mcp import ClientSession, stdio_client, StdioServerParameters
    MCP_AVAILABLE = True
except ImportError as e:
    print(f"MCP import error: {e}")
    MCP_AVAILABLE = False
    ClientSession = None
    stdio_client = None
    StdioServerParameters = None

# Browser-use HTTP server URL (must run: uv run mcp-server-browser-use server)
MCP_BROWSER_USE_HTTP_URL = os.environ.get("MCP_BROWSER_USE_HTTP_URL", "http://127.0.0.1:8383/mcp").strip()
BROWSER_USE_HTTP_UNAVAILABLE_MSG = (
    "Browser-use HTTP server not available. Start it with: uv run mcp-server-browser-use server (in mcp-browser-use directory)."
)

# Import memory system (with error handling)
MEMORY_AVAILABLE = False
MemoryManager = None
memory_import_error = None

try:
    # Check for required dependencies first
    try:
        import numpy
    except ImportError:
        raise ImportError("numpy is required for the memory system. Install with: pip install numpy")
    
    from src.memory import MemoryManager
    MEMORY_AVAILABLE = True
    print("[OK] Memory system imports successful")
except ImportError as e:
    memory_import_error = str(e)
    print(f"[WARN] Memory system not available: {e}")
    if "numpy" in str(e).lower():
        print("   Install numpy with: pip install numpy")
    MEMORY_AVAILABLE = False
    MemoryManager = None
except Exception as e:
    memory_import_error = str(e)
    print(f"[WARN] Memory system not available (unexpected error): {e}")
    MEMORY_AVAILABLE = False
    MemoryManager = None

# Import philosopher mode
try:
    from src.features.philosopher_mode import PhilosopherMode
    PHILOSOPHER_MODE_AVAILABLE = True
    print("[OK] Philosopher mode imports successful")
except ImportError as e:
    print(f"[WARN] Philosopher mode not available: {e}")
    PHILOSOPHER_MODE_AVAILABLE = False
    PhilosopherMode = None
except Exception as e:
    print(f"[WARN] Philosopher mode not available (unexpected error): {e}")
    PHILOSOPHER_MODE_AVAILABLE = False
    PhilosopherMode = None

# Load environment variables from .env file in project root
if DOTENV_AVAILABLE:
    # Load from project root (two levels up from src/servers/)
    env_path = Path(__file__).resolve().parent.parent.parent / '.env'
    if env_path.exists():
        load_dotenv(env_path)
        print(f"✅ Loaded environment variables from {env_path}")
    else:
        print(f"⚠️  No .env file found. Using system environment variables.")
        print(f"   Looked in: {env_path}")
else:
    print("⚠️  python-dotenv not available. Using system environment variables only.")

# Pydantic models for request/response validation
# Note: 'command' is intentionally not accepted from clients; only server-side presets are used.
class ServerConfig(BaseModel):
    model_config = ConfigDict(extra="forbid")  # Reject any extra fields (e.g. 'command') in request body

    id: Optional[str] = None
    name: Optional[str] = None
    preset_id: Optional[str] = None  # Required for add/update; must be in MCP_PRESETS
    apiKey: Optional[str] = None
    model: Optional[str] = None
    url: Optional[str] = None
    wsUrl: Optional[str] = None
    status: Optional[str] = None
    enabled: Optional[bool] = None
    action: Optional[str] = None

class ToolCallRequest(BaseModel):
    toolName: str
    parameters: Optional[Dict[str, Any]] = None

# Pydantic models for file operations
class ReadFileRequest(BaseModel):
    filename: str  # Name of the file to read

class WriteFileRequest(BaseModel):
    filename: str  # Name of the file to write
    content: str  # Content to write to the file
    format: Optional[str] = "txt"  # File format (txt, docx, xlsx, pdf)

class FileResponse(BaseModel):
    success: bool  # Whether the operation was successful
    message: str  # Human-readable message
    data: Optional[Dict[str, Any]] = None  # Optional data payload

class AuthSignupRequest(BaseModel):
    username: str
    password: str


class AuthLoginRequest(BaseModel):
    username: str
    password: str


class AuthTokenResponse(BaseModel):
    access_token: str
    token_type: str = "bearer"
    expires_in: int
    username: str


class AuthUserResponse(BaseModel):
    username: str
    created_at: str


class TelegramChatMessage(BaseModel):
    role: str
    content: str


class TelegramChatRequest(BaseModel):
    message: str
    conversation_id: Optional[str] = None
    user_id: Optional[str] = None
    history: Optional[List[TelegramChatMessage]] = None
    system_prompt: Optional[str] = None
    model: Optional[str] = None
    temperature: Optional[float] = None
    max_output_tokens: Optional[int] = None


class TelegramChatResponse(BaseModel):
    reply: str
    conversation_id: str
    usage: Optional[Dict[str, Any]] = None

# Pydantic models for memory operations
class MemoryStoreRequest(BaseModel):
    text: str
    category: Optional[str] = None
    source: Optional[str] = None
    metadata: Optional[Dict[str, Any]] = None

class MemorySearchRequest(BaseModel):
    query: str
    limit: Optional[int] = None
    similarity_threshold: Optional[float] = None
    category: Optional[str] = None

class MemoryExtractRequest(BaseModel):
    messages: List[Dict[str, str]]
    max_memories: Optional[int] = 3

class MemoryResponse(BaseModel):
    success: bool
    message: str
    data: Optional[Dict[str, Any]] = None

# Pydantic models for philosopher mode
class PhilosopherStartRequest(BaseModel):
    conversation_id: Optional[str] = None
    user_id: Optional[str] = None

class PhilosopherStopRequest(BaseModel):
    conversation_id: Optional[str] = None
    user_id: Optional[str] = None

class PhilosopherContemplateRequest(BaseModel):
    conversation_id: Optional[str] = None
    user_id: Optional[str] = None
    question: Optional[str] = None  # Optional: if not provided, generate one

class PhilosopherResponse(BaseModel):
    success: bool
    message: str
    data: Optional[Dict[str, Any]] = None


class ProxyFetchRequest(BaseModel):
    """Request body for POST /v1/proxy/fetch (avoids URL length limits on iOS Safari)."""
    url: str


# Project root (two levels up from src/servers/proxy_server.py)
_PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent

# Global state (similar to the Node.js version)
mcp_clients = {}
mcp_servers = {}
SERVERS_FILE = _PROJECT_ROOT / "config" / "mcp_servers.json"

# Server-side allowlist: only these presets may be used for MCP server execution.
# Never execute user-supplied command strings; resolve command/args only from here.
MCP_PRESETS = {
    "browser-use": {"type": "inprocess"},  # No subprocess; handled in-process in call_tool/list_tools
    # Future: "stdio": {"type": "stdio", "command": sys.executable, "allowed_args": [["-m", "mcp_server_browser_use"], ...]},
}
TEAM_CONFIG_FILE = _PROJECT_ROOT / "config" / "team-config.json"
# Optional: same system prompt / rules as web UI; when present, used for Telegram (overrides TELEGRAM_SYSTEM_PROMPT env)
CATBOT_SYSTEM_PROMPT_FILE = _PROJECT_ROOT / "config" / "catbot_system_prompt.txt"
SCRATCH_DIR = _PROJECT_ROOT / "scratch"

# Allowed file extensions for scratch file operations (path traversal mitigation)
READ_ALLOWED_EXTENSIONS = {".txt", ".docx", ".xlsx", ".xls", ".pdf", ".png", ".jpg", ".jpeg"}
WRITE_ALLOWED_EXTENSIONS = {".txt", ".docx", ".xlsx", ".xls", ".pdf"}
# Allowed extensions for Google Drive upload (scratch workspace only; path exfiltration mitigation)
DRIVE_UPLOAD_EXTENSIONS = {".txt", ".docx", ".xlsx", ".xls", ".pdf", ".png", ".jpg", ".jpeg"}
# Max file size for read/write in bytes (10MB default), configurable via env
FILE_OPS_MAX_SIZE_BYTES = int(os.getenv("FILE_OPS_MAX_SIZE", "10485760"))

# Telegram chat session storage (simple in-memory cache)
telegram_conversations: Dict[str, List[Dict[str, str]]] = {}

# Philosopher mode state storage (per conversation)
philosopher_mode_active: Dict[str, bool] = {}
philosopher_mode_instances: Dict[str, Any] = {}

# Assistant context: timezone and knowledge-gap awareness (prepended to all chat system prompts)
def _get_assistant_context_block() -> str:
    """Returns context block with server timezone and knowledge-awareness instructions."""
    tz_str = datetime.now().astimezone().strftime("%Z (UTC%z)")
    return (
        f"Context: You are running in timezone {tz_str}. Use this when interpreting dates/times unless the user specifies otherwise.\n"
        "Knowledge awareness: Your training has a cutoff. Acknowledge your knowledge gap; do not assume the current year or recent events. "
        "When the user provides current facts, corrections, or information that differs from your training "
        '(e.g., "it\'s 2025 now", "that API changed"), accept them as authoritative and do not contradict them.\n\n'
    )


# Telegram/OpenAI configuration
def _parse_telegram_history_limit() -> int:
    """Parse TELEGRAM_HISTORY_LIMIT with default 12 on invalid value."""
    try:
        return max(1, int(os.getenv("TELEGRAM_HISTORY_LIMIT", "12")))
    except ValueError:
        return 12


def _parse_telegram_chat_timeout() -> float:
    """Parse TELEGRAM_CHAT_TIMEOUT with default 30 on invalid value."""
    try:
        return max(1.0, float(os.getenv("TELEGRAM_CHAT_TIMEOUT", "30")))
    except ValueError:
        return 30.0


TELEGRAM_DEFAULT_MODEL = os.getenv("TELEGRAM_MODEL") or os.getenv("OPENAI_MODEL") or os.getenv("MCP_LLM_MODEL_NAME", "gpt-4o-mini")
TELEGRAM_SYSTEM_PROMPT_ENV = os.getenv(
    "TELEGRAM_SYSTEM_PROMPT",
    "You are CATBot, a helpful AI assistant that responds concisely for Telegram users.",
)


def _get_telegram_system_prompt_base() -> str:
    """Return the base system prompt for Telegram: config file if present, else TELEGRAM_SYSTEM_PROMPT env."""
    if CATBOT_SYSTEM_PROMPT_FILE.exists():
        try:
            content = CATBOT_SYSTEM_PROMPT_FILE.read_text(encoding="utf-8").strip()
            if content:
                return content
        except Exception as e:
            print(f"Warning: Could not read {CATBOT_SYSTEM_PROMPT_FILE}: {e}")
    return TELEGRAM_SYSTEM_PROMPT_ENV
TELEGRAM_HISTORY_LIMIT = _parse_telegram_history_limit()
TELEGRAM_CHAT_TIMEOUT = _parse_telegram_chat_timeout()
TELEGRAM_OPENAI_BASE_URL = (
    os.getenv("TELEGRAM_OPENAI_BASE_URL")
    or os.getenv("OPENAI_API_BASE")
    or os.getenv("MCP_LLM_OPENAI_ENDPOINT")
    or "https://api.openai.com/v1"
)
TELEGRAM_OPENAI_CHAT_PATH = os.getenv("TELEGRAM_OPENAI_CHAT_PATH", "/chat/completions")
OPENAI_ORG_ID = os.getenv("OPENAI_ORG_ID") or os.getenv("OPENAI_ORGANIZATION")
OPENAI_PROJECT_ID = os.getenv("OPENAI_PROJECT_ID")
# Optional shared secret for bot-to-proxy auth; when set, requests must include X-Telegram-Secret or Authorization: Bearer <secret>
TELEGRAM_SECRET = os.getenv("TELEGRAM_SECRET")

# Auth configuration
AUTH_USERS_FILE = _PROJECT_ROOT / "config" / "auth_users.json"
JWT_SECRET = os.getenv("JWT_SECRET", "change-this-secret-in-production")
JWT_ALGORITHM = "HS256"
JWT_EXPIRATION_SECONDS = int(os.getenv("JWT_EXPIRATION_SECONDS", "3600"))

# Simple in-memory user storage with JSON persistence
users_db: Dict[str, Dict[str, str]] = {}


def _base64url_encode(data: bytes) -> str:
    return base64.urlsafe_b64encode(data).decode().rstrip("=")


def _base64url_decode(data: str) -> bytes:
    padding = "=" * (-len(data) % 4)
    return base64.urlsafe_b64decode(f"{data}{padding}")


def hash_password(password: str, salt: str) -> str:
    hashed = hashlib.pbkdf2_hmac(
        "sha256",
        password.encode("utf-8"),
        salt.encode("utf-8"),
        100_000,
    )
    return _base64url_encode(hashed)


def create_password_record(password: str) -> Dict[str, str]:
    salt = secrets.token_hex(16)
    return {
        "salt": salt,
        "password_hash": hash_password(password, salt),
    }


def verify_password(password: str, salt: str, expected_hash: str) -> bool:
    candidate_hash = hash_password(password, salt)
    return hmac.compare_digest(candidate_hash, expected_hash)


def save_users_db() -> None:
    AUTH_USERS_FILE.write_text(json.dumps(users_db, indent=2), encoding="utf-8")


def load_users_db() -> None:
    global users_db
    if not AUTH_USERS_FILE.exists():
        users_db = {}
        return

    try:
        users_db = json.loads(AUTH_USERS_FILE.read_text(encoding="utf-8"))
    except Exception as e:
        print(f"⚠️ Failed to load users database: {e}")
        users_db = {}


def create_jwt(payload: Dict[str, Any], expires_in: int = JWT_EXPIRATION_SECONDS) -> str:
    now = datetime.now(timezone.utc)
    body = payload.copy()
    body["iat"] = int(now.timestamp())
    body["exp"] = int((now + timedelta(seconds=expires_in)).timestamp())

    header = {"alg": JWT_ALGORITHM, "typ": "JWT"}
    header_b64 = _base64url_encode(json.dumps(header, separators=(",", ":")).encode("utf-8"))
    payload_b64 = _base64url_encode(json.dumps(body, separators=(",", ":")).encode("utf-8"))

    signing_input = f"{header_b64}.{payload_b64}".encode("utf-8")
    signature = hmac.new(JWT_SECRET.encode("utf-8"), signing_input, hashlib.sha256).digest()
    signature_b64 = _base64url_encode(signature)
    return f"{header_b64}.{payload_b64}.{signature_b64}"


def decode_and_validate_jwt(token: str) -> Dict[str, Any]:
    try:
        header_b64, payload_b64, signature_b64 = token.split(".")
    except ValueError as exc:
        raise HTTPException(status_code=401, detail="Invalid token format") from exc

    signing_input = f"{header_b64}.{payload_b64}".encode("utf-8")
    expected_sig = hmac.new(JWT_SECRET.encode("utf-8"), signing_input, hashlib.sha256).digest()
    actual_sig = _base64url_decode(signature_b64)

    if not hmac.compare_digest(expected_sig, actual_sig):
        raise HTTPException(status_code=401, detail="Invalid token signature")

    try:
        payload = json.loads(_base64url_decode(payload_b64).decode("utf-8"))
    except Exception as exc:
        raise HTTPException(status_code=401, detail="Invalid token payload") from exc

    exp = payload.get("exp")
    if not isinstance(exp, int):
        raise HTTPException(status_code=401, detail="Token missing expiration")
    if datetime.now(timezone.utc).timestamp() > exp:
        raise HTTPException(status_code=401, detail="Token expired")

    return payload


def get_current_user_from_headers(authorization: Optional[str], x_auth_token: Optional[str]) -> Dict[str, Any]:
    auth_value = authorization
    if not auth_value and x_auth_token:
        auth_value = f"Bearer {x_auth_token}"

    if not auth_value:
        raise HTTPException(status_code=401, detail="Missing Authorization header")

    # Strip whitespace and handle potential formatting issues
    auth_value = auth_value.strip()
    
    try:
        scheme, token = auth_value.split(" ", 1)
    except ValueError as exc:
        raise HTTPException(status_code=401, detail="Invalid Authorization header") from exc

    # Strip token whitespace as well
    token = token.strip()
    
    if scheme.lower() != "bearer":
        raise HTTPException(status_code=401, detail="Authorization scheme must be Bearer")

    # Debug logging for token issues
    if not token or len(token.split(".")) != 3:
        print(f"🔒 Token format issue - token length: {len(token) if token else 0}, parts: {len(token.split('.')) if token else 0}")
        print(f"   Token preview: {token[:50] if token else 'None'}...")
        raise HTTPException(status_code=401, detail="Invalid token format")

    payload = decode_and_validate_jwt(token)
    username = payload.get("sub")
    if not username or username not in users_db:
        raise HTTPException(status_code=401, detail="Invalid token subject")

    return {"username": username, **users_db[username]}


def get_current_user(
    authorization: Optional[str] = Header(default=None),
    x_auth_token: Optional[str] = Header(default=None, alias="X-Auth-Token"),
) -> Dict[str, Any]:
    return get_current_user_from_headers(authorization, x_auth_token)


def security_log(action: str, user: str, server_id: Optional[str], detail: str) -> None:
    """Log MCP config/connect actions for audit; do not log secrets."""
    ts = datetime.now(timezone.utc).isoformat()
    sid = server_id or ""
    print(f"[SEC] {ts} action={action} user={user} server_id={sid} detail={detail}", flush=True)


# Helper utilities for Telegram integration
def build_openai_url(path: str) -> str:
    """Return absolute URL for OpenAI-compatible endpoints."""

    if path.startswith("http://") or path.startswith("https://"):
        return path

    base = TELEGRAM_OPENAI_BASE_URL.rstrip("/")
    if not path.startswith("/"):
        path = f"/{path}"
    return f"{base}{path}"


def trim_telegram_history(history: List[Dict[str, str]]) -> None:
    """Trim stored history to the configured limit (in-place)."""

    max_messages = max(TELEGRAM_HISTORY_LIMIT, 1) * 2
    if len(history) > max_messages:
        del history[: len(history) - max_messages]


def _validate_telegram_secret(request: Request) -> None:
    """If TELEGRAM_SECRET is set, require X-Telegram-Secret or Authorization Bearer to match; else raise 401."""
    if not TELEGRAM_SECRET:
        return
    secret_header = request.headers.get("X-Telegram-Secret")
    auth_header = request.headers.get("Authorization")
    token = None
    if secret_header is not None and secret_header.strip() == TELEGRAM_SECRET:
        token = TELEGRAM_SECRET
    if auth_header and auth_header.strip().startswith("Bearer "):
        token = auth_header.strip()[7:].strip()
    if token != TELEGRAM_SECRET:
        raise HTTPException(status_code=401, detail="Telegram secret required or invalid")


# Create scratch directory if it doesn't exist
SCRATCH_DIR.mkdir(parents=True, exist_ok=True)
load_users_db()

# Global AutoGen team instance
autogen_team = None

# Global memory manager instance
memory_manager = None

# Initialize memory manager if available
if MEMORY_AVAILABLE:
    try:
        memory_enabled = os.getenv("MEMORY_ENABLED", "true").lower() == "true"
        if memory_enabled:
            memory_manager = MemoryManager()
            print(f"✅ Memory system initialized with {memory_manager.count()} existing memories")
        else:
            print("⚠️  Memory system disabled via MEMORY_ENABLED=false")
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"⚠️  Failed to initialize memory system: {e}")
        print(f"   Full traceback:\n{error_trace}")
        memory_manager = None

# MCP Client Manager class to handle transport lifecycle
class MCPClientManager:
    """Manages MCP client and transport lifecycle."""

    def __init__(self, server_config: Dict[str, Any]):
        self.server_config = server_config
        self.client = None
        self.transport = None
        self.read = None
        self.write = None

    async def connect(self):
        """Connect to MCP server using server-side allowlisted preset only; never execute user-supplied command."""
        if not MCP_AVAILABLE:
            raise Exception("MCP SDK not available")

        preset_id = self.server_config.get("preset_id")
        if not preset_id or preset_id not in MCP_PRESETS:
            raise ValueError(
                "Server has no valid preset_id. Reconfigure the server with a preset (e.g. preset_id: 'browser-use')."
            )
        preset = MCP_PRESETS[preset_id]
        if preset.get("type") == "inprocess":
            # browser-use: connect is handled in connect_server without calling this code path
            raise ValueError(
                f"Preset '{preset_id}' is inprocess; connect is handled separately. Do not call MCPClientManager.connect for this preset."
            )
        if preset.get("type") != "stdio":
            raise ValueError(
                f"Preset '{preset_id}' has no executable command. Reconfigure the server with a valid stdio preset."
            )

        # Resolve command and args only from allowlist; never use server_config['command']
        command = preset.get("command")
        allowed_args = preset.get("allowed_args", [])
        if not command or not allowed_args:
            raise ValueError(f"Preset '{preset_id}' has no command/args defined in MCP_PRESETS.")
        # Use first allowed_args entry for this preset (e.g. ["-m", "mcp_server_browser_use"])
        args = list(allowed_args[0]) if isinstance(allowed_args[0], (list, tuple)) else []

        try:
            print(f"🔧 Creating MCP client with command: {command} and args: {args}")

            # Prepare environment variables from server config (apiKey, model only)
            env = os.environ.copy()
            if self.server_config.get("apiKey"):
                model = self.server_config.get("model", "").lower()
                if "gemini" in model:
                    env["GOOGLE_API_KEY"] = self.server_config["apiKey"]
                    env["MCP_MODEL_PROVIDER"] = "google"
                elif "claude" in model:
                    env["ANTHROPIC_API_KEY"] = self.server_config["apiKey"]
                    env["MCP_MODEL_PROVIDER"] = "anthropic"
                else:
                    env["OPENAI_API_KEY"] = self.server_config["apiKey"]
                    env["MCP_MODEL_PROVIDER"] = "openai"
            if self.server_config.get("model"):
                env["MCP_MODEL_NAME"] = self.server_config["model"]
            env.setdefault("BROWSER_USE_HEADLESS", "true")
            env.setdefault("BROWSER_USE_DISABLE_SECURITY", "false")

            print(f"🔐 Environment variables for MCP server: {list(env.keys())}")

            server_params = StdioServerParameters(command=command, args=args, env=env)

            try:
                transport_cm = stdio_client(server_params)
                async with transport_cm as (self.read, self.write):
                    self.client = ClientSession(self.read, self.write)
                    await self.client.initialize()
                    print("✅ MCP client setup complete")
                    return self.client
            except Exception as e:
                print(f"Failed to create stdio client: {e}")
                raise

        except Exception as e:
            print(f"MCP connection error: {e}")
            raise Exception(f"Failed to connect to MCP server: {str(e)}")

    async def disconnect(self):
        """Disconnect from MCP server."""
        if self.client:
            await self.client.close()
        if self.transport:
            # The transport context manager will handle cleanup
            pass

# FastAPI app
app = FastAPI(title="CATBot Proxy Server", version="2.0.0")

# Startup event to verify app initialization
@app.on_event("startup")
async def startup_event():
    """Log that the application has started successfully."""
    import sys
    print("🚀 FastAPI application startup event fired", flush=True)
    sys.stdout.flush()
    print(f"🚀 App routes registered: {len(app.routes)} routes", flush=True)
    sys.stdout.flush()
    # List all registered routes
    for route in app.routes:
        if hasattr(route, 'path') and hasattr(route, 'methods'):
            print(f"   Route: {list(route.methods)} {route.path}", flush=True)
            sys.stdout.flush()

# Request logging middleware to debug CORS issues
class RequestLoggingMiddleware(BaseHTTPMiddleware):
    """Middleware to log all incoming requests for debugging."""
    async def dispatch(self, request: Request, call_next):
        import sys
        # Safely log the incoming request with error handling
        try:
            method = getattr(request, 'method', 'UNKNOWN')
            path = getattr(request.url, 'path', 'unknown') if hasattr(request, 'url') else 'unknown'
            query = getattr(request.url, 'query', '') if hasattr(request, 'url') and hasattr(request.url, 'query') else ''
            origin = request.headers.get('origin', 'none') if hasattr(request, 'headers') else 'none'
            print(f"🌐 [{method}] {path}?{query}", flush=True)
            sys.stdout.flush()
            print(f"   Origin: {origin}", flush=True)
            sys.stdout.flush()
            if hasattr(request, 'headers'):
                print(f"   Headers: {dict(request.headers)}", flush=True)
                sys.stdout.flush()
        except Exception as log_error:
            # If logging fails, continue anyway - don't break the request
            print(f"⚠️ Error logging request: {log_error}", flush=True)
            sys.stdout.flush()
            import traceback
            print(traceback.format_exc(), flush=True)
            sys.stdout.flush()
        
        try:
            # Process the request
            response = await call_next(request)
            # Log successful response
            try:
                method = getattr(request, 'method', 'UNKNOWN')
                path = getattr(request.url, 'path', 'unknown') if hasattr(request, 'url') else 'unknown'
                status = getattr(response, 'status_code', 'unknown')
                print(f"✅ [{method}] {path} -> {status}", flush=True)
                sys.stdout.flush()
            except Exception as log_err:
                print(f"⚠️ Error logging response: {log_err}", flush=True)
                sys.stdout.flush()
            return response
        except Exception as e:
            # Log the exception
            try:
                method = getattr(request, 'method', 'UNKNOWN')
                path = getattr(request.url, 'path', 'unknown') if hasattr(request, 'url') else 'unknown'
                print(f"❌ [{method}] {path} -> Exception: {e}", flush=True)
                sys.stdout.flush()
                import traceback
                print(traceback.format_exc(), flush=True)
                sys.stdout.flush()
            except Exception:
                print("❌ Error in exception logging", flush=True)
                sys.stdout.flush()
            # Re-raise the exception so it can be handled by exception handlers
            raise

# Add request logging middleware first (before CORS)
app.add_middleware(RequestLoggingMiddleware)

# Add CORS middleware
# Allow specific origins for development and production
# Note: Using allow_credentials=False allows more flexible origin matching
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allow all origins for development (be more permissive to debug)
    allow_credentials=False,  # Set to False to allow more flexible CORS handling
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS", "PATCH"],
    allow_headers=["*", "Authorization", "X-Auth-Token", "Content-Type", "Accept"],  # Explicitly allow Authorization header
    expose_headers=["*"],
)


@app.middleware("http")
async def require_auth_for_v1_routes(request: Request, call_next):
    path = request.url.path

    if request.method == "OPTIONS":
        return await call_next(request)

    # Exempt certain routes from authentication (public endpoints)
    exempt_paths = {
        "/v1/auth/signup",
        "/v1/auth/login",
        "/v1/audio/transcriptions",  # Whisper endpoint - public for audio transcription
        "/v1/proxy/chat/completions",  # Chat completions proxy - public to avoid mixed content
        "/v1/proxy/models",  # Models list proxy - public to avoid mixed content
        "/v1/proxy/autogen",  # AutoGen workflow proxy - public to avoid mixed content
        "/v1/proxy/browser-agent",  # Browser automation proxy - public to avoid mixed content
        "/v1/proxy/deep-research",  # Deep research proxy - public to avoid mixed content
        "/v1/proxy/tts/voices",  # TTS voices endpoint - public
        "/v1/proxy/tts/speech",  # TTS speech endpoint - public
        "/v1/proxy/search",  # Search proxy - public
        "/v1/proxy/news",  # News proxy - public
        "/v1/proxy/fetch",  # Web fetch proxy - public
    }
    # Telegram bot endpoints are unauthenticated (bot uses TELEGRAM_SECRET when set)
    require_auth = (
        path.startswith("/v1/")
        and path not in exempt_paths
        and not path.startswith("/v1/telegram/chat")
    )
    if require_auth:
        try:
            # Get authorization header - FastAPI headers are case-insensitive, but check both for robustness
            # Use get() with case-insensitive lookup
            auth_header = None
            x_auth_token = None
            
            # Try to get authorization header (case-insensitive)
            for key, value in request.headers.items():
                if key.lower() == "authorization":
                    auth_header = value
                    break
                elif key.lower() == "x-auth-token":
                    x_auth_token = value
                    break
            
            # Fallback to direct get if not found in loop
            if not auth_header:
                auth_header = request.headers.get("authorization") or request.headers.get("Authorization")
            if not x_auth_token:
                x_auth_token = request.headers.get("x-auth-token") or request.headers.get("X-Auth-Token")
            
            # Debug logging for auth issues
            if not auth_header and not x_auth_token:
                print(f"🔒 Auth check failed for {path}: No authorization header found")
                print(f"   Available headers: {list(request.headers.keys())}")
            elif auth_header:
                # Log token preview for debugging (first 50 chars)
                token_preview = auth_header[:50] + "..." if len(auth_header) > 50 else auth_header
                print(f"🔒 Auth check for {path}: Found auth header (length: {len(auth_header)}, preview: {token_preview})")
            
            get_current_user_from_headers(
                auth_header,
                x_auth_token,
            )
        except HTTPException as exc:
            print(f"🔒 Auth check failed for {path}: {exc.detail}")
            # Log the actual header value for debugging (truncated)
            auth_debug = request.headers.get("authorization") or request.headers.get("Authorization") or "None"
            if auth_debug != "None":
                print(f"   Auth header value (first 100 chars): {auth_debug[:100]}")
            # Include CORS headers in error response
            cors_headers = build_cors_headers(request)
            return JSONResponse(
                status_code=exc.status_code,
                content={"detail": exc.detail},
                headers=cors_headers
            )

    return await call_next(request)

def build_cors_headers(request: Request) -> Dict[str, str]:
    """Build CORS headers for the request origin. Supports both localhost and remote access."""
    try:
        # Safely get origin from request headers, with fallback
        if hasattr(request, 'headers'):
            origin = request.headers.get("origin")
            # If no origin header (e.g., same-origin request), allow all origins for network access
            if not origin:
                origin = "*"
        else:
            origin = "*"
    except Exception:
        # If we can't access headers, allow all origins for network access
        origin = "*"
    
    return {
        "Access-Control-Allow-Origin": origin,
        "Access-Control-Allow-Methods": "GET, POST, PUT, DELETE, OPTIONS, PATCH",
        "Access-Control-Allow-Headers": "*",
    }

# Global exception handler to ensure CORS headers are always included
@app.exception_handler(HTTPException)
async def http_exception_handler(request: Request, exc: HTTPException):
    """Handle HTTP exceptions and ensure CORS headers are included."""
    # Safely build CORS headers - if this fails, use minimal headers
    try:
        cors_headers = build_cors_headers(request)
    except Exception as header_error:
        print(f"⚠️ Error building CORS headers in HTTPException handler: {header_error}")
        # Use minimal safe headers if build_cors_headers fails
        cors_headers = {
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Methods": "GET, POST, PUT, DELETE, OPTIONS, PATCH",
            "Access-Control-Allow-Headers": "*",
        }
    
    return JSONResponse(
        status_code=exc.status_code,
        content={"detail": exc.detail},
        headers=cors_headers,
    )

@app.exception_handler(RequestValidationError)
async def validation_exception_handler(request: Request, exc: RequestValidationError):
    """Handle validation errors and ensure CORS headers are included."""
    # Safely build CORS headers - if this fails, use minimal headers
    try:
        cors_headers = build_cors_headers(request)
    except Exception as header_error:
        print(f"⚠️ Error building CORS headers in ValidationException handler: {header_error}")
        # Use minimal safe headers if build_cors_headers fails
        cors_headers = {
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Methods": "GET, POST, PUT, DELETE, OPTIONS, PATCH",
            "Access-Control-Allow-Headers": "*",
        }
    
    return JSONResponse(
        status_code=422,
        content={"detail": exc.errors()},
        headers=cors_headers,
    )

@app.exception_handler(Exception)
async def general_exception_handler(request: Request, exc: Exception):
    """Handle all other exceptions and ensure CORS headers are included."""
    import sys
    import traceback
    error_trace = traceback.format_exc()
    print(f"❌ Unhandled exception in general_exception_handler: {exc}", flush=True)
    sys.stdout.flush()
    print(error_trace, flush=True)
    sys.stdout.flush()
    
    # Safely build CORS headers - if this fails, use minimal headers
    try:
        cors_headers = build_cors_headers(request)
    except Exception as header_error:
        print(f"⚠️ Error building CORS headers: {header_error}", flush=True)
        sys.stdout.flush()
        import traceback
        print(traceback.format_exc(), flush=True)
        sys.stdout.flush()
        # Use minimal safe headers if build_cors_headers fails
        cors_headers = {
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Methods": "GET, POST, PUT, DELETE, OPTIONS, PATCH",
            "Access-Control-Allow-Headers": "*",
        }
    
    try:
        response = JSONResponse(
            status_code=500,
            content={"detail": f"Internal server error: {str(exc)}"},
            headers=cors_headers,
        )
        print(f"✅ Created error response with status 500", flush=True)
        sys.stdout.flush()
        return response
    except Exception as response_error:
        print(f"❌ Error creating error response: {response_error}", flush=True)
        sys.stdout.flush()
        import traceback
        print(traceback.format_exc(), flush=True)
        sys.stdout.flush()
        # Last resort - return a simple response
        from fastapi.responses import PlainTextResponse
        return PlainTextResponse(
            content=f"Internal server error: {str(exc)}",
            status_code=500,
            headers=cors_headers,
        )

# Helper function to clean HTML text (similar to Node.js version)
def clean_text(text: str) -> str:
    """Clean HTML text by removing tags and decoding entities."""
    if not text:
        return ""

    # Remove HTML tags
    text = re.sub(r'</?[^>]+(>|$)', '', text)

    # Decode HTML entities
    html_entities = {
        '&amp;': '&', '&lt;': '<', '&gt;': '>', '&quot;': '"', '&#039;': "'",
        '&rsquo;': "'", '&lsquo;': "'", '&rdquo;': '"', '&ldquo;': '"',
        '&ndash;': '-', '&mdash;': '—'
    }

    for entity, replacement in html_entities.items():
        text = text.replace(entity, replacement)

    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

# Helper function to parse dates (similar to Node.js version)
def parse_date(date_str: Optional[str]) -> Optional[float]:
    """Parse date string to timestamp."""
    if not date_str:
        return None
    try:
        return datetime.fromisoformat(date_str.replace('Z', '+00:00')).timestamp()
    except (ValueError, AttributeError):
        return None

# Load servers from disk; migrate legacy 'command' to preset_id and never retain command
def load_servers():
    """Load MCP servers from JSON file. Migrate legacy config: set preset_id from name, drop command."""
    global mcp_servers
    try:
        if SERVERS_FILE.exists():
            with open(SERVERS_FILE, "r", encoding="utf-8") as f:
                servers = json.load(f)
            result = {}
            for server in servers:
                sid = server.get("id")
                if not sid:
                    continue
                # Migrate: if has command but no preset_id, infer preset_id from name
                if server.get("command") and not server.get("preset_id"):
                    if "mcp-browser-use" in (server.get("name") or "").lower():
                        server["preset_id"] = "browser-use"
                    else:
                        server["preset_id"] = None  # Legacy non-browser; not connectable until reconfigured
                # Never retain command in memory
                server.pop("command", None)
                result[sid] = server
            mcp_servers = result
            print(f"Loaded {len(mcp_servers)} MCP servers from disk")
    except Exception as e:
        print(f"No existing servers file found, starting with empty state: {e}")

# Save servers to disk; never persist 'command'
def save_servers():
    """Save MCP servers to disk. Only persist safe keys; never write command."""
    global mcp_servers
    try:
        servers = []
        for s in mcp_servers.values():
            safe = {k: s[k] for k in MCP_SERVER_SAFE_KEYS if k in s}
            servers.append(safe)
        with open(SERVERS_FILE, "w", encoding="utf-8") as f:
            json.dump(servers, f, indent=2, ensure_ascii=False)
        print(f"Saved {len(servers)} MCP servers to disk")
    except Exception as e:
        print(f"Error saving servers to disk: {e}")

# Load AutoGen team from config
def load_autogen_team():
    """Load AutoGen team from team-config.json."""
    global autogen_team
    
    if not AUTOGEN_AVAILABLE:
        print("⚠️  AutoGen not available, skipping team load")
        return None
    
    try:
        if not TEAM_CONFIG_FILE.exists():
            print(f"⚠️  Team config file not found: {TEAM_CONFIG_FILE}")
            return None
            
        print(f"📂 Loading AutoGen team from {TEAM_CONFIG_FILE}...")
        
        with open(TEAM_CONFIG_FILE, 'r', encoding='utf-8') as f:
            team_config = json.load(f)
        
        # Load the team from the configuration using ComponentLoader
        loader = ComponentLoader()
        team = loader.load_component(team_config)
        
        print(f"✅ AutoGen team loaded successfully: {team_config.get('label', 'Unknown')}")
        return team
        
    except Exception as e:
        import traceback
        print(f"❌ Error loading AutoGen team: {e}")
        print(traceback.format_exc())
        return None

# Load servers on startup (with error handling to prevent startup failures)
try:
    load_servers()
    print(f"✅ Loaded {len(mcp_servers)} MCP servers from disk")
except Exception as e:
    import traceback
    print(f"⚠️ Warning: Could not load servers on startup: {e}")
    print(traceback.format_exc())
    # Continue anyway - server should still work without pre-loaded servers

# Load AutoGen team on startup (with error handling to prevent startup failures)
try:
    autogen_team = load_autogen_team()
    if autogen_team is not None:
        print("✅ AutoGen team loaded successfully on startup")
except Exception as e:
    import traceback
    print(f"⚠️ Warning: Could not load AutoGen team on startup: {e}")
    print(traceback.format_exc())
    # Continue anyway - server should still work without AutoGen team
    autogen_team = None

# Web proxy endpoint for fetching content (GET for backward compat, POST for iOS Safari / long URLs)
def _is_dns_or_network_error(exc: BaseException) -> bool:
    """True if the exception is DNS (getaddrinfo) or network unreachable."""
    if isinstance(exc, socket.gaierror):
        return True
    # Windows: OSError can have winerror 11002 (WSAHOST_NOT_FOUND)
    if isinstance(exc, OSError) and getattr(exc, "winerror", None) == 11002:
        return True
    # Check wrapped cause (e.g. httpx wraps socket errors)
    cause = getattr(exc, "__cause__", None)
    if cause and _is_dns_or_network_error(cause):
        return True
    errstr = str(exc).lower()
    if "getaddrinfo failed" in errstr or "name or service not known" in errstr or "nodename nor servname" in errstr or "errno 11002" in errstr:
        return True
    return False


async def _do_proxy_fetch(url: str) -> Dict[str, str]:
    """Shared fetch logic: fetch URL and return dict with content or raise."""
    if not url or not url.strip():
        raise HTTPException(status_code=400, detail="URL parameter is required")
    # Normalize URL (allow without scheme for convenience)
    url = url.strip()
    if not url.startswith(("http://", "https://")):
        url = "https://" + url
    headers = {
        "User-Agent": "Mozilla/5.0 (iPhone; CPU iPhone OS 17_0 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.0 Mobile/15E148 Safari/604.1",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.9",
        "Connection": "keep-alive",
    }
    try:
        async with httpx.AsyncClient(timeout=15.0, follow_redirects=True) as client:
            response = await client.get(url, headers=headers)
        response.raise_for_status()
        return {"content": response.text}
    except HTTPException:
        raise
    except Exception as e:
        # DNS or network failure on the machine running the proxy (e.g. WSL, VPN, no outbound DNS)
        if _is_dns_or_network_error(e):
            raise HTTPException(
                status_code=502,
                detail=(
                    "The proxy server could not resolve the website's hostname (DNS lookup failed). "
                    "This usually means the machine running the proxy has no internet or restricted DNS. "
                    "Ensure the proxy runs on a machine with working internet and DNS (e.g. try pinging the host from that machine)."
                ),
            )
        raise HTTPException(status_code=500, detail=f"Failed to fetch content: {str(e)}")


@app.get("/v1/proxy/fetch")
async def proxy_fetch_get(url: str, request: Request):
    """Fetch web content via GET (query param). Use POST for long URLs (e.g. iOS Safari)."""
    try:
        result = await _do_proxy_fetch(url)
        cors = build_cors_headers(request)
        return JSONResponse(content=result, headers=cors)
    except HTTPException:
        raise
    except Exception as e:
        print(f"Proxy fetch error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch content: {str(e)}")


@app.post("/v1/proxy/fetch")
async def proxy_fetch_post(body: ProxyFetchRequest, request: Request):
    """Fetch web content via POST body. Avoids URL length limits on iOS Safari."""
    try:
        result = await _do_proxy_fetch(body.url)
        cors = build_cors_headers(request)
        return JSONResponse(content=result, headers=cors)
    except HTTPException:
        raise
    except Exception as e:
        print(f"Proxy fetch error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch content: {str(e)}")

# Web search endpoint
@app.get("/v1/proxy/search")
async def proxy_search(query: str):
    """Search the web using Brave Search API or DuckDuckGo fallback."""
    if not query:
        raise HTTPException(status_code=400, detail="Search query is required")

    # Try Brave Search first
    brave_api_key = os.getenv('BRAVE_API_KEY')
    if not brave_api_key:
        print("⚠️  BRAVE_API_KEY not configured. Falling back to DuckDuckGo.")
    else:
        try:
            print(f"🔍 Using Brave Search API for query: {query[:50]}...")
            async with httpx.AsyncClient(timeout=15.0) as client:
                response = await client.get(
                    'https://api.search.brave.com/res/v1/web/search',
                    headers={
                        'Accept': 'application/json',
                        'Accept-Encoding': 'gzip',
                        'X-Subscription-Token': brave_api_key
                    },
                    params={
                        'q': query,
                        'count': 10,
                        'search_lang': 'en',
                        'safesearch': 'moderate',
                        'freshness': 'past_month'
                    }
                )

            if response.status_code == 200:
                data = response.json()
                if data.get('web', {}).get('results'):
                    results = []
                    for result in data['web']['results']:
                        # Parse date information
                        date_str = result.get('age') or result.get('published')
                        parsed_date = parse_date(date_str) if date_str else None

                        results.append({
                            'url': result['url'],
                            'title': clean_text(result.get('title', '')),
                            'snippet': clean_text(result.get('description', '')),
                            'date': parsed_date
                        })

                    # Sort by date (newest first) and filter valid results
                    results = [r for r in results if r['title'] and r['snippet']]
                    results.sort(key=lambda x: parse_date(x.get('date')) or 0, reverse=True)

                    print(f"✅ Brave Search returned {len(results)} results")
                    return {"results": results[:5], "source": "brave"}  # Return top 5 results
                else:
                    print(f"⚠️  Brave Search returned no results in response")
            elif response.status_code == 401:
                print(f"❌ Brave Search API authentication failed (401). Check your BRAVE_API_KEY.")
                raise HTTPException(
                    status_code=500,
                    detail="Brave Search API authentication failed. Please check your BRAVE_API_KEY configuration."
                )
            elif response.status_code == 429:
                print(f"⚠️  Brave Search API rate limit exceeded (429). Falling back to DuckDuckGo.")
            else:
                print(f"⚠️  Brave Search API returned status {response.status_code}. Falling back to DuckDuckGo.")
                try:
                    error_data = response.json()
                    print(f"   Error details: {error_data}")
                except:
                    print(f"   Error text: {response.text[:200]}")

        except httpx.RequestError as e:
            error_msg = str(e) if str(e) else f"Network error: {type(e).__name__}"
            print(f"❌ Brave Search network error: {error_msg}")
            print("   Falling back to DuckDuckGo...")
        except httpx.HTTPStatusError as e:
            print(f"❌ Brave Search HTTP error: {e.response.status_code if e.response else 'Unknown'}")
            print("   Falling back to DuckDuckGo...")
        except HTTPException:
            # Re-raise HTTPExceptions (like auth errors) - don't fall back
            raise
        except Exception as e:
            error_msg = str(e) if str(e) else f"Unknown error: {type(e).__name__}"
            print(f"❌ Brave Search failed: {error_msg}")
            print("   Falling back to DuckDuckGo...")
            import traceback
            print(traceback.format_exc())

    # Fallback to DuckDuckGo (only if Brave Search is not available or failed)
    print("🦆 Falling back to DuckDuckGo search...")
    try:
        search_url = f"https://html.duckduckgo.com/html/?q={query}"

        async with httpx.AsyncClient(timeout=15.0) as client:
            response = await client.get(
                search_url,
                headers={
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                    'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
                    'Accept-Language': 'en-US,en;q=0.5',
                    'Referer': 'https://duckduckgo.com/'
                },
                follow_redirects=True
            )
            
            # Check response status
            if response.status_code != 200:
                raise HTTPException(
                    status_code=500,
                    detail=f"DuckDuckGo search returned HTTP {response.status_code}. The search service may be temporarily unavailable."
                )

        results = []
        html = response.text

        # Extract results using multiple regex patterns (DuckDuckGo HTML structure can vary)
        patterns = [
            # Pattern 1: Modern DuckDuckGo structure
            r'<div class="links_main links_deep result__body">.*?<a class="result__a" href="([^"]+)".*?>(.*?)</a>.*?<a class="result__snippet".*?>(.*?)</a>',
            # Pattern 2: Alternative structure
            r'<div class="result__body">.*?<a class="result__url" href="([^"]+)".*?>(.*?)</a>.*?<div class="result__snippet">(.*?)</div>',
            # Pattern 3: Another variant
            r'<div class="result__body">.*?<a class="result__a" href="([^"]+)".*?>(.*?)</a>.*?<div class="result__snippet">(.*?)</div>',
            # Pattern 4: Try to find any links with result classes
            r'<a[^>]*class="[^"]*result[^"]*"[^>]*href="([^"]+)"[^>]*>(.*?)</a>.*?<div[^>]*class="[^"]*snippet[^"]*"[^>]*>(.*?)</div>',
            # Pattern 5: More generic pattern
            r'<a[^>]*href="([^"]+)"[^>]*class="[^"]*result__a[^"]*"[^>]*>(.*?)</a>.*?<span[^>]*class="[^"]*result__snippet[^"]*"[^>]*>(.*?)</span>'
        ]

        for pattern in patterns:
            matches = re.finditer(pattern, html, re.DOTALL)
            for match in matches:
                if len(results) >= 5:  # Limit to 5 results
                    break

                try:
                    url, title, snippet = match.groups()
                    url = url.replace('&amp;', '&')
                    
                    # Clean up the extracted text
                    title = clean_text(title)
                    snippet = clean_text(snippet)

                    # Skip if URL is invalid or is a DuckDuckGo internal link
                    if url and 'duckduckgo.com' not in url and title and snippet:
                        results.append({
                            'url': url,
                            'title': title,
                            'snippet': snippet
                        })
                except (ValueError, IndexError) as e:
                    # Skip malformed matches
                    continue

            if len(results) >= 5:
                break

        # If no results found with regex, try a simpler approach
        if len(results) == 0:
            # Log the HTML snippet for debugging (first 1000 chars)
            print(f"⚠️  DuckDuckGo search: No results found with regex patterns. HTML preview: {html[:1000]}")
            # Return empty results rather than failing
            return {"results": [], "source": "duckduckgo", "message": "No results found. DuckDuckGo HTML structure may have changed."}

        print(f"✅ DuckDuckGo returned {len(results)} results")
        return {"results": results, "source": "duckduckgo"}

    except httpx.RequestError as e:
        error_msg = str(e) if str(e) else f"Network error: {type(e).__name__}"
        print(f"Search error (network): {error_msg}")
        raise HTTPException(status_code=500, detail=f"Failed to perform search: Network error - {error_msg}")
    except httpx.HTTPStatusError as e:
        error_msg = f"HTTP {e.response.status_code}: {e.response.text[:200]}" if e.response else str(e)
        print(f"Search error (HTTP): {error_msg}")
        raise HTTPException(status_code=500, detail=f"Failed to perform search: {error_msg}")
    except Exception as e:
        error_msg = str(e) if str(e) else f"Unknown error: {type(e).__name__}"
        print(f"Search error: {error_msg}")
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Failed to perform search: {error_msg}")

# News API search endpoint
@app.get("/v1/proxy/news")
async def proxy_news_search(query: str):
    """Search news articles using News API."""
    if not query:
        raise HTTPException(status_code=400, detail="Search query is required")

    # Get News API key from environment variable
    news_api_key = os.getenv('NEWS_API_KEY')
    if not news_api_key:
        raise HTTPException(
            status_code=503,
            detail="NEWS_API_KEY is not configured. Please set it in your .env file."
        )

    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            response = await client.get(
                'https://newsapi.org/v2/everything',
                headers={
                    'Accept': 'application/json',
                },
                params={
                    'q': query,
                    'apiKey': news_api_key,
                    'sortBy': 'publishedAt',
                    'language': 'en',
                    'pageSize': 100  # Maximum allowed by News API
                }
            )

        if response.status_code == 200:
            data = response.json()
            articles = data.get('articles', [])
            
            if not articles:
                return {
                    "success": False,
                    "message": f"No articles found for search term \"{query}\"",
                    "articles": []
                }

            # Format articles for response
            formatted_articles = []
            for article in articles:
                formatted_articles.append({
                    'title': clean_text(article.get('title', '')),
                    'url': article.get('url', ''),
                    'description': clean_text(article.get('description', '')),
                    'publishedAt': article.get('publishedAt', ''),
                    'source': article.get('source', {}).get('name', 'Unknown')
                })

            return {
                "success": True,
                "message": f"Found {len(formatted_articles)} articles",
                "articles": formatted_articles,
                "totalResults": data.get('totalResults', len(formatted_articles))
            }
        else:
            error_data = response.json() if response.headers.get('content-type', '').startswith('application/json') else {}
            error_message = error_data.get('message', f"News API returned status {response.status_code}")
            print(f"News API error: {error_message}")
            raise HTTPException(
                status_code=response.status_code,
                detail=f"News API error: {error_message}"
            )

    except httpx.HTTPStatusError as e:
        print(f"News API HTTP error: {e}")
        raise HTTPException(
            status_code=e.response.status_code,
            detail=f"News API request failed: {str(e)}"
        )
    except Exception as e:
        print(f"News API error: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to fetch news: {str(e)}"
        )

# AutoGen team chat endpoint (integrated directly)
@app.post("/v1/proxy/autogen")
async def autogen_chat(request: Request):
    """Run AutoGen team conversation directly (no separate service needed)."""
    global autogen_team
    
    try:
        body = await request.json()
        input_text = body.get('input')
        
        if not input_text:
            raise HTTPException(status_code=400, detail="Input parameter is required")

        print(f"🤖 AutoGen team request: {input_text[:100]}...")

        # Check if AutoGen is available and team is loaded
        if not AUTOGEN_AVAILABLE:
            raise HTTPException(
                status_code=503,
                detail="AutoGen not available. Please install: pip install autogen-agentchat autogen-ext"
            )
        
        if autogen_team is None:
            # Try to load the team
            print("🔄 Loading AutoGen team for the first time...")
            autogen_team = load_autogen_team()
            if autogen_team is None:
                raise HTTPException(
                    status_code=503,
                    detail="AutoGen team not loaded. Check team-config.json exists and is valid."
                )
        else:
            # Check if the team config file has been modified and reload if needed
            try:
                config_mtime = TEAM_CONFIG_FILE.stat().st_mtime
                if not hasattr(autogen_team, '_config_mtime') or autogen_team._config_mtime != config_mtime:
                    print("🔄 Team config file changed, reloading AutoGen team...")
                    new_team = load_autogen_team()
                    if new_team is not None:
                        autogen_team = new_team
                        autogen_team._config_mtime = config_mtime
                    else:
                        print("⚠️  Failed to reload team, using existing team")
            except Exception as e:
                print(f"⚠️  Error checking team config modification time: {e}")

        try:
            # Run the team conversation
            print(f"🚀 Running AutoGen team with input: {input_text[:100]}...")
            result = await autogen_team.run(task=input_text)
            
            # Extract messages and final answer from result
            messages = []
            if hasattr(result, 'messages'):
                messages = [
                    {
                        "source": msg.source if hasattr(msg, 'source') else 'unknown',
                        "content": msg.content if hasattr(msg, 'content') else str(msg)
                    }
                    for msg in result.messages
                ]
            
            # Format all messages into a readable conversation summary
            conversation_summary = "=== AutoGen Team Workflow ===\n\n"
            
            if messages:
                for i, msg in enumerate(messages, 1):
                    source = msg.get('source', 'unknown')
                    content = msg.get('content', '')
                    conversation_summary += f"[{i}] {source}:\n{content}\n\n"
                
                conversation_summary += "=== End of Workflow ===\n\n"
                conversation_summary += "Please review the above conversation and provide a concise summary of the final result."
            else:
                conversation_summary = "No messages returned from AutoGen team."
            
            print(f"✅ AutoGen team completed with {len(messages)} messages")
            
            return {
                "output": conversation_summary,
                "response": conversation_summary,  # Alternative field name for compatibility
                "messages": messages,
                "message_count": len(messages)
            }
            
        except Exception as e:
            import traceback
            error_trace = traceback.format_exc()
            print(f"❌ AutoGen team execution error: {e}")
            print(f"   Traceback: {error_trace}")
            raise HTTPException(
                status_code=500,
                detail=f"AutoGen team execution failed: {str(e)}"
            )

    except HTTPException:
        raise
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"❌ AutoGen endpoint error: {e}")
        print(f"   Traceback: {error_trace}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to process AutoGen request: {str(e)}"
        )

# Allowed keys when persisting MCP server config (never persist 'command')
MCP_SERVER_SAFE_KEYS = {"id", "name", "preset_id", "apiKey", "model", "url", "wsUrl", "status", "enabled"}


# MCP server management endpoints (all require authentication)
@app.post("/v1/mcp/servers")
async def manage_servers(
    server_config: ServerConfig,
    current_user: Dict[str, Any] = Depends(get_current_user),
):
    """Manage MCP servers (create, update, clear). Never persist or execute client-supplied command."""
    # Log without secrets
    safe_dict = {k: v for k, v in server_config.model_dump().items() if k != "apiKey"}
    print(f"Received server config: {safe_dict}")

    global mcp_clients, mcp_servers

    try:
        # Handle clear action
        if server_config.action == "clear":
            for server_id, client in list(mcp_clients.items()):
                try:
                    await client.close()
                except Exception as e:
                    print(f"Error closing client {server_id}: {e}")
            mcp_servers.clear()
            mcp_clients.clear()
            print("Cleared all MCP servers and clients")
            save_servers()
            security_log("mcp_clear", current_user.get("username", ""), None, "all servers cleared")
            return {"message": "All MCP servers cleared successfully"}

        # Validate required fields
        if not server_config.id or not server_config.name:
            raise HTTPException(status_code=400, detail="Missing required fields: id, name")

        # Resolve preset_id: require or infer from name (browser-use)
        preset_id = server_config.preset_id
        if not preset_id and server_config.name and "mcp-browser-use" in server_config.name.lower():
            preset_id = "browser-use"
        if not preset_id:
            raise HTTPException(
                status_code=400,
                detail="Missing preset_id. Provide preset_id (e.g. 'browser-use') or use a name that implies it.",
            )
        if preset_id not in MCP_PRESETS:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid preset_id '{preset_id}'. Allowed: {list(MCP_PRESETS.keys())}",
            )

        # Build stored server dict from allowed fields only; never store 'command'
        raw = server_config.model_dump()
        stored = {k: raw[k] for k in MCP_SERVER_SAFE_KEYS if k in raw and raw[k] is not None}
        stored["status"] = "disconnected"
        # Preserve connection status when updating
        existing_server = mcp_servers.get(server_config.id)
        if existing_server:
            stored["status"] = existing_server.get("status", "disconnected")
        stored["preset_id"] = preset_id

        if existing_server:
            mcp_servers[server_config.id] = {**existing_server, **stored}
            print(f"Updated MCP server: {server_config.name} ({server_config.id})")
            security_log("mcp_update", current_user.get("username", ""), server_config.id, f"preset_id={preset_id}")
        else:
            mcp_servers[server_config.id] = stored
            print(f"Added MCP server: {server_config.name} ({server_config.id})")
            security_log("mcp_add", current_user.get("username", ""), server_config.id, f"preset_id={preset_id}")

        save_servers()
        return {"message": "Server saved successfully"}

    except HTTPException:
        raise
    except Exception as e:
        print(f"Error saving MCP server: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to save MCP server: {str(e)}")

@app.get("/v1/mcp/servers")
async def get_servers(
    current_user: Dict[str, Any] = Depends(get_current_user),
):
    """Get all MCP servers."""
    try:
        servers = list(mcp_servers.values())
        return {"servers": servers}
    except Exception as e:
        print(f"Error getting MCP servers: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to get MCP servers: {str(e)}")

def setup_browser_llm():
    """Set up the language model for browser automation."""
    model_provider = os.getenv("MCP_MODEL_PROVIDER", "google").lower()
    model_name = os.getenv("MCP_MODEL_NAME", "gemini-flash-latest")
    temperature = float(os.getenv("MCP_TEMPERATURE", "0.1"))

    if model_provider == "openai":
        from langchain_openai import ChatOpenAI
        return ChatOpenAI(
            model=model_name,
            temperature=temperature,
            api_key=os.getenv("OPENAI_API_KEY")
        )
    elif model_provider == "anthropic":
        from langchain_anthropic import ChatAnthropic
        return ChatAnthropic(
            model=model_name,
            temperature=temperature,
            api_key=os.getenv("ANTHROPIC_API_KEY")
        )
    elif model_provider == "google":
        from langchain_google_genai import ChatGoogleGenerativeAI
        return ChatGoogleGenerativeAI(
            model=model_name,
            temperature=temperature,
            api_key=os.getenv("GOOGLE_API_KEY")
        )
    else:
        # Default to Google/Gemini
        from langchain_google_genai import ChatGoogleGenerativeAI
        return ChatGoogleGenerativeAI(
            model=model_name,
            temperature=temperature,
            api_key=os.getenv("GOOGLE_API_KEY", "dummy-key")
        )

async def create_mcp_client(server_config: Dict[str, Any]):
    """Create MCP client connection (similar to Node.js version)."""
    manager = MCPClientManager(server_config)
    client = await manager.connect()
    return client

@app.post("/v1/mcp/servers/{server_id}/connect")
async def connect_server(
    server_id: str,
    current_user: Dict[str, Any] = Depends(get_current_user),
):
    """Connect to an MCP server. Uses only server-side presets; never executes user-supplied command."""
    try:
        print(f"Attempting to connect to server: {server_id}")

        server = mcp_servers.get(server_id)
        if not server:
            print(f"Server not found: {server_id}")
            raise HTTPException(status_code=404, detail="Server not found")

        print(f"Found server: {server.get('name')} ({server_id})")

        # Inprocess preset (e.g. browser-use): no subprocess, mark connected
        preset_id = server.get("preset_id") or (
            "browser-use" if "mcp-browser-use" in (server.get("name") or "").lower() else None
        )
        if preset_id == "browser-use" or "mcp-browser-use" in (server.get("name") or "").lower():
            server["preset_id"] = preset_id or "browser-use"
            server["status"] = "connected"
            mcp_servers[server_id] = server
            print(f"Successfully connected to MCP Browser Use server: {server.get('name')}")
            security_log("mcp_connect", current_user.get("username", ""), server_id, "preset_id=browser-use")
            return {"message": "Server connected successfully"}

        if server_id in mcp_clients:
            print(f"Server already connected: {server_id}")
            raise HTTPException(status_code=409, detail="Server is already connected")

        if not preset_id or preset_id not in MCP_PRESETS:
            raise HTTPException(
                status_code=400,
                detail="Server has no valid preset_id. Reconfigure the server with a preset (e.g. preset_id: 'browser-use').",
            )
        if MCP_PRESETS[preset_id].get("type") != "stdio":
            raise HTTPException(
                status_code=400,
                detail=f"Preset '{preset_id}' does not support stdio connect. Use a stdio preset or browser-use.",
            )

        if not MCP_AVAILABLE:
            raise HTTPException(status_code=503, detail="MCP SDK not available")

        print(f"Creating MCP client for server: {server.get('name')}")
        client = await create_mcp_client(server)
        mcp_clients[server_id] = client

        server["status"] = "connected"
        mcp_servers[server_id] = server

        print(f"Successfully connected to MCP server: {server.get('name')}")
        security_log("mcp_connect", current_user.get("username", ""), server_id, f"preset_id={preset_id}")
        return {"message": "Server connected successfully"}

    except HTTPException:
        raise
    except ValueError as e:
        print(f"Error connecting to MCP server: {e}")
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        print(f"Error connecting to MCP server: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to connect to MCP server: {str(e)}")

@app.post("/v1/mcp/servers/{server_id}/disconnect")
async def disconnect_server(
    server_id: str,
    current_user: Dict[str, Any] = Depends(get_current_user),
):
    """Disconnect from an MCP server."""
    try:
        global mcp_clients, mcp_servers

        # Check if this is the MCP Browser Use server
        server = mcp_servers.get(server_id)
        if server and "mcp-browser-use" in server.get("name", "").lower():
            # Just mark as disconnected since we don't have a real connection
            server["status"] = "disconnected"
            mcp_servers[server_id] = server
            security_log("mcp_disconnect", current_user.get("username", ""), server_id, "browser-use")
            return {"message": "Server disconnected successfully"}

        if not MCP_AVAILABLE:
            raise HTTPException(status_code=503, detail="MCP SDK not available")

        client = mcp_clients.get(server_id)
        if not client:
            raise HTTPException(status_code=404, detail="Server is not connected")

        # Create a temporary manager to handle disconnect
        if server:
            manager = MCPClientManager(server)
            await manager.disconnect()

        del mcp_clients[server_id]

        server = mcp_servers.get(server_id)
        if server:
            server["status"] = "disconnected"
            mcp_servers[server_id] = server

        security_log("mcp_disconnect", current_user.get("username", ""), server_id, "stdio")
        return {"message": "Server disconnected successfully"}

    except HTTPException:
        raise
    except Exception as e:
        print(f"Error disconnecting MCP server: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to disconnect MCP server: {str(e)}")


async def _browser_use_http_list_tools() -> Dict[str, Any]:
    """List tools from the browser-use HTTP MCP server. Raises on connection failure."""
    from fastmcp import Client

    async with Client(MCP_BROWSER_USE_HTTP_URL) as client:
        tools = await client.list_tools()
    # Convert to the shape expected by the proxy API (name, description, inputSchema)
    tools_list = []
    for t in tools:
        entry = {
            "name": getattr(t, "name", "unknown"),
            "description": getattr(t, "description", None) or "",
            "inputSchema": getattr(t, "inputSchema", None) or {"type": "object", "properties": {}},
        }
        tools_list.append(entry)
    return {"tools": tools_list}


async def _browser_use_http_call_tool(tool_name: str, parameters: Dict[str, Any]) -> Dict[str, Any]:
    """Call a tool on the browser-use HTTP MCP server. Returns result with content list; raises on connection failure."""
    from fastmcp import Client

    # Map proxy parameter names to server names (e.g. instruction -> task for run_browser_agent)
    args = dict(parameters)
    if tool_name == "run_browser_agent" and "instruction" in args and "task" not in args:
        args["task"] = args.pop("instruction")

    async with Client(MCP_BROWSER_USE_HTTP_URL) as client:
        result = await client.call_tool(tool_name, args)

    # Build content list from result.content (list of items with .text or str)
    content = []
    for item in getattr(result, "content", []) or []:
        text = getattr(item, "text", None) or str(item)
        content.append({"type": "text", "text": text})
    if getattr(result, "is_error", False) and not content:
        content.append({"type": "text", "text": "Tool returned an error."})
    return {"content": content}


# Browser automation tool endpoint (browser-use via HTTP; other servers via MCP client)
@app.post("/v1/mcp/servers/{server_id}/tools/call")
async def call_tool(server_id: str, request: ToolCallRequest):
    """Call a tool on an MCP server. Browser-use preset uses HTTP client; others use connected MCP client."""
    if not MCP_AVAILABLE:
        raise HTTPException(status_code=503, detail="MCP SDK not available")

    try:
        print(f"🔧 [TOOLS/CALL] Server: {server_id}")

        server = mcp_servers.get(server_id)
        tool_name = request.toolName
        parameters = request.parameters or {}
        print(f"🔍 [TOOLS/CALL] Tool name: {tool_name}")
        print(f"🔍 [TOOLS/CALL] Parameters: {parameters}")

        # Browser-use preset: use HTTP client (no mcp_clients entry)
        if server and "mcp-browser-use" in server.get("name", "").lower():
            try:
                result = await _browser_use_http_call_tool(tool_name, parameters)
                return {"result": result}
            except Exception as e:
                print(f"❌ [TOOLS/CALL] Browser-use HTTP error: {e}")
                raise HTTPException(
                    status_code=503,
                    detail=BROWSER_USE_HTTP_UNAVAILABLE_MSG + " " + str(e),
                )

        client = mcp_clients.get(server_id)
        if not client:
            print(f"❌ [TOOLS/CALL] Server {server_id} not found or not connected")
            raise HTTPException(status_code=404, detail="Server is not connected")

        if not tool_name:
            print("❌ [TOOLS/CALL] toolName is required but missing")
            raise HTTPException(status_code=400, detail="toolName is required")

        result = await client.request(
            method="tools/call",
            params={"name": tool_name, "arguments": parameters},
        )
        return {"result": result}

    except HTTPException:
        raise
    except Exception as e:
        print(f"💥 [TOOLS/CALL] Error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to call tool on MCP server: {str(e)}")

@app.post("/v1/mcp/servers/{server_id}/tools/list")
async def list_tools(server_id: str):
    """List tools available on an MCP server. Browser-use preset uses HTTP client."""
    if not MCP_AVAILABLE:
        raise HTTPException(status_code=503, detail="MCP SDK not available")

    try:
        print(f"🔍 [TOOLS/LIST] Server: {server_id}")

        server = mcp_servers.get(server_id)
        # Browser-use preset: list tools from HTTP server
        if server and "mcp-browser-use" in server.get("name", "").lower():
            try:
                result = await _browser_use_http_list_tools()
                return {"result": result}
            except Exception as e:
                print(f"❌ [TOOLS/LIST] Browser-use HTTP error: {e}")
                raise HTTPException(
                    status_code=503,
                    detail=BROWSER_USE_HTTP_UNAVAILABLE_MSG + " " + str(e),
                )

        global mcp_clients
        client = mcp_clients.get(server_id)
        if not client:
            print(f"❌ [TOOLS/LIST] Server {server_id} not found or not connected")
            raise HTTPException(status_code=404, detail="Server is not connected")

        result = await client.request(
            method="tools/list",
            params={},
        )

        print(f"📨 [TOOLS/LIST] Raw response from MCP server: {result}")

        # Validate response structure
        if not result:
            print("❌ [TOOLS/LIST] No result returned from MCP server")
        elif 'tools' not in result:
            print(f"❌ [TOOLS/LIST] Missing 'tools' field in response: {list(result.keys())}")
        elif not isinstance(result['tools'], list):
            print(f"❌ [TOOLS/LIST] 'tools' field is not an array: {type(result['tools'])}")
        else:
            print(f"✅ [TOOLS/LIST] Found {len(result['tools'])} tools in response")
            for i, tool in enumerate(result['tools']):
                print(f"  Tool {i}: {tool.get('name', 'unnamed')}")
                if 'name' not in tool:
                    print(f"    ❌ Missing name for tool {i}")
                if 'description' not in tool:
                    print(f"    ⚠️  Missing description for tool {i}")
                if 'inputSchema' not in tool:
                    print(f"    ⚠️  Missing inputSchema for tool {i}")
                else:
                    schema = tool['inputSchema']
                    print(f"    ✅ inputSchema type: {schema.get('type')}")
                    if 'properties' in schema:
                        print(f"    ✅ Has {len(schema['properties'])} properties")
                    else:
                        print("    ⚠️  No properties in inputSchema")

        return {"result": result}

    except HTTPException:
        raise
    except Exception as e:
        print(f"💥 [TOOLS/LIST] Error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to list tools on MCP server: {str(e)}")

# Root endpoint
@app.get("/")
async def root():
    """Root endpoint."""
    return {"message": "CATBot Proxy Server", "version": "2.0.0"}

@app.post("/v1/auth/signup", response_model=AuthTokenResponse)
async def auth_signup(request: AuthSignupRequest):
    """Create a new user account and return a signed JWT."""
    username = request.username.strip().lower()
    password = request.password

    if len(username) < 3:
        raise HTTPException(status_code=400, detail="username must be at least 3 characters")
    if len(password) < 8:
        raise HTTPException(status_code=400, detail="password must be at least 8 characters")
    if username in users_db:
        raise HTTPException(status_code=409, detail="username already exists")

    password_record = create_password_record(password)
    users_db[username] = {
        **password_record,
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    save_users_db()

    token = create_jwt({"sub": username})
    return AuthTokenResponse(
        access_token=token,
        expires_in=JWT_EXPIRATION_SECONDS,
        username=username,
    )


@app.post("/v1/auth/login", response_model=AuthTokenResponse)
async def auth_login(request: AuthLoginRequest):
    """Authenticate a user and return a signed JWT."""
    username = request.username.strip().lower()
    user = users_db.get(username)
    if not user:
        raise HTTPException(status_code=401, detail="Invalid credentials")

    if not verify_password(request.password, user["salt"], user["password_hash"]):
        raise HTTPException(status_code=401, detail="Invalid credentials")

    token = create_jwt({"sub": username})
    return AuthTokenResponse(
        access_token=token,
        expires_in=JWT_EXPIRATION_SECONDS,
        username=username,
    )


@app.get("/v1/auth/me", response_model=AuthUserResponse)
async def auth_me(current_user: Dict[str, Any] = Depends(get_current_user)):
    """Return the profile of the authenticated user based on JWT bearer token."""
    return AuthUserResponse(
        username=current_user["username"],
        created_at=current_user["created_at"],
    )


@app.post("/v1/telegram/chat", response_model=TelegramChatResponse)
async def telegram_chat_endpoint(raw_request: Request, request: TelegramChatRequest):
    """Process a Telegram chat message via OpenAI-compatible API."""
    _validate_telegram_secret(raw_request)

    message_text = (request.message or "").strip()
    if not message_text:
        raise HTTPException(status_code=400, detail="message is required")

    # Prefer OPENAI_API_KEY; fall back to MCP_LLM_OPENAI_API_KEY for single .env setups
    api_key = os.getenv("OPENAI_API_KEY") or os.getenv("MCP_LLM_OPENAI_API_KEY")
    if not api_key:
        raise HTTPException(
            status_code=503,
            detail="OPENAI_API_KEY or MCP_LLM_OPENAI_API_KEY is not configured on the server",
        )

    conversation_id = request.conversation_id or request.user_id or "default"
    history = telegram_conversations.setdefault(conversation_id, [])

    if request.history is not None:
        history = [
            {"role": msg.role, "content": msg.content}
            for msg in request.history
            if msg.content
        ]
        telegram_conversations[conversation_id] = history
        trim_telegram_history(history)

    history.append({"role": "user", "content": message_text})
    trim_telegram_history(history)

    system_prompt = request.system_prompt
    if system_prompt is None:
        system_prompt = _get_telegram_system_prompt_base()

    # Prepend assistant context (timezone + knowledge awareness)
    system_prompt = _get_assistant_context_block() + system_prompt

    # Retrieve relevant memories if memory system is available
    memory_context = ""
    if MEMORY_AVAILABLE and memory_manager:
        try:
            # Search for relevant memories based on the current message
            # Use a lower threshold (0.3) for automatic retrieval to catch more relevant memories
            relevant_memories = await memory_manager.search_memories(
                query=message_text,
                limit=5,
                similarity_threshold=0.3,  # Lower threshold for better recall
            )
            
            # Log search results for debugging
            if relevant_memories:
                print(f"Found {len(relevant_memories)} relevant memories for query: '{message_text[:50]}...'")
                for mem in relevant_memories:
                    print(f"  - {mem.get('text', '')} (similarity: {mem.get('similarity', 0):.3f})")
            else:
                print(f"No memories found for query: '{message_text[:50]}...' (threshold: 0.3)")
            
            # Build memory context if memories found
            if relevant_memories:
                memory_context = "\n\nRelevant context from previous conversations:\n"
                for i, mem in enumerate(relevant_memories, 1):
                    memory_context += f"{i}. {mem.get('text', '')}\n"
                memory_context += "\nUse this context to provide more personalized and relevant responses."
        except Exception as e:
            print(f"Warning: Failed to retrieve memories: {e}")
            import traceback
            print(traceback.format_exc())

    # Add memory context to system prompt
    if memory_context:
        system_prompt = system_prompt + memory_context

    messages: List[Dict[str, str]] = []
    if system_prompt:
        messages.append({"role": "system", "content": system_prompt})
    messages.extend(history)

    model_name = request.model or TELEGRAM_DEFAULT_MODEL
    if not model_name:
        raise HTTPException(status_code=400, detail="No model configured for Telegram chat")

    payload: Dict[str, Any] = {
        "model": model_name,
        "messages": messages,
    }

    if request.temperature is not None:
        payload["temperature"] = request.temperature

    if request.max_output_tokens is not None:
        payload["max_tokens"] = request.max_output_tokens

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }

    if OPENAI_ORG_ID:
        headers["OpenAI-Organization"] = OPENAI_ORG_ID

    if OPENAI_PROJECT_ID:
        headers["OpenAI-Project"] = OPENAI_PROJECT_ID

    url = build_openai_url(TELEGRAM_OPENAI_CHAT_PATH)

    try:
        async with httpx.AsyncClient(timeout=TELEGRAM_CHAT_TIMEOUT) as client:
            response = await client.post(url, headers=headers, json=payload)
    except httpx.RequestError as exc:
        print(f"Telegram chat request error: {exc}")
        raise HTTPException(status_code=502, detail="Failed to contact language model service") from exc

    if response.status_code != 200:
        print(f"Telegram chat API error {response.status_code}: {response.text}")
        detail = response.text
        try:
            error_json = response.json()
            detail = (
                error_json.get("error", {}).get("message")
                or error_json.get("message")
                or detail
            )
        except ValueError:
            pass
        raise HTTPException(status_code=response.status_code, detail=detail)

    data = response.json()
    reply = None
    choices = data.get("choices") or []
    if choices:
        message = choices[0].get("message") or {}
        reply = message.get("content")

    if not reply:
        reply = "I couldn't generate a response right now. Please try again shortly."

    history.append({"role": "assistant", "content": reply})
    trim_telegram_history(history)

    # Extract and store memories if memory system is available and auto-extract is enabled
    if MEMORY_AVAILABLE and memory_manager:
        auto_extract = os.getenv("MEMORY_AUTO_EXTRACT", "true").lower() == "true"
        if auto_extract:
            try:
                # Extract memories from the conversation (last few messages)
                # Include both user message and assistant response
                recent_messages = history[-4:] if len(history) >= 4 else history
                await memory_manager.extract_memories_from_conversation(
                    messages=recent_messages,
                    max_memories=3,
                )
            except Exception as e:
                print(f"Warning: Failed to extract memories: {e}")

    usage = data.get("usage") if isinstance(data, dict) else None

    return TelegramChatResponse(
        reply=reply,
        conversation_id=conversation_id,
        usage=usage,
    )


@app.delete("/v1/telegram/chat/{conversation_id}")
async def telegram_clear_conversation(request: Request, conversation_id: str):
    """Clear cached Telegram conversation history for a user."""
    _validate_telegram_secret(request)

    removed = telegram_conversations.pop(conversation_id, None) is not None
    return {"conversation_id": conversation_id, "cleared": removed}

# ============================================================================
# MEMORY SYSTEM ENDPOINTS
# ============================================================================

@app.post("/v1/memory/store", response_model=MemoryResponse)
async def store_memory(request: MemoryStoreRequest):
    """Store a memory explicitly."""
    if not MEMORY_AVAILABLE or not memory_manager:
        error_detail = "Memory system is not available."
        if not MEMORY_AVAILABLE:
            error_detail += f" Import failed: {memory_import_error}. Check /v1/memory/status for details."
        elif not memory_manager:
            error_detail += " Memory manager initialization failed. Check server logs and /v1/memory/status endpoint."
        raise HTTPException(
            status_code=503,
            detail=error_detail
        )
    
    try:
        # Store the memory
        memory_id = await memory_manager.store_memory(
            text=request.text,
            category=request.category,
            source=request.source or "explicit",
            metadata=request.metadata,
        )
        
        return MemoryResponse(
            success=True,
            message=f"Memory stored successfully",
            data={"memory_id": memory_id}
        )
    except Exception as e:
        print(f"Error storing memory: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to store memory: {str(e)}")

@app.post("/v1/memory/search", response_model=MemoryResponse)
async def search_memories(request: MemorySearchRequest):
    """Search memories by query."""
    if not MEMORY_AVAILABLE or not memory_manager:
        raise HTTPException(
            status_code=503,
            detail="Memory system is not available. Check MEMORY_ENABLED setting."
        )
    
    try:
        # Search for relevant memories
        results = await memory_manager.search_memories(
            query=request.query,
            limit=request.limit,
            similarity_threshold=request.similarity_threshold,
            category=request.category,
        )
        
        return MemoryResponse(
            success=True,
            message=f"Found {len(results)} relevant memories",
            data={"memories": results, "count": len(results)}
        )
    except Exception as e:
        print(f"Error searching memories: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to search memories: {str(e)}")

@app.get("/v1/memory/list", response_model=MemoryResponse)
async def list_memories(limit: Optional[int] = None):
    """List recent memories."""
    if not MEMORY_AVAILABLE or not memory_manager:
        raise HTTPException(
            status_code=503,
            detail="Memory system is not available. Check MEMORY_ENABLED setting."
        )
    
    try:
        # List memories
        memories = memory_manager.list_memories(limit=limit)
        
        return MemoryResponse(
            success=True,
            message=f"Retrieved {len(memories)} memories",
            data={"memories": memories, "count": len(memories), "total": memory_manager.count()}
        )
    except Exception as e:
        print(f"Error listing memories: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to list memories: {str(e)}")

@app.get("/v1/memory/{memory_id}", response_model=MemoryResponse)
async def get_memory(memory_id: str):
    """Get a specific memory by ID."""
    if not MEMORY_AVAILABLE or not memory_manager:
        raise HTTPException(
            status_code=503,
            detail="Memory system is not available. Check MEMORY_ENABLED setting."
        )
    
    try:
        # Get memory by ID
        memory = memory_manager.get_memory(memory_id)
        
        if not memory:
            raise HTTPException(status_code=404, detail=f"Memory {memory_id} not found")
        
        return MemoryResponse(
            success=True,
            message="Memory retrieved successfully",
            data={"memory": memory}
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error getting memory: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to get memory: {str(e)}")

@app.post("/v1/memory/extract", response_model=MemoryResponse)
async def extract_memories(request: MemoryExtractRequest):
    """Extract and store memories from a conversation."""
    if not MEMORY_AVAILABLE or not memory_manager:
        error_detail = "Memory system is not available."
        if not MEMORY_AVAILABLE:
            error_detail += f" Import failed: {memory_import_error}. Check /v1/memory/status for details."
        elif not memory_manager:
            error_detail += " Memory manager initialization failed. Check server logs and /v1/memory/status endpoint."
        raise HTTPException(
            status_code=503,
            detail=error_detail
        )
    
    try:
        # Check if auto-extract is enabled
        auto_extract = os.getenv("MEMORY_AUTO_EXTRACT", "true").lower() == "true"
        if not auto_extract:
            return MemoryResponse(
                success=False,
                message="Automatic memory extraction is disabled via MEMORY_AUTO_EXTRACT=false",
                data={"extracted": 0}
            )
        
        # Extract memories from conversation
        max_memories = request.max_memories or 1
        memory_ids = await memory_manager.extract_memories_from_conversation(
            messages=request.messages,
            max_memories=max_memories,
        )
        
        return MemoryResponse(
            success=True,
            message=f"Extracted and stored {len(memory_ids)} memories",
            data={"extracted": len(memory_ids), "memory_ids": memory_ids}
        )
    except Exception as e:
        print(f"Error extracting memories: {e}")
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Failed to extract memories: {str(e)}")

@app.get("/v1/memory/status")
async def memory_status():
    """Get the status of the memory system."""
    status = {
        "available": MEMORY_AVAILABLE,
        "enabled": os.getenv("MEMORY_ENABLED", "true").lower() == "true",
        "initialized": memory_manager is not None,
        "memory_count": memory_manager.count() if memory_manager else 0,
    }
    if not MEMORY_AVAILABLE:
        status["error"] = f"Memory module not available (import failed: {memory_import_error})"
        if memory_import_error and "numpy" in memory_import_error.lower():
            status["fix"] = "Install numpy with: pip install numpy"
    elif not status["enabled"]:
        status["error"] = "Memory system disabled via MEMORY_ENABLED=false"
    elif not status["initialized"]:
        status["error"] = "Memory manager failed to initialize (check server logs)"
    return status

@app.delete("/v1/memory/{memory_id}", response_model=MemoryResponse)
async def delete_memory(memory_id: str):
    """Delete a memory by ID."""
    if not MEMORY_AVAILABLE or not memory_manager:
        error_detail = "Memory system is not available."
        if not MEMORY_AVAILABLE:
            error_detail += " Memory module import failed."
        elif not memory_manager:
            error_detail += " Memory manager initialization failed. Check server logs and /v1/memory/status endpoint."
        raise HTTPException(
            status_code=503,
            detail=error_detail
        )
    
    try:
        # Delete memory
        deleted = memory_manager.delete_memory(memory_id)
        
        if not deleted:
            raise HTTPException(status_code=404, detail=f"Memory {memory_id} not found")
        
        return MemoryResponse(
            success=True,
            message=f"Memory {memory_id} deleted successfully",
            data={"memory_id": memory_id}
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error deleting memory: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to delete memory: {str(e)}")

# ============================================================================
# END MEMORY SYSTEM ENDPOINTS
# ============================================================================

# ============================================================================
# PHILOSOPHER MODE HELPER FUNCTIONS
# ============================================================================

async def get_all_available_tools() -> List[Dict]:
    """Get all available tools from all connected MCP servers and built-in proxy tools."""
    print(f"[PHILOSOPHER] get_all_available_tools called - MCP_AVAILABLE: {MCP_AVAILABLE}")
    
    all_tools = []
    
    # Add built-in proxy tools (web search, web scraper, news API)
    # These are always available if the server is running
    
    # 1. Web Search Tool
    all_tools.append({
        "name": "web_search",
        "description": "Search the web using Brave Search API (with DuckDuckGo fallback). Returns top search results with URLs, titles, and snippets.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "The search query to execute (e.g., 'latest AI developments 2024')",
                }
            },
            "required": ["query"]
        },
        "server_id": "proxy_server"
    })
    print("[PHILOSOPHER] Added web_search tool")
    
    # 2. Web Scraper/Fetcher Tool
    all_tools.append({
        "name": "web_scraper",
        "description": "Fetch and scrape content from a web URL. Returns the HTML content of the webpage.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "url": {
                    "type": "string",
                    "description": "The URL to fetch and scrape (must include http:// or https://)",
                }
            },
            "required": ["url"]
        },
        "server_id": "proxy_server"
    })
    print("[PHILOSOPHER] Added web_scraper tool")
    
    # 3. News API Tool (only if API key is configured)
    news_api_key = os.getenv('NEWS_API_KEY')
    if news_api_key:
        all_tools.append({
            "name": "news_search",
            "description": "Search for recent news articles using News API. Returns articles with titles, URLs, descriptions, and publication dates.",
            "inputSchema": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "The search query for news articles (e.g., 'artificial intelligence')",
                    }
                },
                "required": ["query"]
            },
            "server_id": "proxy_server"
        })
        print("[PHILOSOPHER] Added news_search tool")
    else:
        print("[PHILOSOPHER] NEWS_API_KEY not configured, skipping news_search tool")
    
    # Add MCP tools if MCP is available
    if MCP_AVAILABLE:
        # Debug: Check what servers are available
        print(f"[PHILOSOPHER] Checking MCP tools - mcp_clients: {len(mcp_clients)} clients, mcp_servers: {len(mcp_servers)} servers")
        print(f"[PHILOSOPHER] Connected client IDs: {list(mcp_clients.keys())}")
        print(f"[PHILOSOPHER] Server IDs in mcp_servers: {list(mcp_servers.keys())}")
        
        # First, check mcp_servers for connected browser-use servers
        # (Browser-use servers are marked as connected but may not be in mcp_clients)
        for server_id, server in mcp_servers.items():
            server_status = server.get("status", "disconnected")
            server_name = server.get("name", "").lower()
            
            # Check if this is a connected browser-use server
            if server_status == "connected" and "mcp-browser-use" in server_name:
                print(f"[PHILOSOPHER] Found connected browser-use server: {server_id}")
                # Add browser automation tool
                all_tools.append({
                    "name": "run_browser_agent",
                    "description": "Control a web browser using natural language commands. Executes browser automation tasks and returns results.",
                    "inputSchema": {
                        "type": "object",
                        "properties": {
                            "instruction": {
                                "type": "string",
                                "description": "Natural language instruction for browser automation (e.g., 'Navigate to Google and search for cats')",
                            },
                            "max_steps": {
                                "type": "integer",
                                "description": "Maximum number of steps the agent should take",
                                "default": 10
                            },
                            "use_vision": {
                                "type": "boolean",
                                "description": "Whether to use vision for understanding page content",
                                "default": True
                            }
                        },
                        "required": ["instruction"]
                    },
                    "server_id": server_id
                })
        
        # Get tools from each connected MCP client (non-browser-use servers)
        for server_id, client in mcp_clients.items():
            try:
                print(f"[PHILOSOPHER] Processing MCP server: {server_id}")
                # Check if this is the browser-use server (shouldn't be, but check anyway)
                server = mcp_servers.get(server_id)
                print(f"[PHILOSOPHER] Server config for {server_id}: {server}")
                
                if server and "mcp-browser-use" in server.get("name", "").lower():
                    # Skip - already handled above
                    print(f"[PHILOSOPHER] Skipping browser-use server {server_id} (already handled)")
                    continue
                else:
                    # Get tools from MCP server
                    print(f"[PHILOSOPHER] Requesting tools/list from MCP server {server_id}")
                    result = await client.request(method="tools/list", params={})
                    print(f"[PHILOSOPHER] tools/list response from {server_id}: {result}")
                    if result and "tools" in result:
                        print(f"[PHILOSOPHER] Found {len(result['tools'])} tools from server {server_id}")
                        for tool in result["tools"]:
                            tool["server_id"] = server_id
                            all_tools.append(tool)
                    else:
                        print(f"[PHILOSOPHER] No tools in response from server {server_id}")
            except Exception as e:
                print(f"[PHILOSOPHER] Error getting tools from server {server_id}: {e}")
                import traceback
                print(traceback.format_exc())
                continue
    else:
        print("[PHILOSOPHER] MCP not available, skipping MCP tools")
    
    print(f"[PHILOSOPHER] Total tools collected: {len(all_tools)}")
    return all_tools

async def execute_tool_for_philosopher(tool_name: str, parameters: Dict) -> str:
    """Execute a tool for philosopher mode. Returns result as string."""
    print(f"[PHILOSOPHER] Executing tool: {tool_name} with parameters: {parameters}")
    
    # Handle built-in proxy server tools
    if tool_name == "web_search":
        try:
            query = parameters.get("query", "")
            if not query:
                return "Error: 'query' parameter is required for web_search"
            
            # Call the proxy search endpoint
            result = await proxy_search(query)
            
            # Format results as readable text
            if result and "results" in result:
                results = result["results"]
                if not results:
                    return f"No search results found for query: {query}"
                
                formatted_results = []
                for i, item in enumerate(results, 1):
                    formatted_results.append(
                        f"{i}. {item.get('title', 'No title')}\n"
                        f"   URL: {item.get('url', 'No URL')}\n"
                        f"   {item.get('snippet', 'No description')}"
                    )
                
                return f"Search results for '{query}':\n\n" + "\n\n".join(formatted_results)
            else:
                return f"Search returned no results for query: {query}"
        except HTTPException as e:
            return f"Error executing web_search: {e.detail}"
        except Exception as e:
            return f"Error executing web_search: {str(e)}"
    
    elif tool_name == "web_scraper":
        try:
            url = parameters.get("url", "")
            if not url:
                return "Error: 'url' parameter is required for web_scraper"
            
            # Call the proxy fetch endpoint
            result = await _do_proxy_fetch(url)
            
            # Return the content (may be HTML)
            if result and "content" in result:
                content = result["content"]
                # Truncate if too long (keep first 10000 characters)
                if len(content) > 10000:
                    return f"Web content from {url} (truncated):\n\n{content[:10000]}...\n\n[Content truncated to 10000 characters]"
                return f"Web content from {url}:\n\n{content}"
            else:
                return f"No content retrieved from URL: {url}"
        except HTTPException as e:
            return f"Error executing web_scraper: {e.detail}"
        except Exception as e:
            return f"Error executing web_scraper: {str(e)}"
    
    elif tool_name == "news_search":
        try:
            query = parameters.get("query", "")
            if not query:
                return "Error: 'query' parameter is required for news_search"
            
            # Call the proxy news search endpoint
            result = await proxy_news_search(query)
            
            # Format results as readable text
            if result and "articles" in result:
                articles = result["articles"]
                if not articles:
                    return f"No news articles found for query: {query}"
                
                formatted_articles = []
                for i, article in enumerate(articles[:10], 1):  # Limit to top 10
                    formatted_articles.append(
                        f"{i}. {article.get('title', 'No title')}\n"
                        f"   Source: {article.get('source', 'Unknown')}\n"
                        f"   Published: {article.get('publishedAt', 'Unknown date')}\n"
                        f"   URL: {article.get('url', 'No URL')}\n"
                        f"   {article.get('description', 'No description')}"
                    )
                
                total = result.get("totalResults", len(articles))
                return f"News articles for '{query}' (showing {len(formatted_articles)} of {total}):\n\n" + "\n\n".join(formatted_articles)
            else:
                return f"News search returned no articles for query: {query}"
        except HTTPException as e:
            return f"Error executing news_search: {e.detail}"
        except Exception as e:
            return f"Error executing news_search: {str(e)}"
    
    # Handle MCP tools
    elif MCP_AVAILABLE:
        # Find which server has this tool
        all_tools = await get_all_available_tools()
        server_id = None
        
        for tool in all_tools:
            if tool.get("name") == tool_name:
                server_id = tool.get("server_id")
                break
        
        if not server_id:
            return f"Error: Tool '{tool_name}' not found on any connected server"
        
        # Execute the tool using the existing call_tool logic
        try:
            request = ToolCallRequest(toolName=tool_name, parameters=parameters)
            result = await call_tool(server_id, request)
            
            # Extract result content
            if result and "result" in result:
                result_data = result["result"]
                if "content" in result_data:
                    content = result_data["content"]
                    if isinstance(content, list):
                        # Extract text from content array
                        text_parts = []
                        for item in content:
                            if isinstance(item, dict) and item.get("type") == "text":
                                text_parts.append(item.get("text", ""))
                        return "\n".join(text_parts)
                    return str(content)
                return str(result_data)
            
            return str(result)
        except Exception as e:
            return f"Error executing tool {tool_name}: {str(e)}"
    else:
        return f"Error: Tool '{tool_name}' is not a built-in tool and MCP is not available"

# ============================================================================
# PHILOSOPHER MODE ENDPOINTS
# ============================================================================

@app.post("/v1/philosopher/start", response_model=PhilosopherResponse)
async def philosopher_start(request: PhilosopherStartRequest):
    """Enable philosopher mode for a conversation."""
    if not PHILOSOPHER_MODE_AVAILABLE or not PhilosopherMode:
        raise HTTPException(
            status_code=503,
            detail="Philosopher mode is not available. Check server logs for details."
        )
    
    # Check if philosopher mode is enabled
    philosopher_enabled = os.getenv("PHILOSOPHER_MODE_ENABLED", "true").lower() == "true"
    if not philosopher_enabled:
        raise HTTPException(
            status_code=503,
            detail="Philosopher mode is disabled via PHILOSOPHER_MODE_ENABLED=false"
        )
    
    # Get conversation ID
    conversation_id = request.conversation_id or request.user_id or "default"
    
    # Check if already active
    if philosopher_mode_active.get(conversation_id, False):
        return PhilosopherResponse(
            success=True,
            message="Philosopher mode is already active for this conversation",
            data={"conversation_id": conversation_id, "active": True}
        )
    
    try:
        # Get API configuration
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise HTTPException(status_code=503, detail="OPENAI_API_KEY is not configured")
        
        api_base = os.getenv("OPENAI_API_BASE", "https://api.openai.com/v1")
        # Use OPENAI_MODEL directly (not TELEGRAM_DEFAULT_MODEL) since philosopher mode is not Telegram-specific
        model = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
        
        # Create philosopher mode instance with tool support
        philosopher = PhilosopherMode(
            api_key=api_key,
            api_base=api_base,
            model=model,
            memory_manager=memory_manager if MEMORY_AVAILABLE else None,
            max_cycles=int(os.getenv("PHILOSOPHER_MAX_CYCLES", "10")),
            similarity_threshold=float(os.getenv("PHILOSOPHER_SIMILARITY_THRESHOLD", "0.3")),
            memory_limit=int(os.getenv("PHILOSOPHER_MEMORY_LIMIT", "10")),
            conversation_history_limit=int(os.getenv("PHILOSOPHER_CONVERSATION_HISTORY_LIMIT", "3")),
            tool_executor=execute_tool_for_philosopher,
            get_tools_func=get_all_available_tools,
            diversification_threshold=int(os.getenv("PHILOSOPHER_DIVERSIFICATION_THRESHOLD", "7")),
        )
        
        # Store instance and activate mode
        philosopher_mode_instances[conversation_id] = philosopher
        philosopher_mode_active[conversation_id] = True
        
        return PhilosopherResponse(
            success=True,
            message="Philosopher mode activated",
            data={"conversation_id": conversation_id, "active": True}
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error starting philosopher mode: {e}")
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Failed to start philosopher mode: {str(e)}")

@app.post("/v1/philosopher/stop", response_model=PhilosopherResponse)
async def philosopher_stop(request: PhilosopherStopRequest):
    """Disable philosopher mode for a conversation."""
    # Get conversation ID
    conversation_id = request.conversation_id or request.user_id or "default"
    
    # Check if active
    if not philosopher_mode_active.get(conversation_id, False):
        return PhilosopherResponse(
            success=True,
            message="Philosopher mode is not active for this conversation",
            data={"conversation_id": conversation_id, "active": False}
        )
    
    try:
        # Deactivate mode
        philosopher_mode_active[conversation_id] = False
        # Remove instance (optional - could keep for reuse)
        philosopher_mode_instances.pop(conversation_id, None)
        
        return PhilosopherResponse(
            success=True,
            message="Philosopher mode deactivated",
            data={"conversation_id": conversation_id, "active": False}
        )
    except Exception as e:
        print(f"Error stopping philosopher mode: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to stop philosopher mode: {str(e)}")

@app.get("/v1/philosopher/status")
async def philosopher_status(conversation_id: Optional[str] = None, user_id: Optional[str] = None):
    """Check if philosopher mode is active for a conversation."""
    # Get conversation ID
    conv_id = conversation_id or user_id or "default"
    
    is_active = philosopher_mode_active.get(conv_id, False)
    
    return {
        "active": is_active,
        "conversation_id": conv_id,
        "available": PHILOSOPHER_MODE_AVAILABLE,
        "enabled": os.getenv("PHILOSOPHER_MODE_ENABLED", "true").lower() == "true"
    }

@app.post("/v1/philosopher/contemplate", response_model=PhilosopherResponse)
async def philosopher_contemplate(request: PhilosopherContemplateRequest):
    """Execute a single contemplation cycle."""
    if not PHILOSOPHER_MODE_AVAILABLE or not PhilosopherMode:
        raise HTTPException(
            status_code=503,
            detail="Philosopher mode is not available. Check server logs for details."
        )
    
    # Get conversation ID
    conversation_id = request.conversation_id or request.user_id or "default"
    
    # Check if mode is active
    if not philosopher_mode_active.get(conversation_id, False):
        raise HTTPException(
            status_code=400,
            detail="Philosopher mode is not active for this conversation. Start it first with /v1/philosopher/start"
        )
    
    # Get philosopher instance
    philosopher = philosopher_mode_instances.get(conversation_id)
    if not philosopher:
        raise HTTPException(
            status_code=500,
            detail="Philosopher mode instance not found. Try restarting philosopher mode."
        )
    
    try:
        # Generate question if not provided
        if request.question:
            question = request.question
        else:
            question = await philosopher.generate_contemplation_question()
            if not question:
                raise HTTPException(status_code=500, detail="Failed to generate contemplation question")
        
        # Execute contemplation
        result = await philosopher.contemplate_question(question)
        
        # Store contemplation in memory
        memory_id = await philosopher.store_contemplation(
            question=result["question"],
            conclusion=result["conclusion"],
            cycle_count=result["cycle_count"]
        )
        
        return PhilosopherResponse(
            success=True,
            message="Contemplation completed",
            data={
                "question": result["question"],
                "conclusion": result["conclusion"],
                "contemplation_steps": result["contemplation_steps"],
                "cycle_count": result["cycle_count"],
                "memory_id": memory_id,
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error during contemplation: {e}")
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Failed to contemplate: {str(e)}")

# ============================================================================
# END PHILOSOPHER MODE ENDPOINTS
# ============================================================================

# Simple test endpoint to verify requests are reaching the server
@app.get("/test")
async def test_endpoint():
    """Simple test endpoint that should always work."""
    import sys
    print("🧪 TEST endpoint called", flush=True)
    sys.stdout.flush()
    return {"message": "test successful", "timestamp": time.time()}

# Health check endpoint
@app.get("/health")
async def health_check():
    """Health check endpoint."""
    # Add explicit logging to debug
    import sys
    print("🏥 Health check endpoint called", flush=True)
    sys.stdout.flush()
    try:
        result = {"status": "healthy", "timestamp": time.time()}
        print(f"🏥 Health check returning: {result}", flush=True)
        sys.stdout.flush()
        return result
    except Exception as e:
        import traceback
        print(f"❌ Health check error: {e}", flush=True)
        print(traceback.format_exc(), flush=True)
        sys.stdout.flush()
        raise

# Models list proxy endpoint to handle CORS and mixed content
@app.get("/v1/proxy/models")
async def proxy_models(request: Request, endpoint: Optional[str] = None):
    """
    Proxy models list requests to handle CORS and avoid mixed content issues.
    Routes requests to the OpenAI-compatible API endpoint specified in the request.
    """
    try:
        # Get the endpoint from query parameter or use default
        if not endpoint:
            endpoint = request.query_params.get('endpoint', '')
        
        # If still no endpoint, use default from environment or localhost
        if not endpoint:
            endpoint = os.getenv('OPENAI_API_BASE', 'http://localhost:1234/v1/models')
        else:
            # Ensure the endpoint includes the full path if not already present
            if not endpoint.endswith('/models'):
                endpoint = endpoint.rstrip('/') + '/models'
        
        # Get Authorization header from the request
        auth_header = request.headers.get('Authorization', '')
        
        # Build headers for the forwarded request
        headers = {}
        if auth_header:
            headers['Authorization'] = auth_header
        
        # Add organization/project headers if present in original request
        org_header = request.headers.get('OpenAI-Organization')
        if org_header:
            headers['OpenAI-Organization'] = org_header
        
        project_header = request.headers.get('OpenAI-Project')
        if project_header:
            headers['OpenAI-Project'] = project_header
        
        print(f"📋 Proxying models list request to: {endpoint}")
        
        # Forward the request to the LLM service
        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.get(
                endpoint,
                headers=headers
            )
        
        print(f"✅ Models list response status: {response.status_code}")
        
        # Check if the response is successful
        if response.status_code != 200:
            print(f"❌ LLM service returned error: {response.status_code}")
            print(f"   Response text: {response.text[:500]}")
            return JSONResponse(
                content=response.json() if response.headers.get('content-type', '').startswith('application/json') else {"error": response.text},
                status_code=response.status_code
            )
        
        # Return the JSON response
        try:
            response_data = response.json()
            return JSONResponse(content=response_data, status_code=200)
        except Exception as json_error:
            print(f"❌ Failed to parse JSON response: {json_error}")
            return JSONResponse(
                content={"error": "Invalid JSON response from LLM service"},
                status_code=500
            )
    
    except httpx.ConnectError as e:
        print(f"❌ Connection error: Could not connect to LLM service")
        raise HTTPException(
            status_code=503,
            detail=f"Could not connect to LLM service. Please check the endpoint configuration."
        )
    except Exception as e:
        print(f"❌ Models list proxy error: {e}")
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Failed to proxy models list request: {str(e)}")

# Browser automation proxy endpoints to handle CORS and mixed content
@app.post("/v1/proxy/browser-agent")
async def proxy_browser_agent(request: Request):
    """
    Proxy browser automation requests to the MCP browser server.
    Routes requests to avoid mixed content issues with HTTPS.
    """
    try:
        # Get the request body
        body = await request.json()
        
        # Get MCP browser server URL from environment or use default
        # Try multiple connection methods for better reliability
        mcp_browser_url = os.getenv('MCP_BROWSER_SERVER_URL', None)
        
        # If not set, try to determine the best URL to use
        if not mcp_browser_url:
            # Try to get local IP address
            try:
                s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
                s.connect(("8.8.8.8", 80))
                local_ip = s.getsockname()[0]
                s.close()
                # Try using the local IP first, then fall back to 127.0.0.1
                mcp_browser_url = f"http://{local_ip}:5001"
                print(f"   Detected local IP: {local_ip}, using {mcp_browser_url}")
            except Exception:
                # Fall back to 127.0.0.1
                mcp_browser_url = "http://127.0.0.1:5001"
                print(f"   Using default: {mcp_browser_url}")
        else:
            print(f"   Using configured MCP_BROWSER_SERVER_URL: {mcp_browser_url}")
        
        endpoint = f"{mcp_browser_url.rstrip('/')}/api/browser-agent"
        
        print(f"🌐 Proxying browser-agent request to: {endpoint}")
        print(f"   Request body keys: {list(body.keys()) if isinstance(body, dict) else 'N/A'}")
        
        # First, try a quick health check to see if the server is responsive
        health_endpoint = f"{mcp_browser_url.rstrip('/')}/api/health"
        health_check_passed = False
        try:
            async with httpx.AsyncClient(timeout=5.0) as health_client:
                health_response = await health_client.get(health_endpoint)
                if health_response.status_code == 200:
                    print(f"   ✅ MCP browser server health check passed")
                    health_check_passed = True
                else:
                    print(f"   ⚠️  MCP browser server health check returned: {health_response.status_code}")
        except Exception as health_err:
            print(f"   ⚠️  MCP browser server health check failed: {health_err}")
            # If health check fails with IP, try 127.0.0.1 as fallback
            if not mcp_browser_url.startswith("http://127.0.0.1"):
                print(f"   Trying fallback to 127.0.0.1...")
                mcp_browser_url = "http://127.0.0.1:5001"
                endpoint = f"{mcp_browser_url.rstrip('/')}/api/browser-agent"
                try:
                    async with httpx.AsyncClient(timeout=5.0) as health_client:
                        health_response = await health_client.get(f"{mcp_browser_url.rstrip('/')}/api/health")
                        if health_response.status_code == 200:
                            print(f"   ✅ MCP browser server health check passed with 127.0.0.1")
                            health_check_passed = True
                except Exception:
                    print(f"   ⚠️  Health check also failed with 127.0.0.1")
        
        if not health_check_passed:
            print(f"   ⚠️  Warning: Health check failed, but continuing with request")
        
        # Create a timeout configuration - longer for browser automation tasks
        # Use a timeout that allows for long-running browser tasks
        # connect: time to establish connection (10s)
        # read: time to read response (3 hours for browser automation)
        # write: time to write request (10s)
        # pool: time to get connection from pool (10s)
        timeout = httpx.Timeout(connect=10.0, read=10800.0, write=10.0, pool=10.0)
        
        # Forward the request to the MCP browser server
        async with httpx.AsyncClient(timeout=timeout, follow_redirects=True) as client:
            try:
                print(f"   Sending POST request to MCP browser server...")
                response = await client.post(
                    endpoint,
                    json=body,
                    headers={'Content-Type': 'application/json'}
                )
                print(f"   Received response from MCP browser server")
            except httpx.ConnectError as conn_err:
                print(f"❌ Connection error to MCP browser server: {conn_err}")
                print(f"   Attempted endpoint: {endpoint}")
                print(f"   Please ensure the MCP browser server is running: python start_mcp_browser_server.py")
                raise HTTPException(
                    status_code=503,
                    detail="Could not connect to MCP browser server. Please ensure it's running on port 5001."
                )
            except httpx.ReadTimeout as timeout_err:
                print(f"❌ Read timeout from MCP browser server: {timeout_err}")
                print(f"   The browser automation task may be taking longer than expected.")
                print(f"   This could indicate:")
                print(f"   1. The MCP browser server is processing a long-running task")
                print(f"   2. The MCP browser server is not responding")
                print(f"   3. Network connectivity issues")
                raise HTTPException(
                    status_code=504,
                    detail="Browser automation task timed out after 10 minutes. The task may be taking longer than expected. Please try again or check the MCP browser server logs."
                )
            except httpx.TimeoutException as timeout_err:
                print(f"❌ General timeout from MCP browser server: {timeout_err}")
                raise HTTPException(
                    status_code=504,
                    detail="Browser automation task timed out. Please try again or check the MCP browser server logs."
                )
        
        print(f"✅ Browser-agent response status: {response.status_code}")
        
        # Return the response
        if response.status_code != 200:
            error_content = response.json() if response.headers.get('content-type', '').startswith('application/json') else {"error": response.text}
            print(f"   Error response: {error_content}")
            return JSONResponse(
                content=error_content,
                status_code=response.status_code
            )
        
        return JSONResponse(content=response.json(), status_code=200)
    
    except HTTPException:
        # Re-raise HTTP exceptions as-is
        raise
    except Exception as e:
        print(f"❌ Browser-agent proxy error: {e}")
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Failed to proxy browser-agent request: {str(e)}")

@app.post("/v1/proxy/deep-research")
async def proxy_deep_research(request: Request):
    """
    Proxy deep research requests to the MCP browser server.
    Routes requests to avoid mixed content issues with HTTPS.
    """
    try:
        # Get the request body
        body = await request.json()
        
        # Get MCP browser server URL from environment or use default
        # Try multiple connection methods for better reliability
        mcp_browser_url = os.getenv('MCP_BROWSER_SERVER_URL', None)
        
        # If not set, try to determine the best URL to use
        if not mcp_browser_url:
            # Try to get local IP address
            try:
                s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
                s.connect(("8.8.8.8", 80))
                local_ip = s.getsockname()[0]
                s.close()
                # Try using the local IP first, then fall back to 127.0.0.1
                mcp_browser_url = f"http://{local_ip}:5001"
                print(f"   Detected local IP: {local_ip}, using {mcp_browser_url}")
            except Exception:
                # Fall back to 127.0.0.1
                mcp_browser_url = "http://127.0.0.1:5001"
                print(f"   Using default: {mcp_browser_url}")
        else:
            print(f"   Using configured MCP_BROWSER_SERVER_URL: {mcp_browser_url}")
        
        endpoint = f"{mcp_browser_url.rstrip('/')}/api/deep-research"
        
        print(f"🔬 Proxying deep-research request to: {endpoint}")
        print(f"   Request body keys: {list(body.keys()) if isinstance(body, dict) else 'N/A'}")
        
        # Create a timeout configuration - even longer for deep research tasks
        timeout = httpx.Timeout(connect=10.0, read=10800.0, write=10.0, pool=10.0)  # 3 hours for deep research
        
        # Forward the request to the MCP browser server
        async with httpx.AsyncClient(timeout=timeout) as client:
            try:
                response = await client.post(
                    endpoint,
                    json=body,
                    headers={'Content-Type': 'application/json'}
                )
            except httpx.ConnectError as conn_err:
                print(f"❌ Connection error to MCP browser server: {conn_err}")
                print(f"   Attempted endpoint: {endpoint}")
                print(f"   Please ensure the MCP browser server is running: python start_mcp_browser_server.py")
                raise HTTPException(
                    status_code=503,
                    detail="Could not connect to MCP browser server. Please ensure it's running on port 5001."
                )
            except httpx.ReadTimeout as timeout_err:
                print(f"❌ Read timeout from MCP browser server: {timeout_err}")
                print(f"   The deep research task may be taking longer than expected.")
                raise HTTPException(
                    status_code=504,
                    detail="Deep research task timed out. The task may be taking longer than expected. Please try again or check the MCP browser server logs."
                )
        
        print(f"✅ Deep-research response status: {response.status_code}")
        
        # Return the response
        if response.status_code != 200:
            error_content = response.json() if response.headers.get('content-type', '').startswith('application/json') else {"error": response.text}
            print(f"   Error response: {error_content}")
            return JSONResponse(
                content=error_content,
                status_code=response.status_code
            )
        
        return JSONResponse(content=response.json(), status_code=200)
    
    except HTTPException:
        # Re-raise HTTP exceptions as-is
        raise
    except Exception as e:
        print(f"❌ Deep-research proxy error: {e}")
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Failed to proxy deep-research request: {str(e)}")

# Chat completions proxy endpoint to handle CORS and mixed content
@app.post("/v1/proxy/chat/completions")
async def proxy_chat_completions(request: Request):
    """
    Proxy chat completions requests to handle CORS and avoid mixed content issues.
    Routes requests to the OpenAI-compatible API endpoint specified in the request.
    """
    try:
        # Get the request body
        body = await request.json()
        
        # Get the endpoint from query parameter or request body, or use default
        endpoint = request.query_params.get('endpoint', '')
        if not endpoint:
            # Try to get from body (some clients might send it)
            endpoint = body.get('_endpoint', '')
        
        # If still no endpoint, use default from environment or localhost
        if not endpoint:
            endpoint = os.getenv('OPENAI_API_BASE', 'http://localhost:1234/v1/chat/completions')
        else:
            # Ensure the endpoint includes the full path if not already present
            if not endpoint.endswith('/chat/completions'):
                endpoint = endpoint.rstrip('/') + '/chat/completions'
        
        # Remove internal endpoint parameter from body before forwarding
        body_clean = {k: v for k, v in body.items() if k != '_endpoint'}
        
        # Get Authorization header from the request
        auth_header = request.headers.get('Authorization', '')
        
        # Build headers for the forwarded request
        headers = {
            'Content-Type': 'application/json'
        }
        if auth_header:
            headers['Authorization'] = auth_header
        
        # Add organization/project headers if present in original request
        org_header = request.headers.get('OpenAI-Organization')
        if org_header:
            headers['OpenAI-Organization'] = org_header
        
        project_header = request.headers.get('OpenAI-Project')
        if project_header:
            headers['OpenAI-Project'] = project_header
        
        print(f"💬 Proxying chat completions request to: {endpoint}")
        print(f"   Model: {body_clean.get('model', 'unknown')}")
        
        # Forward the request to the LLM service
        async with httpx.AsyncClient(timeout=120.0) as client:
            response = await client.post(
                endpoint,
                json=body_clean,
                headers=headers
            )
        
        print(f"✅ Chat completions response status: {response.status_code}")
        
        # Check if the response is successful
        if response.status_code != 200:
            print(f"❌ LLM service returned error: {response.status_code}")
            print(f"   Response text: {response.text[:500]}")
            return JSONResponse(
                content=response.json() if response.headers.get('content-type', '').startswith('application/json') else {"error": response.text},
                status_code=response.status_code
            )
        
        # Return the JSON response
        try:
            response_data = response.json()
            return JSONResponse(content=response_data, status_code=200)
        except Exception as json_error:
            print(f"❌ Failed to parse JSON response: {json_error}")
            return JSONResponse(
                content={"error": "Invalid JSON response from LLM service"},
                status_code=500
            )
    
    except httpx.ConnectError as e:
        print(f"❌ Connection error: Could not connect to LLM service")
        raise HTTPException(
            status_code=503,
            detail=f"Could not connect to LLM service. Please check the endpoint configuration."
        )
    except Exception as e:
        print(f"❌ Chat completions proxy error: {e}")
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Failed to proxy chat completions request: {str(e)}")

# OPTIONS handler for Whisper endpoint to handle CORS preflight
@app.options("/v1/audio/transcriptions")
async def proxy_whisper_options(request: Request):
    """Handle CORS preflight requests for Whisper endpoint."""
    return JSONResponse(
        content={},
        status_code=200,
        headers={
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Methods": "POST, OPTIONS",
            "Access-Control-Allow-Headers": "*",
            "Access-Control-Max-Age": "3600",
        }
    )

# Whisper proxy endpoint to handle CORS
@app.post("/v1/audio/transcriptions")
async def proxy_whisper(request: Request):
    """Proxy Whisper transcription requests to handle CORS."""
    try:
        # Get the form data from the request
        form_data = await request.form()
        
        # Get the Whisper endpoint (defaulting to localhost:8001)
        whisper_endpoint = os.getenv('WHISPER_ENDPOINT', 'http://localhost:8001/v1/audio/transcriptions')
        
        print(f"📝 Proxying Whisper request to: {whisper_endpoint}")
        
        # Get Authorization header from the request
        auth_header = request.headers.get('Authorization', '')
        
        # Prepare the files and data for forwarding
        files = {}
        data = {}
        
        for key, value in form_data.items():
            if hasattr(value, 'read'):  # This is a file
                # Read the file content
                file_content = await value.read()
                files[key] = (value.filename, file_content, value.content_type)
                print(f"  📎 File: {value.filename} ({len(file_content)} bytes)")
            else:  # This is a regular form field
                data[key] = value
                print(f"  📄 Field: {key} = {value}")
        
        # Forward the request to the Whisper service
        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.post(
                whisper_endpoint,
                files=files,
                data=data,
                headers={'Authorization': auth_header} if auth_header else {}
            )
        
        print(f"✅ Whisper response status: {response.status_code}")
        print(f"📄 Response content type: {response.headers.get('content-type', 'unknown')}")
        
        # Check if the response is successful
        if response.status_code != 200:
            print(f"❌ Whisper service returned error: {response.status_code}")
            print(f"   Response text: {response.text}")
            return JSONResponse(
                content={"error": f"Whisper service error: {response.text}"},
                status_code=response.status_code
            )
        
        # Try to parse the JSON response
        try:
            response_data = response.json()
            print(f"✅ Parsed JSON response: {response_data}")
            return JSONResponse(content=response_data, status_code=200)
        except Exception as json_error:
            print(f"❌ Failed to parse JSON response: {json_error}")
            print(f"   Raw response text: {response.text[:200]}")
            # Return the raw text if JSON parsing fails
            return JSONResponse(
                content={"text": response.text},
                status_code=200
            )
    
    except httpx.ConnectError as e:
        print(f"❌ Connection error: Could not connect to Whisper service at {whisper_endpoint}")
        print(f"   Make sure the Whisper service is running on port 8001")
        raise HTTPException(
            status_code=503,
            detail=f"Could not connect to Whisper service. Make sure it's running on port 8001."
        )
    except Exception as e:
        print(f"❌ Whisper proxy error: {e}")
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Failed to proxy Whisper request: {str(e)}")

# TTS voices proxy endpoint to handle CORS
@app.get("/v1/proxy/tts/voices")
async def proxy_tts_voices(endpoint: str):
    """Proxy TTS voices requests to handle CORS. Tries /voices first, then /v1/audio/voices."""
    if not endpoint:
        raise HTTPException(status_code=400, detail="Endpoint parameter is required")
    
    try:
        # Normalize the endpoint URL (remove trailing slash)
        base_endpoint = endpoint.rstrip('/')
        
        # Extract base URL (origin: protocol + host + port) to avoid path duplication
        try:
            # Parse the endpoint URL to extract the origin (protocol + host + port)
            if not base_endpoint.startswith('http://') and not base_endpoint.startswith('https://'):
                # If no protocol, assume http://
                base_endpoint = f"http://{base_endpoint}"
            
            parsed_url = urlparse(base_endpoint)
            # Reconstruct base URL with scheme, hostname, and port (if present)
            base_url = f"{parsed_url.scheme}://{parsed_url.netloc}"
        except Exception as e:
            # Fallback to simple string replacement if URL parsing fails
            # Extract origin manually using regex
            match = re.match(r'^(https?://[^/]+)', base_endpoint)
            if match:
                base_url = match.group(1)
            else:
                # Last resort: remove /v1 and any path
                base_url = base_endpoint.split('/')[0] if '/' in base_endpoint else base_endpoint
                if not base_url.startswith('http'):
                    base_url = f"http://{base_url}"
        
        # Try /voices first (Chatterbox style)
        voices_url_primary = f"{base_url}/voices"
        print(f"🎤 Trying primary TTS voices endpoint: {voices_url_primary}")
        
        response = None
        response_data = None
        
        async with httpx.AsyncClient(timeout=10.0) as client:
            try:
                # Try the primary endpoint first
                response = await client.get(
                    voices_url_primary,
                    headers={
                        'Content-Type': 'application/json',
                        'Accept': 'application/json'
                    }
                )
                
                print(f"✅ Primary TTS voices response status: {response.status_code}")
                
                # If primary endpoint succeeds, use it
                if response.status_code == 200:
                    try:
                        response_data = response.json()
                        print(f"✅ Parsed TTS voices JSON response from primary endpoint")
                        return JSONResponse(content=response_data, status_code=200)
                    except Exception as json_error:
                        print(f"❌ Failed to parse JSON response: {json_error}")
                        print(f"   Raw response text: {response.text[:200]}")
                        # Return the raw text if JSON parsing fails
                        return JSONResponse(
                            content={"text": response.text},
                            status_code=200
                        )
            except (httpx.ConnectError, httpx.HTTPStatusError) as e:
                # Primary endpoint failed, try fallback
                print(f"⚠️ Primary endpoint failed: {e}")
                response = None
            
            # If primary failed, try /v1/audio/voices (OpenAI-compatible style)
            if not response or response.status_code != 200:
                voices_url_fallback = f"{base_url}/v1/audio/voices"
                print(f"🎤 Trying fallback TTS voices endpoint: {voices_url_fallback}")
                
                try:
                    response = await client.get(
                        voices_url_fallback,
                        headers={
                            'Content-Type': 'application/json',
                            'Accept': 'application/json'
                        }
                    )
                    
                    print(f"✅ Fallback TTS voices response status: {response.status_code}")
                    
                    # Check if the fallback response is successful
                    if response.status_code == 200:
                        try:
                            response_data = response.json()
                            print(f"✅ Parsed TTS voices JSON response from fallback endpoint")
                            return JSONResponse(content=response_data, status_code=200)
                        except Exception as json_error:
                            print(f"❌ Failed to parse JSON response: {json_error}")
                            print(f"   Raw response text: {response.text[:200]}")
                            # Return the raw text if JSON parsing fails
                            return JSONResponse(
                                content={"text": response.text},
                                status_code=200
                            )
                    else:
                        # Fallback also failed
                        print(f"❌ Fallback TTS service returned error: {response.status_code}")
                        print(f"   Response text: {response.text[:200]}")
                        raise HTTPException(
                            status_code=response.status_code,
                            detail=f"TTS service error: {response.text[:200]}"
                        )
                except httpx.ConnectError as e:
                    print(f"❌ Connection error: Could not connect to TTS service at {voices_url_fallback}")
                    raise HTTPException(
                        status_code=503,
                        detail=f"Could not connect to TTS service. Tried {voices_url_primary} and {voices_url_fallback}"
                    )
                except httpx.HTTPStatusError as e:
                    print(f"❌ HTTP error from fallback TTS service: {e.response.status_code}")
                    raise HTTPException(
                        status_code=e.response.status_code,
                        detail=f"TTS service returned error: {str(e)}"
                    )
        
        # If we get here, both attempts failed
        if response and response.status_code != 200:
            raise HTTPException(
                status_code=response.status_code,
                detail=f"TTS service error: {response.text[:200]}"
            )
        else:
            raise HTTPException(
                status_code=503,
                detail=f"Could not connect to TTS service. Tried {voices_url_primary} and {base_url}/v1/audio/voices"
            )
    
    except HTTPException:
        # Re-raise HTTP exceptions as-is
        raise
    except httpx.ConnectError as e:
        print(f"❌ Connection error: Could not connect to TTS service at {endpoint}")
        raise HTTPException(
            status_code=503,
            detail=f"Could not connect to TTS service at {endpoint}"
        )
    except httpx.HTTPStatusError as e:
        print(f"❌ HTTP error from TTS service: {e.response.status_code}")
        raise HTTPException(
            status_code=e.response.status_code,
            detail=f"TTS service returned error: {str(e)}"
        )
    except Exception as e:
        print(f"❌ TTS voices proxy error: {e}")
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Failed to proxy TTS voices request: {str(e)}")

# TTS speech proxy endpoint to handle CORS and streaming
@app.post("/v1/proxy/tts/speech")
async def proxy_tts_speech(request: Request, endpoint: Optional[str] = None):
    """Proxy TTS speech requests to handle CORS. Supports streaming responses."""
    try:
        # Get endpoint from query parameter or try to extract from request
        if not endpoint:
            # Try to get from query params
            endpoint = request.query_params.get('endpoint', '')
        
        # If still no endpoint, try to get from TTS_ENDPOINT env var or use default
        if not endpoint:
            tts_endpoint = os.getenv('TTS_ENDPOINT', 'http://localhost:4123/v1')
            endpoint = tts_endpoint.rstrip('/')
        
        # Normalize the endpoint URL (remove trailing slash)
        base_endpoint = endpoint.rstrip('/')
        
        # Extract base URL (origin: protocol + host + port) to avoid path duplication
        try:
            # Parse the endpoint URL to extract the origin (protocol + host + port)
            if not base_endpoint.startswith('http://') and not base_endpoint.startswith('https://'):
                # If no protocol, assume http://
                base_endpoint = f"http://{base_endpoint}"
            
            parsed_url = urlparse(base_endpoint)
            # Reconstruct base URL with scheme, hostname, and port (if present)
            base_url = f"{parsed_url.scheme}://{parsed_url.netloc}"
        except Exception as e:
            # Fallback to simple string replacement if URL parsing fails
            # Extract origin manually using regex
            match = re.match(r'^(https?://[^/]+)', base_endpoint)
            if match:
                base_url = match.group(1)
            else:
                # Last resort: remove /v1 and any path
                base_url = base_endpoint.split('/')[0] if '/' in base_endpoint else base_endpoint
                if not base_url.startswith('http'):
                    base_url = f"http://{base_url}"
        
        # Construct the speech endpoint URL
        speech_url = f"{base_url}/v1/audio/speech"
        
        print(f"🎤 Proxying TTS speech request to: {speech_url}")
        
        # Get the request body
        try:
            request_body = await request.json()
        except Exception:
            request_body = {}
        
        # Get headers from the original request
        # Forward Accept header to support both SSE (Chatterbox) and binary audio (VibeVoice/OpenAI-compatible)
        # If client requests SSE, forward it; otherwise let TTS service decide (defaults to binary audio)
        accept_header = request.headers.get('Accept', '')
        forward_headers = {
            'Content-Type': 'application/json',
        }
        
        # Forward Accept header if present (allows Chatterbox to return SSE, VibeVoice will ignore and return binary)
        if accept_header:
            forward_headers['Accept'] = accept_header
        
        # Forward Authorization header if present
        auth_header = request.headers.get('Authorization')
        if auth_header:
            forward_headers['Authorization'] = auth_header
        
        # Create a streaming response generator
        # We need to peek at the response to get the actual content-type
        # Since we can't change media_type after StreamingResponse is created,
        # we'll use a wrapper that can handle both SSE and binary audio
        
        class TTSStreamWrapper:
            """Wrapper to handle both SSE and binary audio responses."""
            def __init__(self):
                self.content_type = None
                self.first_chunk = True
                
            async def __aiter__(self):
                try:
                    async with httpx.AsyncClient(timeout=60.0) as client:
                        async with client.stream(
                            'POST',
                            speech_url,
                            json=request_body,
                            headers=forward_headers
                        ) as response:
                            print(f"✅ TTS speech response status: {response.status_code}")
                            
                            # Log request details for debugging
                            print(f"📤 TTS request body: {json.dumps(request_body, indent=2)[:500]}")
                            print(f"📤 TTS request headers: {forward_headers}")
                            
                            # Capture the actual content type from the TTS service
                            self.content_type = response.headers.get('content-type', 'audio/mpeg')
                            print(f"📦 TTS response content-type: {self.content_type}")
                            
                            # Check if the response is successful
                            if response.status_code != 200:
                                error_text = await response.aread()
                                print(f"❌ TTS service returned error: {response.status_code}")
                                print(f"   Response text: {error_text[:200]}")
                                # Yield error as bytes
                                if isinstance(error_text, bytes):
                                    yield error_text
                                else:
                                    yield str(error_text).encode('utf-8')
                                return
                            
                            # Stream the response chunk by chunk as bytes
                            # This preserves binary audio data (MP3, WAV, etc.) or SSE text
                            async for chunk in response.aiter_bytes():
                                if chunk:
                                    yield chunk
                                    
                except httpx.ConnectError as e:
                    print(f"❌ Connection error: Could not connect to TTS service at {speech_url}")
                    error_msg = f"Error: Could not connect to TTS service at {speech_url}"
                    yield error_msg.encode('utf-8')
                except Exception as e:
                    print(f"❌ TTS speech proxy error: {e}")
                    import traceback
                    print(traceback.format_exc())
                    error_msg = f"Error: Failed to proxy TTS speech request: {str(e)}"
                    yield error_msg.encode('utf-8')
        
        stream_wrapper = TTSStreamWrapper()
        
        # Determine initial content-type based on Accept header
        # If client requests SSE, use that; otherwise default to audio/mpeg
        # The actual content-type will be in the response headers
        initial_content_type = 'audio/mpeg'  # Default
        if accept_header and 'text/event-stream' in accept_header:
            initial_content_type = 'text/event-stream'  # Use SSE if requested
        
        # Use StreamingResponse with the initial content-type
        # The actual content-type from the TTS service will be in the response headers
        # JavaScript will check res.headers.get('content-type') to get the actual type
        # Note: FastAPI's StreamingResponse sets Content-Type based on media_type parameter,
        # but the actual response from the TTS service should have its own content-type header
        # that JavaScript can read. However, if Chatterbox is returning audio/mpeg instead of
        # text/event-stream, that means Chatterbox itself is not respecting the stream_format: 'sse'
        # parameter or the Accept header.
        
        response_obj = StreamingResponse(
            stream_wrapper,
            media_type=initial_content_type,  # Initial guess based on Accept header
            status_code=200
        )
        
        return response_obj
    
    except Exception as e:
        print(f"❌ TTS speech proxy error: {e}")
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Failed to proxy TTS speech request: {str(e)}")

# ============================================================================
# FILE OPERATIONS ENDPOINTS
# ============================================================================

def resolve_scratch_path(filename: str, allowed_extensions: Optional[Set[str]] = None) -> Path:
    """
    Resolve a user-supplied filename to a path under SCRATCH_DIR.
    Rejects absolute paths, traversal (..), and disallowed extensions.
    Returns the canonical path for safe I/O. Raises HTTPException 400 on invalid input.
    """
    # Reject empty or whitespace-only filename
    if not filename or not filename.strip():
        raise HTTPException(status_code=400, detail="Invalid filename")
    # Reject absolute paths (Unix / or Windows drive/root)
    if os.path.isabs(filename) or filename.startswith("/") or filename.startswith("\\"):
        raise HTTPException(status_code=400, detail="Invalid filename")
    # Reject path traversal components
    parts = Path(filename).parts
    if ".." in parts:
        raise HTTPException(status_code=400, detail="Invalid filename")
    # Build candidate path and resolve to canonical form
    candidate = SCRATCH_DIR / filename
    try:
        resolved = candidate.resolve()
    except (OSError, RuntimeError) as e:
        raise HTTPException(status_code=400, detail="Invalid filename") from e
    # Enforce containment under SCRATCH_DIR (Python 3.9+ relative_to)
    root = SCRATCH_DIR.resolve()
    try:
        resolved.relative_to(root)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid filename")
    # Optional: enforce extension allowlist
    if allowed_extensions is not None:
        suffix = resolved.suffix.lower()
        if suffix not in allowed_extensions:
            raise HTTPException(status_code=400, detail="Invalid filename")
    return resolved


def read_text_file(filepath: Path) -> str:
    """Read a plain text file and return its content"""
    try:
        # Read the file with UTF-8 encoding
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        # Fallback to latin-1 if UTF-8 fails
        with open(filepath, 'r', encoding='latin-1') as f:
            return f.read()

def read_docx_file(filepath: Path) -> str:
    """Read a Word document and return its text content"""
    # Load the document using python-docx
    doc = Document(filepath)
    # Extract text from all paragraphs
    paragraphs = [para.text for para in doc.paragraphs]
    # Join paragraphs with newlines
    return '\n'.join(paragraphs)

def read_xlsx_file(filepath: Path) -> str:
    """Read an Excel file and return its content as formatted text"""
    # Load the workbook
    wb = openpyxl.load_workbook(filepath, data_only=True)
    result = []
    
    # Process each sheet in the workbook
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        result.append(f"=== Sheet: {sheet_name} ===\n")
        
        # Process each row in the sheet
        for row in sheet.iter_rows(values_only=True):
            # Filter out None values and convert to strings
            row_data = [str(cell) if cell is not None else '' for cell in row]
            # Join cells with tabs for better formatting
            result.append('\t'.join(row_data))
        
        result.append('\n')  # Add blank line between sheets
    
    # Join all lines with newlines
    return '\n'.join(result)

def read_pdf_file(filepath: Path) -> str:
    """Read a PDF file and return its text content"""
    result = []
    # Open the PDF file in binary mode
    with open(filepath, 'rb') as f:
        # Create a PDF reader object
        pdf_reader = PyPDF2.PdfReader(f)
        # Extract text from each page
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text = page.extract_text()
            result.append(f"=== Page {page_num + 1} ===\n{text}\n")
    
    # Join all pages with newlines
    return '\n'.join(result)

def read_png_file(filepath: Path) -> Dict[str, Any]:
    """Read a PNG image and return metadata and base64-encoded data"""
    # Open the image using PIL
    img = Image.open(filepath)
    
    # Get image metadata
    metadata = {
        'width': img.width,
        'height': img.height,
        'format': img.format,
        'mode': img.mode
    }
    
    # Convert image to base64 for transmission
    buffered = BytesIO()
    img.save(buffered, format=img.format)
    img_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')
    
    return {
        'metadata': metadata,
        'data': img_base64,
        'description': f"Image: {img.width}x{img.height} pixels, format: {img.format}"
    }

def write_text_file(filepath: Path, content: str) -> None:
    """Write content to a plain text file"""
    # Write the content with UTF-8 encoding
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(content)

def write_docx_file(filepath: Path, content: str) -> None:
    """Write content to a Word document"""
    # Create a new document
    doc = Document()
    
    # Split content into paragraphs and add each to the document
    for paragraph in content.split('\n'):
        doc.add_paragraph(paragraph)
    
    # Save the document
    doc.save(filepath)

def write_xlsx_file(filepath: Path, content: str) -> None:
    """Write content to an Excel file"""
    # Create a new workbook
    wb = openpyxl.Workbook()
    # Get the active sheet
    ws = wb.active
    ws.title = "Sheet1"
    
    # Split content into rows
    rows = content.split('\n')
    
    # Write each row to the Excel file
    for row_idx, row_content in enumerate(rows, start=1):
        # Split row by tabs or commas
        if '\t' in row_content:
            cells = row_content.split('\t')
        else:
            cells = row_content.split(',')
        
        # Write each cell
        for col_idx, cell_content in enumerate(cells, start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=cell_content.strip())
            
            # Apply formatting to the first row (header)
            if row_idx == 1:
                cell.font = Font(bold=True)
                cell.alignment = Alignment(horizontal='center')
    
    # Auto-adjust column widths
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)  # Cap at 50 characters
        ws.column_dimensions[column_letter].width = adjusted_width
    
    # Save the workbook
    wb.save(filepath)

def write_pdf_file(filepath: Path, content: str) -> None:
    """Write content to a PDF file using reportlab"""
    try:
        # Import reportlab for PDF creation
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas as pdf_canvas
        from reportlab.lib.units import inch
        
        # Create a canvas for PDF generation
        c = pdf_canvas.Canvas(str(filepath), pagesize=letter)
        width, height = letter
        
        # Set up text formatting
        text_object = c.beginText(1 * inch, height - 1 * inch)
        text_object.setFont("Helvetica", 12)
        
        # Split content into lines and add to PDF
        lines = content.split('\n')
        for line in lines:
            # Handle long lines by wrapping
            if len(line) > 80:
                words = line.split(' ')
                current_line = ''
                for word in words:
                    if len(current_line + word) < 80:
                        current_line += word + ' '
                    else:
                        text_object.textLine(current_line)
                        current_line = word + ' '
                if current_line:
                    text_object.textLine(current_line)
            else:
                text_object.textLine(line)
            
            # Add new page if needed (simple check)
            if text_object.getY() < 1 * inch:
                c.drawText(text_object)
                c.showPage()
                text_object = c.beginText(1 * inch, height - 1 * inch)
                text_object.setFont("Helvetica", 12)
        
        # Draw the text and save the PDF
        c.drawText(text_object)
        c.save()
        
    except ImportError:
        # If reportlab is not available, raise an error
        raise HTTPException(
            status_code=500,
            detail="PDF writing requires reportlab library. Install with: pip install reportlab"
        )

@app.post("/v1/files/read", response_model=FileResponse)
async def read_file(
    request: ReadFileRequest,
    current_user: Dict[str, Any] = Depends(get_current_user),
):
    """
    Read a file from the scratch directory
    Supports: txt, docx, xlsx, pdf, png
    """
    if not FILE_OPS_AVAILABLE:
        raise HTTPException(
            status_code=503,
            detail="File operations not available. Install: pip install python-docx openpyxl PyPDF2 reportlab Pillow"
        )
    # Resolve path with containment and extension checks (blocks path traversal)
    filepath = resolve_scratch_path(request.filename, READ_ALLOWED_EXTENSIONS)
    try:
        # Check if file exists
        if not filepath.exists():
            return FileResponse(
                success=False,
                message=f"File not found: {request.filename}"
            )
        # Enforce max file size before reading
        if filepath.stat().st_size > FILE_OPS_MAX_SIZE_BYTES:
            return FileResponse(
                success=False,
                message="File too large"
            )
        # Determine file extension
        file_ext = filepath.suffix.lower()
        # Read file based on extension
        if file_ext == '.txt':
            content = read_text_file(filepath)
            return FileResponse(
                success=True,
                message=f"Successfully read {request.filename}",
                data={'content': content, 'type': 'text'}
            )
        
        elif file_ext == '.docx':
            content = read_docx_file(filepath)
            return FileResponse(
                success=True,
                message=f"Successfully read {request.filename}",
                data={'content': content, 'type': 'text'}
            )
        
        elif file_ext in ['.xlsx', '.xls']:
            content = read_xlsx_file(filepath)
            return FileResponse(
                success=True,
                message=f"Successfully read {request.filename}",
                data={'content': content, 'type': 'text'}
            )
        
        elif file_ext == '.pdf':
            content = read_pdf_file(filepath)
            return FileResponse(
                success=True,
                message=f"Successfully read {request.filename}",
                data={'content': content, 'type': 'text'}
            )
        
        elif file_ext in ['.png', '.jpg', '.jpeg']:
            image_data = read_png_file(filepath)
            return FileResponse(
                success=True,
                message=f"Successfully read {request.filename}",
                data={'content': image_data['description'], 'type': 'image', 'image_data': image_data}
            )
        
        else:
            # Unsupported file type
            return FileResponse(
                success=False,
                message=f"Unsupported file type: {file_ext}. Supported types: txt, docx, xlsx, pdf, png"
            )
    
    except Exception as e:
        # Handle any errors during file reading
        return FileResponse(
            success=False,
            message=f"Error reading file: {str(e)}"
        )

@app.post("/v1/files/write", response_model=FileResponse)
async def write_file(
    request: WriteFileRequest,
    current_user: Dict[str, Any] = Depends(get_current_user),
):
    """
    Write content to a file in the scratch directory
    Supports: txt, docx, xlsx, pdf
    """
    if not FILE_OPS_AVAILABLE:
        raise HTTPException(
            status_code=503,
            detail="File operations not available. Install: pip install python-docx openpyxl PyPDF2 reportlab Pillow"
        )
    # Enforce max content size before processing
    content_bytes = request.content.encode("utf-8")
    if len(content_bytes) > FILE_OPS_MAX_SIZE_BYTES:
        return FileResponse(success=False, message="Content too large")
    # Build logical filename with extension from format if missing
    logical_name = request.filename.strip()
    if not Path(logical_name).suffix:
        logical_name = f"{logical_name}.{request.format.lower()}"
    # Resolve path with containment and extension checks (blocks path traversal)
    filepath = resolve_scratch_path(logical_name, WRITE_ALLOWED_EXTENSIONS)
    file_ext = filepath.suffix.lower()
    try:
        # Write file based on extension
        if file_ext == '.txt':
            write_text_file(filepath, request.content)
        
        elif file_ext == '.docx':
            write_docx_file(filepath, request.content)
        
        elif file_ext in ['.xlsx', '.xls']:
            write_xlsx_file(filepath, request.content)
        
        elif file_ext == '.pdf':
            write_pdf_file(filepath, request.content)
        
        else:
            # Unsupported file type
            return FileResponse(
                success=False,
                message=f"Unsupported file type for writing: {file_ext}. Supported types: txt, docx, xlsx, pdf"
            )
        
        # Return success response
        return FileResponse(
            success=True,
            message=f"Successfully wrote {filepath.name}",
            data={'filepath': str(filepath), 'size': filepath.stat().st_size}
        )
    
    except Exception as e:
        # Handle any errors during file writing
        return FileResponse(
            success=False,
            message=f"Error writing file: {str(e)}"
        )

@app.get("/v1/files/list")
async def list_files(
    current_user: Dict[str, Any] = Depends(get_current_user),
):
    """List all files in the scratch directory"""
    try:
        # Get all files in scratch directory
        files = []
        for file in SCRATCH_DIR.iterdir():
            if file.is_file():
                files.append({
                    'name': file.name,
                    'size': file.stat().st_size,
                    'modified': file.stat().st_mtime,
                    'extension': file.suffix
                })
        
        # Sort files by modification time (newest first)
        files.sort(key=lambda x: x['modified'], reverse=True)
        
        return {
            'success': True,
            'count': len(files),
            'files': files,
            'scratch_dir': str(SCRATCH_DIR)
        }
    
    except Exception as e:
        # Handle any errors during directory listing
        return {
            'success': False,
            'message': f"Error listing files: {str(e)}"
        }

@app.delete("/v1/files/delete/{filename}")
async def delete_file(
    filename: str,
    current_user: Dict[str, Any] = Depends(get_current_user),
):
    """Delete a file from the scratch directory"""
    # Resolve path with containment and extension checks (blocks path traversal)
    filepath = resolve_scratch_path(filename, READ_ALLOWED_EXTENSIONS)
    try:
        # Check if file exists
        if not filepath.exists():
            return FileResponse(
                success=False,
                message=f"File not found: {filename}"
            )
        
        # Delete the file
        filepath.unlink()
        
        return FileResponse(
            success=True,
            message=f"Successfully deleted {filename}"
        )
    
    except Exception as e:
        # Handle any errors during file deletion
        return FileResponse(
            success=False,
            message=f"Error deleting file: {str(e)}"
        )

# ============================================================================
# END FILE OPERATIONS ENDPOINTS
# ============================================================================

# ============================================================================
# GOOGLE DRIVE UPLOAD ENDPOINT
# ============================================================================

@app.post("/v1/proxy/upload-to-drive")
async def upload_to_drive(request: Request, current_user: Dict[str, Any] = Depends(get_current_user)):
    """
    Upload a file from the scratch directory to Google Drive using service account credentials from .env file.
    filePath must be a filename relative to the scratch directory (e.g. report.docx). Path traversal and absolute paths are rejected.
    Credentials are read from environment variables:
    - GOOGLE_DRIVE_PROJECT_ID
    - GOOGLE_DRIVE_PRIVATE_KEY_ID
    - GOOGLE_DRIVE_PRIVATE_KEY
    - GOOGLE_DRIVE_CLIENT_EMAIL
    - GOOGLE_DRIVE_FOLDER_ID
    """
    folder_id = None
    file_path_obj = None
    try:
        # Get form data from request
        form_data = await request.form()
        # Get file path from form data (must be scratch-relative filename; validated below)
        file_path = form_data.get('filePath')
        if not file_path or not str(file_path).strip():
            raise HTTPException(status_code=400, detail="filePath is required")
        # Resolve to path under SCRATCH_DIR only; rejects absolute, traversal, disallowed extensions
        file_path_obj = resolve_scratch_path(str(file_path).strip(), DRIVE_UPLOAD_EXTENSIONS)
        if not file_path_obj.exists():
            raise HTTPException(status_code=404, detail=f"File not found: {file_path}")
        
        # Get optional file name for Drive
        file_name = form_data.get('fileName')
        
        # Get folder ID from form data or environment variable
        folder_id = form_data.get('folderId') or os.getenv('GOOGLE_DRIVE_FOLDER_ID')
        if not folder_id:
            raise HTTPException(status_code=400, detail="folderId is required (provide in request or set GOOGLE_DRIVE_FOLDER_ID in .env)")
        
        # Read Google Drive credentials from environment variables (loaded from .env file)
        project_id = os.getenv('GOOGLE_DRIVE_PROJECT_ID')
        private_key_id = os.getenv('GOOGLE_DRIVE_PRIVATE_KEY_ID')
        private_key = os.getenv('GOOGLE_DRIVE_PRIVATE_KEY')
        client_email = os.getenv('GOOGLE_DRIVE_CLIENT_EMAIL')
        
        # Validate that all required credentials are present
        if not all([project_id, private_key_id, private_key, client_email]):
            missing = [k for k, v in {
                'GOOGLE_DRIVE_PROJECT_ID': project_id,
                'GOOGLE_DRIVE_PRIVATE_KEY_ID': private_key_id,
                'GOOGLE_DRIVE_PRIVATE_KEY': private_key,
                'GOOGLE_DRIVE_CLIENT_EMAIL': client_email
            }.items() if not v]
            raise HTTPException(
                status_code=500,
                detail=f"Missing Google Drive credentials in .env file: {', '.join(missing)}"
            )
        
        # Construct credentials object from environment variables
        # Replace \\n with actual newlines in private key
        private_key_formatted = private_key.replace('\\n', '\n')
        
        credentials_dict = {
            "type": "service_account",
            "project_id": project_id,
            "private_key_id": private_key_id,
            "private_key": private_key_formatted,
            "client_email": client_email,
            "auth_uri": "https://accounts.google.com/o/oauth2/auth",
            "token_uri": "https://oauth2.googleapis.com/token",
            "auth_provider_x509_cert_url": "https://www.googleapis.com/oauth2/v1/certs",
            "client_x509_cert_url": f"https://www.googleapis.com/robot/v1/metadata/x509/{client_email}"
        }
        
        # Try to import Google Drive API libraries
        try:
            from google.oauth2 import service_account
            from googleapiclient.discovery import build
            from googleapiclient.http import MediaFileUpload
            GOOGLE_DRIVE_AVAILABLE = True
        except ImportError:
            GOOGLE_DRIVE_AVAILABLE = False
            raise HTTPException(
                status_code=503,
                detail="Google Drive API libraries not available. Install with: pip install google-auth google-auth-oauthlib google-auth-httplib2 google-api-python-client"
            )
        
        # Authenticate with Google Drive using service account credentials
        credentials = service_account.Credentials.from_service_account_info(
            credentials_dict,
            scopes=['https://www.googleapis.com/auth/drive.file']
        )
        
        # Build the Drive service
        drive_service = build('drive', 'v3', credentials=credentials)
        
        # Determine file name to use
        upload_file_name = file_name if file_name else file_path_obj.name
        
        # Prepare file metadata
        file_metadata = {
            'name': upload_file_name,
            'parents': [folder_id] if folder_id else []
        }
        
        # Upload file to Google Drive
        media = MediaFileUpload(
            str(file_path_obj),
            resumable=True
        )
        
        # Perform the upload
        file = drive_service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id, name, webViewLink'
        ).execute()
        
        # Audit log: user, filename, folder_id, success, file_id
        user_sub = current_user.get("sub") or "unknown"
        print(f"[AUDIT] upload-to-drive user={user_sub} filename={file_path_obj.name} folder_id={folder_id} success=true file_id={file.get('id')}")
        # Return success response with file ID
        return {
            'success': True,
            'fileId': file.get('id'),
            'fileName': file.get('name'),
            'webViewLink': file.get('webViewLink'),
            'message': f"File successfully uploaded to Google Drive with ID: {file.get('id')}"
        }
        
    except HTTPException as exc:
        # Audit log on auth/path/not-found/credential errors
        user_sub = current_user.get("sub") if current_user else "unknown"
        filename_log = file_path_obj.name if file_path_obj else "n/a"
        print(f"[AUDIT] upload-to-drive user={user_sub} filename={filename_log} folder_id={folder_id or 'n/a'} success=false status={exc.status_code} detail={exc.detail}")
        # Re-raise HTTP exceptions
        raise
    except Exception as e:
        # Audit log on unexpected errors
        user_sub = current_user.get("sub") if current_user else "unknown"
        print(f"[AUDIT] upload-to-drive user={user_sub} folder_id={folder_id or 'n/a'} success=false error={str(e)}")
        # Handle any other errors
        print(f"❌ Google Drive upload error: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to upload file to Google Drive: {str(e)}"
        )

# ============================================================================
# END GOOGLE DRIVE UPLOAD ENDPOINT
# ============================================================================

# ============================================================================
# SSL CERTIFICATE UTILITIES
# ============================================================================

def get_local_ip() -> Optional[str]:
    """
    Get the local IP address of this machine.
    Returns the IP address or None if unable to determine.
    """
    try:
        # Connect to a remote address to determine local IP (doesn't actually send data)
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        ip = s.getsockname()[0]
        s.close()
        return ip
    except Exception:
        return None

def find_mkcert_certificates() -> Tuple[Optional[str], Optional[str]]:
    """
    Find mkcert-generated certificates in project root or certs/ directory.
    Returns a tuple of (cert_file, key_file) or (None, None) if not found.
    mkcert creates files like: anton.local+2.pem, anton.local+2-key.pem, etc.
    """
    # Search in certs/ first, then project root
    search_dirs = [_PROJECT_ROOT / "certs", _PROJECT_ROOT]
    for search_dir in search_dirs:
        if not search_dir.exists():
            continue
        cert_files = list(search_dir.glob("anton.local*.pem"))
        # Filter out key files and find matching pairs
        cert_key_pairs = []
        for cert_path in cert_files:
            if "-key" in cert_path.name:
                continue
            key_path = cert_path.parent / (cert_path.stem + "-key.pem")
            if key_path.exists():
                cert_key_pairs.append((str(cert_path), str(key_path), cert_path.stat().st_mtime))
        if cert_key_pairs:
            cert_key_pairs.sort(key=lambda x: x[2], reverse=True)
            return cert_key_pairs[0][0], cert_key_pairs[0][1]
    return None, None

def get_ssl_certificates() -> Tuple[Optional[str], Optional[str]]:
    """
    Get SSL certificate files for HTTPS server.
    First tries to find mkcert certificates in certs/ or project root.
    Returns a tuple of (cert_file, key_file) or (None, None) if not found.
    """
    # Try to find mkcert certificates first
    cert_file, key_file = find_mkcert_certificates()
    if cert_file and key_file and os.path.exists(cert_file) and os.path.exists(key_file):
        print(f"[SSL] Found mkcert certificate: {cert_file}")
        return cert_file, key_file
    
    # Fall back to default certificate file names in certs/ or project root
    for base in [_PROJECT_ROOT / "certs", _PROJECT_ROOT]:
        default_cert = base / "anton.local+2.pem"
        default_key = base / "anton.local+2-key.pem"
        if default_cert.exists() and default_key.exists():
            print(f"[SSL] Using default certificate: {default_cert}")
            return str(default_cert), str(default_key)
    
    # Return None if no certificates found
    print("[SSL] No SSL certificates found. Server will run without HTTPS.")
    return None, None

# ============================================================================
# END SSL CERTIFICATE UTILITIES
# ============================================================================

if __name__ == "__main__":
    # Start the server
    print("[START] Starting CATBot Proxy Server with File Operations...")
    print(f"[INFO] Scratch directory: {SCRATCH_DIR}")
    
    # Get SSL certificates for HTTPS
    cert_file, key_file = get_ssl_certificates()
    
    # Configure uvicorn with SSL if certificates are available
    if cert_file and key_file:
        print(f"[SSL] Starting HTTPS server on port 8002")
        print(f"[SSL] Certificate: {cert_file}")
        print(f"[SSL] Key: {key_file}")
        uvicorn.run(
            "src.servers.proxy_server:app",
            host="0.0.0.0",
            port=8002,
            reload=True,
            log_level="info",
            ssl_keyfile=key_file,
            ssl_certfile=cert_file
        )
    else:
        print("[WARN] Starting HTTP server (no SSL certificates found)")
        print("[INFO] To enable HTTPS, ensure mkcert certificate files are in certs/ directory")
        uvicorn.run(
            "src.servers.proxy_server:app",
            host="0.0.0.0",
            port=8002,
            reload=True,
            log_level="info"
        )
